from __future__ import annotations

import re
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))

from normalize_phase2_report_symbols import audit, normalize
from formalize_phase2_project_naming import formalize_report


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / (
    "docs/weekly_reports/"
    "CVS-RFFI_Phase2详细复现报告1_qKNN首次出现括号定义版_截至20260817.docx"
)
GROUPED_SOURCE = ROOT / (
    "docs/weekly_reports/"
    "CVS-RFFI_Phase2详细复现报告1_qKNN结果按配置拆分版_截至20260817.docx"
)


def visible_text(paragraph) -> str:
    return "".join(paragraph._p.itertext())


def math_text(paragraph) -> str:
    values = paragraph._p.xpath(
        ".//m:oMath//m:t/text() | .//m:oMathPara//m:t/text()"
    )
    return "".join(values)


def result_numbers(doc: Document) -> list[str]:
    pattern = re.compile(r"[+-]?\d+(?:\.\d+)?(?:%|pp|s)")
    return [
        token
        for table in doc.tables
        for row in table.rows[1:]
        for cell in row.cells
        for token in pattern.findall(cell.text)
    ]


def subscript_pairs(doc: Document) -> set[tuple[str, str]]:
    pairs: set[tuple[str, str]] = set()
    for node in doc.element.xpath(".//m:sSub"):
        base = "".join(
            node.xpath('./*[local-name()="e"]//*[local-name()="t"]/text()')
        )
        sub = "".join(
            node.xpath('./*[local-name()="sub"]//*[local-name()="t"]/text()')
        )
        pairs.add((base, sub))
    return pairs


class NormalizePhase2ReportSymbolsTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.tempdir = tempfile.TemporaryDirectory()
        cls.output = Path(cls.tempdir.name) / "normalized.docx"
        normalize(SOURCE, cls.output)
        cls.source_doc = Document(str(SOURCE))
        cls.output_doc = Document(str(cls.output))

    @classmethod
    def tearDownClass(cls) -> None:
        cls.tempdir.cleanup()

    def test_uses_one_symbol_family_per_concept(self) -> None:
        body = "\n".join(visible_text(p) for p in self.output_doc.paragraphs)
        self.assertIn("分类概率统一记为", body)
        self.assertIn("qKNN名称中的K不表示K-shot", body)
        self.assertNotIn("分别表示适应/注册前后的预测", body)

        domain_math = math_text(self.output_doc.paragraphs[17])
        self.assertIn("src", domain_math)
        self.assertIn("tgt", domain_math)

        metric_math = "\n".join(
            math_text(self.output_doc.paragraphs[index])
            for index in (46, 48, 50, 54, 58, 61, 64, 65)
        )
        self.assertIn("DA0", metric_math)
        self.assertIn("REG0", metric_math)
        self.assertIn("DA1", metric_math)
        self.assertIn("REG1", metric_math)
        self.assertNotIn("pre", metric_math)
        self.assertNotIn("post", metric_math)
        all_math = "\n".join(
            math_text(paragraph)
            for paragraph in self.output_doc.paragraphs
        )
        self.assertNotIn("Cold", all_math)
        self.assertNotIn("Ctnew", all_math)
        self.assertNotIn("Aoldpost", all_math)

        pairs = subscript_pairs(self.output_doc)
        self.assertIn(("P", "src"), pairs)
        self.assertIn(("P", "tgt"), pairs)
        self.assertNotIn(("P", "s"), pairs)
        self.assertNotIn(("P", "t"), pairs)
        self.assertNotIn(("Y", "S"), pairs)
        self.assertNotIn(("Y", "T"), pairs)

        mri_math = "\n".join(
            math_text(self.output_doc.paragraphs[index]) for index in (82, 84, 86, 88)
        )
        self.assertIn("π", mri_math)
        self.assertIn("src", mri_math)
        self.assertIn("tgt", mri_math)
        self.assertIn("ε₀", math_text(self.output_doc.paragraphs[82]))
        domain_math = "\n".join(
            math_text(self.output_doc.paragraphs[index])
            for index in (81, 84, 85, 86, 88, 93, 94, 97, 98, 99, 100, 102)
        )
        self.assertNotIn("srcource", domain_math)
        self.assertNotIn("tgtarget", domain_math)
        self.assertNotIn("srcupport", domain_math)
        self.assertIn("Ksrc,src", math_text(self.output_doc.paragraphs[98]))
        self.assertIn("Ktgt,tgt", math_text(self.output_doc.paragraphs[98]))
        self.assertIn("Ksrc,tgt", math_text(self.output_doc.paragraphs[98]))

        mopc_math = "\n".join(
            math_text(self.output_doc.paragraphs[index])
            for index in (139, 141, 142, 143, 145, 147, 149)
        )
        self.assertIn("Tₜₑₘₚ", mopc_math)
        self.assertIn("ξ", mopc_math)
        self.assertNotIn("τ", mopc_math)
        self.assertIn("Bp", mopc_math)
        self.assertNotIn("Bμ", mopc_math)
        self.assertIn("ℒHR=j=1J", mopc_math)
        self.assertNotIn("ℒHR=g=", mopc_math)

    def test_preserves_results_references_and_document_scale(self) -> None:
        self.assertEqual(len(self.source_doc.tables), len(self.output_doc.tables))
        self.assertEqual(
            sum(len(table.rows) for table in self.source_doc.tables),
            sum(len(table.rows) for table in self.output_doc.tables),
        )
        self.assertEqual(result_numbers(self.source_doc), result_numbers(self.output_doc))
        self.assertEqual(
            [p.text for p in self.source_doc.paragraphs if p.text.startswith("[")],
            [p.text for p in self.output_doc.paragraphs if p.text.startswith("[")],
        )
        result = audit(self.output, SOURCE)
        self.assertEqual(result["tables"], 17)
        self.assertEqual(result["rows"], 153)
        self.assertEqual(result["references"], 5)
        self.assertGreaterEqual(result["omml"], 309)

    def test_audit_accepts_ertb_idr_definition_instead_of_qknn_distinction(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "ertb-idr.docx"
            formalize_report(GROUPED_SOURCE, output)
            result = audit(output, allow_result_regrouping=True)
            self.assertEqual(result["tables"], 52)


if __name__ == "__main__":
    unittest.main()
