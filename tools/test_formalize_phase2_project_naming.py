from __future__ import annotations

import tempfile
import unittest
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from formalize_phase2_project_naming import formalize_report
from normalize_phase2_report_symbols import result_number_tokens


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/weekly_reports/CVS-RFFI_Phase2详细复现报告1_qKNN结果按配置拆分版_截至20260817.docx"

PROJECT_LINE = (
    "项目名称：星地跨域少样本持续射频指纹识别"
    "（Ground-to-Satellite Cross-Domain Few-Shot Continual Radio-Frequency "
    "Fingerprint Identification, CVS-RFFI）"
)
PHASE1_LINE = (
    "Phase1方法：地面跨接收机域泛化射频指纹表征方法"
    "（Ground-Based Cross-Receiver Domain-Generalized Radio-Frequency Fingerprint "
    "Representation Learning）"
)
PHASE2_LINE = (
    "Phase2方法：星载少样本域适应与新类增量注册方法"
    "（Spaceborne Few-Shot Domain Adaptation and New-Class Incremental Registration for RFFI）"
)
ERTB_LINE = (
    "Phase2对比方法：高效稳健任务均衡增量判别注册方法"
    "（Efficient Robust Task-Balanced Incremental Discriminant Registration, "
    "ERTB-IDR；对应D92 E0，而非原始D92）"
)
CORE_LINE = (
    "核心任务：Phase1学习跨接收机稳定表征；Phase2在LEO弱信道下使用少量target support，"
    "同时完成旧类域适应与新类增量注册。"
)


def paragraph_text(paragraph) -> str:
    nodes = paragraph._p.xpath(".//w:t | .//m:t")
    return "".join(node.text or "" for node in nodes)


def all_visible_text(document: Document) -> str:
    values = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.append(paragraph_text(cell.paragraphs[0]))
                for paragraph in cell.paragraphs[1:]:
                    values.append(paragraph_text(paragraph))
    return "\n".join(values)


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph_text(paragraph) == text:
            return paragraph
    raise AssertionError(f"paragraph not found: {text}")


def red_bold_text(paragraph) -> str:
    values = []
    for run in paragraph._p.xpath("./w:r"):
        text = "".join(node.text or "" for node in run.xpath(".//w:t"))
        properties = run.find(qn("w:rPr"))
        if properties is None:
            continue
        bold = properties.find(qn("w:b"))
        color = properties.find(qn("w:color"))
        if bold is not None and color is not None and color.get(qn("w:val")) == "C00000":
            values.append(text)
    return "".join(values)


class FormalizePhase2ProjectNamingTest(unittest.TestCase):
    def test_formal_names_replace_internal_experiment_code(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "formal.docx"
            formalize_report(SOURCE, output)
            result = Document(output)
            body = all_visible_text(result)

            self.assertNotIn("ADV3B02", body)
            self.assertNotIn("Phase1域泛化基座 backbone", body)
            self.assertIn("Phase1域泛化基座", body)
            self.assertIn("星载少样本域适应与新类增量注册方法", body)
            self.assertNotIn("qKNN", body)
            self.assertIn("ERTB-IDR", body)
            self.assertIn("D92 E0_FULL_ONLY", body)
            self.assertIn("关闭原D92的Fisher/Pareto安全门和K折full/block双几何融合", body)
            for line in (PROJECT_LINE, PHASE1_LINE, PHASE2_LINE, ERTB_LINE, CORE_LINE):
                self.assertIn(line, body)

    def test_project_and_phase_lines_are_red_and_bold(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "formal.docx"
            formalize_report(SOURCE, output)
            result = Document(output)
            for line in (PROJECT_LINE, PHASE1_LINE, PHASE2_LINE, ERTB_LINE, CORE_LINE):
                paragraph = find_paragraph(result, line)
                self.assertEqual(red_bold_text(paragraph), line)

            stage2b = next(
                paragraph for paragraph in result.paragraphs
                if paragraph_text(paragraph).startswith("Stage2-B要回答的问题是")
            )
            self.assertIn("Stage2-B要回答的问题是", red_bold_text(stage2b))
            stage2c = next(
                paragraph for paragraph in result.paragraphs
                if paragraph_text(paragraph).startswith("Stage2-C在同一target receiver中注册新类")
            )
            self.assertIn("Stage2-C在同一target receiver中注册新类", red_bold_text(stage2c))

    def test_experiment_content_and_document_structure_are_preserved(self):
        source = Document(SOURCE)
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "formal.docx"
            formalize_report(SOURCE, output)
            result = Document(output)

            self.assertEqual(len(result.tables), 52)
            self.assertEqual(len(result.paragraphs), len(source.paragraphs) + 5)
            self.assertEqual(
                Counter(result_number_tokens(result)),
                Counter(result_number_tokens(source)),
            )
            source_refs = [p.text for p in source.paragraphs if p.text.startswith("[")]
            result_refs = [p.text for p in result.paragraphs if p.text.startswith("[")]
            self.assertEqual(result_refs, source_refs)
            self.assertEqual(
                len(result.element.xpath(".//m:oMath")),
                len(source.element.xpath(".//m:oMath")),
            )

    def test_k20_new3_configuration_starts_on_a_fresh_page(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "formal.docx"
            formalize_report(SOURCE, output)
            result = Document(output)

            caption = next(
                paragraph for paragraph in result.paragraphs
                if paragraph_text(paragraph) == "配置：K=20，新类数=3"
            )
            self.assertTrue(caption.paragraph_format.page_break_before)


if __name__ == "__main__":
    unittest.main()
