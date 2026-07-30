from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PANDOC = Path(r"C:\Users\lh594\.local\bin\pandoc.exe")


FRONT = r"""# 1.核心概念及其与Phase2阶段的对应

本章先区分域适应、少样本学习和类增量学习，再说明它们在CVS-RFFI Phase2中的具体位置。三者描述的不是同一维度：域适应关注数据分布变化，少样本学习关注目标任务的标注数量，类增量学习关注标签空间和模型状态随时间扩展。

## 1.1三个概念的核心区别

**域适应（domain adaptation，DA）：类别语义保持不变，但source域与target域的数据分布不同；模型利用协议允许的target信息降低target风险。**

**少样本学习（few-shot learning，FSL）：目标任务中每个类别只有少量带标签support；模型必须利用已有先验，在独立query上实现泛化。**

**类增量学习（class-incremental learning，CIL）：新类别分阶段加入，模型状态持续保存；每次更新后都要在全部已学习类别中统一预测。**

**少样本类增量学习（few-shot class-incremental learning，FSCIL）：类增量过程中的每个新类只有K-shot support，因而必须同时解决新类欠学习和旧类遗忘。**

**核心区别：DA回答“换了域如何适应”，FSL回答“标注很少如何学习”，CIL回答“类别持续增加如何保持旧知识”。Stage2-C同时包含receiver域偏移、K-shot新类和标签空间扩展，因此属于跨域FSCIL。**

## 1.2RFFI任务的三轴谱系

为避免把“样本少”“域变化”和“类别逐步增加”混成同一问题，表1从三个相互独立的轴定位RFFI任务：类别轴描述训练、测试及增量阶段的类别关系；域轴描述接收机、信道或采集条件是否变化；时间轴描述任务是否顺序到达，以及更新后的模型状态是否持续保存。

**读表原则：K-shot只描述监督预算，不单独决定任务类型。只有类别轴、域轴和时间轴同时确定后，才能判断一个任务属于闭集低样本分类、FSL、域适应、CIL还是跨域FSCIL。**

**表1 RFFI任务在类别轴、域轴和时间轴上的定位**

|任务类型|类别轴|域轴|时间轴|监督预算|冻结后推理类别空间|
|---|---|---|---|---|---|
|低样本闭集RFFI|$\mathcal C_{\mathrm{tr}}=\mathcal C_{\mathrm{te}}$|$d_{\mathrm{tr}}=d_{\mathrm{te}}$|单任务；状态不累积|每个已知类的训练样本较少|固定已知类|
|标准RFFI-FSL|$\mathcal C_B\cap\mathcal C_N=\varnothing$|$d_B=d_N$|episodic；episode间不累积|novel类每类$K$个support|当前episode的novel类|
|广义RFFI-FSL|$\mathcal C_B\cap\mathcal C_N=\varnothing$|$d_B=d_N$|episodic；episode间不累积|novel类每类$K$个support|base类与当前novel类|
|RFFI域适应|$\mathcal C_S=\mathcal C_T$|$d_S\ne d_T$|单次迁移；类别不扩展|target标签可为0、少量或充分|与source相同的发射机集合|
|RFFI少样本域适应|$\mathcal C_S=\mathcal C_T$|$d_S\ne d_T$|单次迁移；类别不扩展|每个已知类$K$个target support|与source相同的发射机集合|
|跨域RFFI-FSL|$\mathcal C_B\cap\mathcal C_N=\varnothing$|$d_B\ne d_T$|episodic；通常不持久化|target novel类每类$K$个support|标准FSL为novel类；广义FSL为base类与novel类|
|RFFI-CIL|$\mathcal C^{(\le t-1)}\subset\mathcal C^{(\le t)}$|通常$d_t=d_{t-1}$|session顺序到达；状态持续保存|新增类标签数量不限定为少样本|全部累计已注册类|
|RFFI-FSCIL|$\mathcal C^{(\le t-1)}\subset\mathcal C^{(\le t)}$|通常$d_t=d_{t-1}$|session顺序到达；状态持续保存|每个增量新类$K$个support|全部累计已注册类|
|RFFI域增量学习（Domain-IL）|$\mathcal C_t=\mathcal C_{t-1}$|$d_t\ne d_{t-1}$|域顺序到达；状态持续保存|新域标签可有可无|固定发射机集合|
|跨域RFFI-FSCIL|$\mathcal C^{(\le t-1)}\subset\mathcal C^{(\le t)}$|$d_t\ne d_{t-1}$|session顺序到达；状态持续保存|target新类每类$K$个support；旧类监督按协议提供|全部累计已注册旧类与新类|
|开放世界跨域RFFI-FSCIL（扩展）|类别增长且存在$\mathcal C_{\mathrm{unknown}}$|域随部署阶段变化|注册与拒识顺序执行；状态持续保存|已注册新类为K-shot；unknown无标签|已注册旧类与新类，并拒识未注册unknown|

三个边界尤其重要：低样本闭集RFFI的训练类和测试类相同，不等于novel-class FSL；episodic FSL的episode彼此独立，不保存不断扩大的标签空间，因此不等于CIL；一次性域适应只处理当前target域，而Domain-IL要求多个域按顺序到达并持续保存模型状态。

**CVS阶段定位：Stage2-A是零标签target迁移参考，只有实际利用无标签target更新时才属于UDA；Stage2-B对应RFFI少样本监督域适应；Stage2-C在同一row中提供旧类与新类K-shot target support，对应跨域单步FSCIL与新类注册，连续执行多个增量session后才构成完整FSCIL；Phase3才扩展到未注册unknown拒识。**

## 1.3RFFI中的类别轴与域轴

RFFI接收信号可抽象为

$$
x=\Psi\!\left(\mathcal R_d\!\left[\mathcal H_d\!\left(\mathcal T_y(s)\right)\right]+n_d\right),
$$

其中，$\mathcal T_y$表示发射机硬件非理想性，决定身份类别；$\mathcal H_d$和$\mathcal R_d$分别表示信道和接收机链路，决定域条件；$\Psi$表示接收后的同步、裁剪、归一化或时频变换。**发射机身份变化属于类别变化，receiver或信道变化属于域变化，两者不能混为同一任务。**

旧类、增量新类和累计已注册类别分别记为

$$
\mathcal C_{\mathrm{old}},\qquad
\mathcal C_t^{\mathrm{new}},\qquad
\mathcal C^{(\le t)}
=\mathcal C_{\mathrm{old}}\cup\bigcup_{i=1}^{t}\mathcal C_i^{\mathrm{new}}.
$$

## 1.4少样本学习的严格任务定义

少样本学习不是泛指“总数据量较小”，而是指模型已经从base数据或预训练状态中获得先验，在目标任务每类只有少量带标签样本时，仍能泛化到未参与适配的query。

一次$N$-way $K$-shot任务记为$\tau=(S_\tau,Q_\tau)$。support包含$N$个类别，每类$K$个独立带标签样本：

$$
S_\tau
=\bigcup_{c\in\mathcal C_\tau}
\left\{(x_{c,k},c)\right\}_{k=1}^{K},
\qquad
|\mathcal C_\tau|=N,
\qquad
|S_\tau|=NK.
$$

学习算法读取先验状态$\Omega_0$和support，生成当前任务预测器：

$$
h_\tau=\mathcal A(S_\tau;\Omega_0),\qquad
\widehat y_q=\arg\max_{c\in\mathcal C_\tau}h_\tau(x_q)_c,
\quad(x_q,y_q)\in Q_\tau.
$$

support与query必须样本级互斥：

$$
S_\tau\cap Q_\tau=\varnothing.
$$

标准novel-class FSL还要求base类与novel类互斥，query通常只在当前novel类中竞争。若训练与测试类别相同，只是每类样本较少，更准确的名称是low-shot closed-set classification。ProtoNet[1]属于度量型FSL方法，但把prototype估计器用于旧类target support，并不会自动把任务变成标准novel-class FSL。

**K-shot统计的是每类K个独立物理IQ记录；同一IQ的FFT、裁剪、均衡或数据增强view均不增加K。query只用于冻结后评价，不能训练、调参、早停、设阈值或回滚。**

## 1.5域适应及Stage2-B定位

source域与target域可写为

$$
\mathcal D_s=(\mathcal X,P_s(X,Y)),\qquad
\mathcal D_t=(\mathcal X,P_t(X,Y)).
$$

跨接收机RFFI保持旧类标签空间不变，但类别条件分布发生变化：

$$
\mathcal C_s=\mathcal C_t=\mathcal C_{\mathrm{old}},
\qquad
P_s(X\mid Y)\neq P_t(X\mid Y).
$$

域适应的目标是利用source知识和允许读取的target数据，降低

$$
R_t(h)=\mathbb E_{(X,Y)\sim P_t}\!\left[\ell(h(X),Y)\right].
$$

Stage2-B每个旧类仅提供K个target-old support，因此属于**少样本监督域适应**：少样本描述target标签预算，域适应才是任务本质。Phase1域泛化训练不能读取未来target receiver；Phase2域适应发生在部署以后，可以读取协议允许的target support。是否接触target数据，是DG与DA的关键边界。

## 1.6类增量、FSCIL与新类注册

类增量学习要求新类按阶段到达，并把更新后的状态持续保存：

$$
\mathcal C^{(\le t)}
=\mathcal C^{(\le t-1)}\cup\mathcal C_t^{\mathrm{new}},
\qquad
\widehat y
=\arg\max_{c\in\mathcal C^{(\le t)}}h_t(x)_c.
$$

推理时不提供task ID，模型必须让旧类与新类在同一标签空间竞争。当每个增量新类只有K个support时，任务成为FSCIL。一次新类集合注册只能评价single-session/one-step类扩展；若要声明持续FSCIL，还需连续执行多个增量session，并在每个session后评价累计类别和遗忘。

CVS中的“新类注册”是部署操作：利用带标签新类support建立prototype、分类权重、adapter或其他持久状态，使该身份从未注册unknown转为seen-new class。应严格区分

$$
\text{unknown rejection}\neq\text{new-class registration}.
$$

unknown尚未获得可信标签，不能直接进入注册集合；只有获得合法标签和support后，才能成为已注册新类。

## 1.7与CVS-RFFI Phase2阶段的对应

|阶段|域是否变化|类别是否增加|可用target标签|准确任务定位|冻结后query范围|
|---|---|---|---|---|---|
|Stage2-A|是|否|无|zero-label transfer/reference；仅在实际使用无标签target更新时才属于UDA|旧类target query|
|Stage2-B|是|否|旧类K-shot|少样本监督域适应|旧类target query|
|Stage2-C|是|是|旧类与新类均为K-shot target support|跨域单步FSCIL/新类注册；连续多session时为完整FSCIL|旧类＋已注册新类统一query|
|Phase3|是|可能|unknown无真值|开集拒识与开放世界扩展|旧类＋已注册新类＋未注册unknown|

**Stage2-A不属于K-shot学习；Stage2-B类别不增加，不是标准novel-class FSL；Stage2-C同时扩大标签空间并要求旧新统一竞争，是少样本类增量任务。**

## 1.8样本角色与成功条件

- **旧类$\mathcal C_{\mathrm{old}}$：**Phase1已经见过的发射机；更换receiver不改变其身份。
- **新类$\mathcal C_t^{\mathrm{new}}$：**Phase1未见、在Stage2-C通过合法support注册的发射机。
- **未注册unknown：**尚未获得可信标签或尚未加入已注册类别集合的发射机。
- **support：**唯一允许参与适应、注册和状态更新的带标签样本。
- **query：**模型冻结后用于评价的独立样本，不得影响predictor。
- **K-shot：**每类K个互不重复的物理接收观测。

"""


REFERENCES = r"""[1] SNELL J, SWERSKY K, ZEMEL R S. Prototypical Networks for Few-shot Learning[C]//Advances in Neural Information Processing Systems 30. 2017. https://papers.nips.cc/paper/6996-prototypical-networks-for-few-shot-learning

[2] YANG L, LI Q, REN X, et al. Mitigating Receiver Impact on Radio Frequency Fingerprint Identification via Domain Adaptation[J]. IEEE Internet of Things Journal, 2024, 11(13):24024-24034. DOI:10.1109/JIOT.2024.3389491.

[3] FENG J, FANG S, FAN Y. Cross-Receiver Radio Frequency Fingerprint Identification Based on Domain Adaptation With Dynamic Distribution Alignment[J]. IEEE Internet of Things Journal, 2025, 12(16):33202-33214. DOI:10.1109/JIOT.2025.3573713.

[4] LIU Y, WANG J, LI J, NIU S, SONG H. Class-Incremental Learning for Wireless Device Identification in IoT[J]. IEEE Internet of Things Journal, 2021, 8(23):17227-17235. DOI:10.1109/JIOT.2021.3078407.

[5] LI D, CHEN Z, SHAO M, et al. Non-Exemplar Class-Incremental Learning via Prototype Correction and Hierarchical Regularization for Specific Emitter Identification[J]. IEEE Transactions on Intelligent Transportation Systems, 2025, 26(8):12632-12646. DOI:10.1109/TITS.2025.3559174."""


def run_pandoc_to_markdown(source: Path, destination: Path) -> None:
    subprocess.run(
        [
            str(PANDOC),
            str(source),
            "-f",
            "docx",
            "-t",
            "gfm+tex_math_dollars",
            "--wrap=none",
            "-o",
            str(destination),
        ],
        check=True,
    )


def build_markdown(converted: str) -> str:
    marker = "# 3.数据、权限与评价框架"
    if marker not in converted:
        raise RuntimeError(f"missing body marker: {marker}")
    tail = marker + converted.split(marker, 1)[1]

    replacements = {
        "# 3.数据、权限与评价框架": "# 2.数据、权限与评价框架",
        "## 3.": "## 2.",
        "# 4.Stage2-B：跨接收机旧类域适应仿真实验": "# 3.Stage2-B：跨接收机旧类域适应仿真实验",
        "## 4.": "## 3.",
        "# 5.Stage2-C：少样本类增量仿真实验": "# 4.Stage2-C：少样本类增量仿真实验",
        "## 5.": "## 4.",
        "# 6.统一比较与下一步": "# 5.统一比较与下一步",
        "## 6.": "## 5.",
        "第5节": "第4节",
        "第4.1节": "第3.1节",
    }
    for old, new in replacements.items():
        tail = tail.replace(old, new)

    citation_map = {"3": "1", "9": "2", "10": "3", "11": "4", "12": "5"}

    def replace_citation(match: re.Match[str]) -> str:
        escaped_number = match.group(1)
        plain_number = match.group(2)
        number = citation_map[escaped_number or plain_number]
        if escaped_number is not None:
            return rf"\[{number}\]"
        return f"[{number}]"

    tail = re.sub(
        r"\\\[(3|9|10|11|12)\\\]|\[(3|9|10|11|12)\]",
        replace_citation,
        tail,
    )

    if "# 参考文献" not in tail or "# 附录A" not in tail:
        raise RuntimeError("reference or appendix marker missing")
    before_refs, after_refs = tail.split("# 参考文献", 1)
    _, appendix = after_refs.split("# 附录A", 1)
    tail = (
        before_refs.rstrip()
        + "\n\n# 参考文献\n\n"
        + REFERENCES
        + "\n\n# 附录A"
        + appendix
    )
    optimized = FRONT.rstrip() + "\n\n" + tail.lstrip()
    return normalize_strong_delimiters(optimized)


def normalize_strong_delimiters(markdown: str) -> str:
    normalized_lines: list[str] = []
    for line in markdown.splitlines():
        parts = line.split("**")
        if len(parts) < 3:
            normalized_lines.append(line)
            continue
        rebuilt = parts[0]
        for index, part in enumerate(parts[1:], start=1):
            rebuilt += "**"
            if index % 2 == 0 and part and not part[0].isspace():
                rebuilt += " "
            rebuilt += part
        normalized_lines.append(rebuilt)
    return "\n".join(normalized_lines) + "\n"


def set_rfonts(r_pr, latin: str = "Times New Roman", east_asia: str = "宋体") -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def style_run(run, size: float | None = None) -> None:
    r_pr = run._element.get_or_add_rPr()
    set_rfonts(r_pr)
    if size is not None:
        run.font.size = Pt(size)


def clear_story(story) -> None:
    paragraphs = list(story.paragraphs)
    if not paragraphs:
        story.add_paragraph()
        return
    paragraphs[0].clear()
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        tr_pr.append(marker)


def set_cell_margins(
    cell,
    *,
    top: int = 70,
    start: int = 90,
    bottom: int = 70,
    end: int = 90,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_twips: list[int]) -> None:
    if len(widths_twips) != len(table.columns):
        raise ValueError("column width count does not match table")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    total_width = sum(widths_twips)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def postprocess_docx(path: Path) -> None:
    doc = Document(path)
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        section.different_first_page_header_footer = False
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            clear_story(story)

    for style in doc.styles:
        if style.type not in {
            WD_STYLE_TYPE.PARAGRAPH,
            WD_STYLE_TYPE.CHARACTER,
            WD_STYLE_TYPE.TABLE,
        }:
            continue
        r_pr = style._element.get_or_add_rPr()
        set_rfonts(r_pr)

    style_specs = {
        "Normal": (11, False, None),
        "Body Text": (11, False, None),
        "Heading 1": (16, True, RGBColor(31, 78, 121)),
        "Heading 2": (13, True, RGBColor(31, 78, 121)),
        "Heading 3": (11.5, True, RGBColor(64, 64, 64)),
    }
    for name, (size, bold, color) in style_specs.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = bold
        if color is not None:
            style.font.color.rgb = color

    key_prefixes = (
        "域适应（domain adaptation，DA）：",
        "少样本学习（few-shot learning，FSL）：",
        "类增量学习（class-incremental learning，CIL）：",
        "少样本类增量学习（few-shot class-incremental learning，FSCIL）：",
        "核心区别：",
        "读表原则：",
        "CVS阶段定位：",
        "Stage2-A不属于K-shot学习；",
    )

    first_heading_seen = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1":
            if first_heading_seen:
                paragraph.paragraph_format.page_break_before = True
            else:
                first_heading_seen = True
        size = 11
        if paragraph.style.name == "Heading 1":
            size = 16
        elif paragraph.style.name == "Heading 2":
            size = 13
        elif paragraph.style.name == "Heading 3":
            size = 11.5
        for run in paragraph.runs:
            style_run(run, size=size)
        if text.startswith("表1 RFFI任务在类别轴、域轴和时间轴上的定位"):
            paragraph.paragraph_format.keep_with_next = True
        if text.startswith(key_prefixes):
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)

    for table_index, table in enumerate(doc.tables):
        table_font_size = 8.2 if table_index == 0 else 9
        for paragraph in iter_table_paragraphs(table):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            for run in paragraph.runs:
                style_run(run, size=table_font_size)
        if table.rows:
            set_repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                set_cell_shading(cell, "1F4E79" if table_index == 0 else "E7E6E6")
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        if table_index == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)
        for row in table.rows:
            set_row_cant_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
        if table_index == 0:
            usable_width = int(
                doc.sections[0].page_width.twips
                - doc.sections[0].left_margin.twips
                - doc.sections[0].right_margin.twips
            )
            fractions = (0.17, 0.16, 0.11, 0.16, 0.18, 0.22)
            widths = [int(usable_width * fraction) for fraction in fractions[:-1]]
            widths.append(usable_width - sum(widths))
            set_table_geometry(table, widths)
            highlighted_tasks = {"RFFI少样本域适应", "跨域RFFI-FSCIL"}
            for row_index, row in enumerate(table.rows[1:], start=1):
                fill = "F7F7F7"
                if 4 <= row_index <= 6:
                    fill = "EAF2F8"
                elif row_index >= 7:
                    fill = "FDF2E9"
                for column_index, cell in enumerate(row.cells):
                    set_cell_shading(cell, fill)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if column_index in {1, 2}
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        if row.cells[0].text.strip() in highlighted_tasks:
                            run.font.color.rgb = RGBColor(192, 0, 0)
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() in {"Stage2-A", "Stage2-B", "Stage2-C"}:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(192, 0, 0)

    doc.core_properties.title = "CVS-RFFI Phase2阶段工作详细报告正文精简优化版"
    doc.core_properties.subject = "Phase2概念、方法与仿真实验"
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-md", type=Path, required=True)
    args = parser.parse_args()

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    args.output_md.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="phase2_body_optimized_") as temp_dir:
        converted = Path(temp_dir) / "source.md"
        run_pandoc_to_markdown(args.input, converted)
        optimized = build_markdown(converted.read_text(encoding="utf-8"))
        args.output_md.write_text(optimized, encoding="utf-8", newline="\n")
        subprocess.run(
            [
                str(PANDOC),
                str(args.output_md),
                "-f",
                "gfm+tex_math_dollars",
                "-o",
                str(args.output_docx),
                "--reference-doc",
                str(args.input),
            ],
            check=True,
        )
    postprocess_docx(args.output_docx)
    print(args.output_docx)


if __name__ == "__main__":
    main()
