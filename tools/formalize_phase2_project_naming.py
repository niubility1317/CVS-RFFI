from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PROJECT_LINE = (
    "项目名称：星地跨域少样本持续射频指纹识别"
    "（Ground-to-Satellite Cross-Domain Few-Shot Continual Radio-Frequency "
    "Fingerprint Identification, CVS-RFFI）"
)
PHASE1_LINE = (
    "Phase1方法：地面跨接收机域泛化射频指纹表征方法"
    "（Ground-Based Cross-Receiver Domain-Generalized Radio-Frequency Fingerprint "
    "Representation Learning）"
)
PHASE2_LINE = (
    "Phase2方法：星载少样本域适应与新类增量注册方法"
    "（Spaceborne Few-Shot Domain Adaptation and New-Class Incremental Registration for RFFI）"
)
ERTB_LINE = (
    "Phase2对比方法：高效稳健任务均衡增量判别注册方法"
    "（Efficient Robust Task-Balanced Incremental Discriminant Registration, "
    "ERTB-IDR；对应D92 E0，而非原始D92）"
)
CORE_LINE = (
    "核心任务：Phase1学习跨接收机稳定表征；Phase2在LEO弱信道下使用少量target support，"
    "同时完成旧类域适应与新类增量注册。"
)
KEY_RED = RGBColor(192, 0, 0)


def all_paragraphs(document: Document) -> Iterable:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def visible_text(paragraph) -> str:
    nodes = paragraph._p.xpath(".//w:t | .//m:t")
    return "".join(node.text or "" for node in nodes)


def set_run_fonts(
    run,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    size: float = 10.5,
) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")


def clear_paragraph_content(paragraph) -> None:
    properties = paragraph._p.find(qn("w:pPr"))
    for child in list(paragraph._p):
        if child is not properties:
            paragraph._p.remove(child)


def rebuild_paragraph(paragraph, segments: list[tuple[str, bool]]) -> None:
    clear_paragraph_content(paragraph)
    for text, emphasized in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        set_run_fonts(
            run,
            bold=emphasized,
            color=KEY_RED if emphasized else None,
        )


def insert_key_line(document: Document, anchor, text: str, *, size: float, keep_next: bool):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_fonts(run, bold=True, color=KEY_RED, size=size)
    anchor.addnext(paragraph._p)
    return paragraph._p


def replace_plain(document: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(document):
        for node in paragraph._p.xpath(".//w:t"):
            if node.text and old in node.text:
                count += node.text.count(old)
                node.text = node.text.replace(old, new)
    return count


def find_body_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if visible_text(paragraph).startswith(prefix):
            return paragraph
    raise RuntimeError(f"paragraph not found: {prefix}")


def formalize_report(source: Path | str, output: Path | str) -> None:
    source_path = Path(source).resolve()
    output_path = Path(output).resolve()
    if source_path == output_path:
        raise ValueError("source and output must differ")
    if not source_path.is_file():
        raise FileNotFoundError(source_path)

    document = Document(str(source_path))

    anchor = document.paragraphs[0]._p
    anchor = insert_key_line(document, anchor, PROJECT_LINE, size=11.5, keep_next=True)
    anchor = insert_key_line(document, anchor, PHASE1_LINE, size=9.5, keep_next=True)
    anchor = insert_key_line(document, anchor, PHASE2_LINE, size=9.5, keep_next=True)
    anchor = insert_key_line(document, anchor, ERTB_LINE, size=9.5, keep_next=True)
    anchor = insert_key_line(document, anchor, CORE_LINE, size=10.5, keep_next=False)
    paragraph_wrapper = next(p for p in document.paragraphs if p._p is anchor)
    paragraph_wrapper.paragraph_format.space_after = Pt(6)

    phase_bridge = find_body_paragraph(document, "两轮实验统一使用ADV3B02")
    rebuild_paragraph(
        phase_bridge,
        [
            (
                "两轮实验统一使用Phase1方法——地面跨接收机域泛化射频指纹表征方法"
                "（以下简称Phase1域泛化基座）形成的冻结checkpoint（训练后冻结的模型参数快照）。"
                "deployment bundle是与checkpoint共同封存、可供Phase2只读使用的部署状态；"
                "backbone指把接收IQ映射为身份特征的主干网络，adapter是附加的小型可训练适配模块，"
                "prototype是某一发射机类别的特征中心。统一基座可避免backbone差异干扰方法比较。",
                False,
            )
        ],
    )

    ertb_intro = find_body_paragraph(document, "qKNN（quantized K-nearest neighbors")
    rebuild_paragraph(
        ertb_intro,
        [
            (
                "ERTB-IDR（Efficient Robust Task-Balanced Incremental Discriminant "
                "Registration，高效稳健任务均衡增量判别注册方法；对应D92 E0_FULL_ONLY，"
                "而非原始D92）是本报告实际采用的项目对比方法。它读取Phase1域泛化基座和"
                "当前row的target support：首先从固定LEO接收IQ提取288维联合特征，再使用"
                "ground-spectrum Cauchy稳健中心减弱接收机/信道扰动；随后分别估计旧类任务"
                "与新类任务的自动收缩协方差，并以固定等权形成任务均衡的共享判别几何；最后"
                "仅执行一次full主几何LDA闭式拟合，编译面对全部已注册类的统一仿射分类头。"
                "ERTB-IDR关闭原D92的Fisher/Pareto安全门和K折full/block双几何融合，不进行"
                "梯度训练，也不保存原始exemplar。Stage2-B的S2B-old表示仅使用旧类target "
                "support构造的DA1_REG0状态；Stage2-C再加入新类support形成DA1_REG1状态。",
                False,
            )
        ],
    )

    replacements = (
        ("ADV3B02", "Phase1域泛化基座"),
        ("同一Phase1域泛化基座地面域泛化checkpoint", "同一Phase1域泛化基座checkpoint"),
        ("Phase1域泛化基座增量更新", "Phase1基座更新"),
        ("Phase1域泛化基座 backbone", "Phase1域泛化基座backbone"),
        ("编码器接口替换为Phase1域泛化基座的160维", "编码器接口接入Phase1域泛化基座输出的160维"),
        ("qKNN", "ERTB-IDR"),
        ("fc_bf_fp→zero-bias Fingerprints基座", "zero-bias Fingerprints基座"),
        (
            "COMPLETED_DIAGNOSTIC_NEGATIVE_NOT_PROMOTABLE",
            "“诊断完成但结果为负，不具备晋级条件”",
        ),
        (
            "DIAGNOSTIC_NEW_CLASS_NO_LEO_NON_FORMAL",
            "“无LEO新类诊断（非正式结果）”",
        ),
    )
    for old, new in replacements:
        replace_plain(document, old, new)

    if "ADV3B02" in "\n".join(visible_text(p) for p in all_paragraphs(document)):
        raise RuntimeError("internal experiment code remains")

    stage2b = find_body_paragraph(document, "Stage2-B要回答的问题是")
    rebuild_paragraph(stage2b, [(visible_text(stage2b), True)])

    stage2c = find_body_paragraph(document, "Stage2-C在同一target receiver中注册新类")
    stage2c_text = visible_text(stage2c)
    stage2c_key = (
        "Stage2-C在同一target receiver中注册新类，并要求旧类与新类在同一输出空间统一竞争，"
        "因此它是跨域FSCIL。"
    )
    if not stage2c_text.startswith(stage2c_key):
        raise RuntimeError("Stage2-C key sentence changed unexpectedly")
    rebuild_paragraph(
        stage2c,
        [(stage2c_key, True), (stage2c_text[len(stage2c_key) :], False)],
    )

    conclusion = find_body_paragraph(document, "ERTB-IDR在五个共同切片上的旧新调和均值")
    conclusion_text = visible_text(conclusion)
    conclusion_key = "正式结论保持“诊断完成但结果为负，不具备晋级条件”，不能表述为ERTB-IDR已完成晋级。"
    if conclusion_key not in conclusion_text:
        raise RuntimeError("qKNN conclusion sentence not found")
    before, after = conclusion_text.split(conclusion_key, 1)
    rebuild_paragraph(
        conclusion,
        [(before, False), (conclusion_key, True), (after, False)],
    )

    # Keep the K=20/new=3 result block away from the preceding page boundary.
    # Native Word otherwise places part of this caption above the top margin.
    k20_new3_caption = find_body_paragraph(document, "配置：K=20，新类数=3")
    k20_new3_caption.paragraph_format.page_break_before = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Formalize Phase2 project naming and key emphasis")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    formalize_report(args.source, args.output)


if __name__ == "__main__":
    main()
