from __future__ import annotations

import argparse
import tempfile
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image, ImageDraw, ImageFont

from revise_phase2_report_numbering_and_results import (
    HEADING_BLUE,
    TEXT_RED,
    add_result_table_before,
    append_paragraph_before,
    clean_text,
    find_paragraph,
    normalize_all_visible_run_fonts,
    remove_between,
    replace_paragraph_text,
    set_repeat_table_header,
    set_row_cant_split,
    style_run,
)


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


STAGE2C_FORMAL = [
    ["K", "新类数", "方法", "注册后旧类A_old", "新类A_new", "联合H_old,new", "遗忘F_old"],
    ["5", "1", "CSIL", "83.70%", "0.00%", "0.00%", "0.00%"],
    ["5", "1", "MoPC-HR", "87.47%", "0.00%", "0.00%", "0.00%"],
    ["5", "3", "CSIL", "83.70%", "0.00%", "0.00%", "0.00%"],
    ["5", "3", "MoPC-HR", "87.47%", "0.02%", "0.04%", "0.00%"],
    ["5", "5", "MoPC-HR", "77.44%", "25.65%", "37.97%", "10.02%"],
    ["5", "10", "MoPC-HR", "63.80%", "22.08%", "32.30%", "23.67%"],
    ["5", "20", "CSIL", "0.78%", "5.17%", "0.54%", "82.92%"],
    ["5", "25", "MoPC-HR", "58.73%", "14.82%", "23.43%", "28.73%"],
    ["10", "1", "CSIL", "83.70%", "0.00%", "0.00%", "0.00%"],
    ["10", "1", "MoPC-HR", "87.47%", "0.00%", "0.00%", "0.00%"],
    ["10", "3", "CSIL", "83.70%", "0.00%", "0.00%", "0.00%"],
    ["10", "3", "MoPC-HR", "66.78%", "50.40%", "56.78%", "20.69%"],
    ["10", "5", "MoPC-HR", "47.38%", "50.52%", "48.05%", "40.09%"],
    ["10", "10", "MoPC-HR", "44.69%", "33.15%", "37.57%", "42.78%"],
    ["10", "20", "CSIL", "8.33%", "5.49%", "3.17%", "75.37%"],
    ["10", "25", "MoPC-HR", "44.99%", "22.31%", "29.32%", "42.48%"],
    ["20", "1", "CSIL", "83.70%", "0.00%", "0.00%", "0.00%"],
    ["20", "1", "MoPC-HR", "60.76%", "96.53%", "72.69%", "26.71%"],
    ["20", "3", "CSIL", "0.08%", "34.04%", "0.14%", "83.62%"],
    ["20", "3", "MoPC-HR", "48.96%", "71.47%", "57.23%", "38.51%"],
    ["20", "5", "MoPC-HR", "33.28%", "61.31%", "41.09%", "54.19%"],
    ["20", "10", "MoPC-HR", "34.07%", "41.36%", "35.48%", "53.40%"],
    ["20", "20", "CSIL", "59.84%", "5.64%", "9.96%", "23.86%"],
    ["20", "25", "MoPC-HR", "36.66%", "27.86%", "30.23%", "50.81%"],
]


COMMON_SLICES = [
    ["共同切片", "方法", "初始A_old", "注册后A_old", "A_new", "H_old,new", "F_old"],
    ["K1/new20", "ERTB-IDR", "68.144%", "44.033%", "27.150%", "33.410%", "24.111pp"],
    ["K1/new20", "CSIL官方流程", "42.833%", "42.833%", "0.000%", "0.000%", "0.000pp"],
    ["K1/new20", "MoPC-HR官方流程", "45.322%", "40.722%", "1.363%", "2.603%", "4.600pp"],
    ["K5/new20", "ERTB-IDR", "81.267%", "63.711%", "58.883%", "60.955%", "17.556pp"],
    ["K5/new20", "CSIL官方流程", "42.833%", "0.200%", "5.557%", "0.316%", "42.633pp"],
    ["K5/new20", "MoPC-HR官方流程", "45.322%", "13.511%", "17.433%", "14.309%", "31.811pp"],
    ["K10/new5", "ERTB-IDR", "86.111%", "76.189%", "74.133%", "74.803%", "9.922pp"],
    ["K10/new5", "CSIL官方流程", "42.833%", "0.689%", "20.413%", "1.264%", "42.144pp"],
    ["K10/new5", "MoPC-HR官方流程", "45.322%", "9.322%", "49.547%", "14.947%", "36.000pp"],
    ["K10/new10", "ERTB-IDR", "86.111%", "72.533%", "66.353%", "69.106%", "13.578pp"],
    ["K10/new10", "CSIL官方流程", "42.833%", "0.000%", "10.460%", "0.000%", "42.833pp"],
    ["K10/new10", "MoPC-HR官方流程", "45.322%", "9.500%", "32.900%", "13.770%", "35.822pp"],
    ["K10/new20", "ERTB-IDR", "86.111%", "71.333%", "68.150%", "69.555%", "14.778pp"],
    ["K10/new20", "CSIL官方流程", "42.833%", "38.222%", "1.660%", "2.979%", "4.611pp"],
    ["K10/new20", "MoPC-HR官方流程", "45.322%", "7.611%", "25.187%", "10.695%", "37.711pp"],
]


SCENARIO_TABLE = [
    ["对象", "场景定义", "CVS研究任务"],
    ["NTN终端／UE", "经过运营登记和网络认证、直接向卫星发送上行信号的具体设备。", "把具体物理发射RF链作为身份类别，而不是把账号或协议地址直接当作射频类别。"],
    ["目标星载接收域", "由卫星接收前端、采样链、轨道几何、LEO弱信道和时间状态共同形成的目标分布。", "Stage2-B用已登记旧终端的K-shot support校准接收域；Stage2-C在同一目标域注册新终端。"],
    ["运营方／核心网", "提供订阅认证、设备登记、授权窗口和审计记录。", "为support提供可信逻辑身份和采集授权，但不替代物理RF链判定。"],
    ["星载或网关CVS", "读取目标接收域I/Q，运行冻结表征、轻量适应、注册和独立query推理。", "输出已注册身份、物理链不一致或defer，并保存可回滚的版本化状态。"],
]


REQUIREMENT_TABLE = [
    ["要求", "最低条件", "不满足时的处理"],
    ["目标域可观测", "能够从目标卫星接收链读取固定数字I/Q，并记录receiver ID、时间、频率和链路质量。", "透明转发载荷或无数字I/Q接口时，转到可信网关处理，不能声称星上适应。"],
    ["可信support", "每类K个独立物理发送事件，来自认证、授权且带challenge／nonce的registration episode。", "标签来源、独立性或质量门失败时拒绝注册；历史query不得转作support。"],
    ["算法权限", "更新只读取冻结bundle和合法support；query零更新，并对全部已注册类逐样本统一竞争。", "禁止用query真值、old／new角色、类别配额、跨query重排或测试反馈调参。"],
    ["资源与可靠性", "报告模型大小、峰值RAM、状态字节、推理／更新时间、能耗、原子提交和回滚点。", "超出星载预算时采用ground-assisted方案；不得用资源不可行的方法支撑星上部署结论。"],
    ["安全与审计", "模型、注册凭证和状态更新可签名、可追溯、带有效期，并能抵抗重放和错误标签。", "发现冲突时保持旧状态、输出defer并转入运营复核。"],
    ["证据等级", "地面代理、硬件在环和受控在轨数据分层报告，明确接收链与信道来源。", "当前地面数据与LEO压力代理不能表述为真实Iridium在轨验证。"],
]


ENROLLMENT_DELTA_TABLE = [
    ["环节", "常规NTN已有流程", "CVS／RFFI新增流程"],
    ["身份前提", "5G-AKA／EAP-AKA′等网络认证建立可信logical identity和安全上下文。", "不重复主认证；在既有认证结果之后申请一次Physical-RF Enrollment Episode。"],
    ["事件授权", "按正常接入和业务策略授权通信。", "新增registration_event_id、receiver_id、registration_window、K、challenge／nonce、bundle版本和授权签名。"],
    ["信号采集", "终端发送正常接入或业务流量。", "在受控窗口发送K个独立registration burst；目标接收链冻结I/Q并执行SNR、AGC、clipping、同步和独立性质量门。"],
    ["模型操作", "建立网络安全上下文，不生成物理RF类别。", "旧类执行Stage2-B目标域校准；新物理RF链执行Stage2-C few-shot enrollment。"],
    ["状态提交", "保存订阅、会话和安全上下文。", "新增physical_RF_chain_id、support_event_id[]、prototype／model delta、new_model_hash、原子提交和回滚指针。"],
    ["正常运行", "承载正常NTN业务。", "RFFI inference only；普通业务query不能自动成为support，也不能更新prototype、adapter或threshold。"],
]


RESOURCE_OVERHEAD_TABLE = [
    ["新增开销", "主要来源", "必须报告的计量", "控制方式"],
    ["空口与终端能耗", "授权交互、registration window和K个受控上行burst。", "信令字节、burst时长、K、终端发射能量和重试次数。", "事件触发而非持续采集；复用既有认证结果；质量门通过后停止。"],
    ["星载计算", "K次特征提取、support质量检查、Stage2-B校准或Stage2-C注册。", "MACs／FLOPs、峰值RAM、更新时间、能耗和最坏时延。", "冻结主干，优先闭式头、prototype或小型adapter；超预算时转ground-assisted。"],
    ["持久状态", "新类prototype／分类头、receiver-specific delta、版本与回滚副本。", "bundle字节、每类新增字节、接收域状态字节和双版本峰值。", "量化、低秩／紧凑状态、原子提交后释放旧临时状态。"],
    ["星地通信", "授权、审计元数据、模型delta或特征回传。", "上／下行字节、峰值带宽、断链重试和完成时延。", "避免持续回传原始I/Q；只传必要特征、凭证、delta和摘要。"],
    ["审计与安全", "support ID、签名、哈希、有效期和操作日志。", "日志字节、签名验证时延和保留周期。", "元数据最小化、批量签名和明确的数据保留策略。"],
]


def accept_tracked_changes(source: Path, output: Path) -> None:
    """Accept visible insertions/moves and remove deletions before local editing."""
    with ZipFile(source, "r") as archive_in, ZipFile(output, "w", ZIP_DEFLATED) as archive_out:
        for item in archive_in.infolist():
            data = archive_in.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    archive_out.writestr(item, data)
                    continue
                for local_name in ("del", "moveFrom"):
                    for node in list(root.xpath(f".//w:{local_name}", namespaces={"w": WORD_NS})):
                        parent = node.getparent()
                        if parent is not None:
                            parent.remove(node)
                for local_name in ("ins", "moveTo"):
                    for node in list(root.xpath(f".//w:{local_name}", namespaces={"w": WORD_NS})):
                        parent = node.getparent()
                        if parent is None:
                            continue
                        position = parent.index(node)
                        for child in list(node):
                            parent.insert(position, child)
                            position += 1
                        parent.remove(node)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            archive_out.writestr(item, data)


def get_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def heat_color(value: float | None, maximum: float = 80.0) -> tuple[int, int, int]:
    if value is None:
        return (235, 238, 242)
    ratio = max(0.0, min(1.0, value / maximum))
    start = (250, 235, 232)
    middle = (255, 246, 204)
    end = (91, 155, 213)
    if ratio < 0.5:
        local = ratio / 0.5
        return tuple(round(start[i] + (middle[i] - start[i]) * local) for i in range(3))
    local = (ratio - 0.5) / 0.5
    return tuple(round(middle[i] + (end[i] - middle[i]) * local) for i in range(3))


def draw_matrix(
    draw: ImageDraw.ImageDraw,
    *,
    x: int,
    y: int,
    width: int,
    row_labels: list[str],
    column_labels: list[str],
    values: list[list[float | None]],
    title: str,
) -> int:
    title_font = get_font(30, bold=True)
    label_font = get_font(23, bold=True)
    value_font = get_font(24, bold=True)
    note_font = get_font(19)
    row_label_width = 135
    header_height = 58
    row_height = 86
    grid_width = width - row_label_width
    cell_width = grid_width / len(column_labels)
    draw.text((x, y), title, font=title_font, fill=(31, 78, 121))
    y += 48
    for column_index, label in enumerate(column_labels):
        cx = x + row_label_width + column_index * cell_width
        draw.rectangle((cx, y, cx + cell_width, y + header_height), fill=(231, 238, 247), outline=(130, 145, 160), width=2)
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((cx + (cell_width - (bbox[2] - bbox[0])) / 2, y + 14), label, font=label_font, fill=(31, 78, 121))
    draw.rectangle((x, y, x + row_label_width, y + header_height), fill=(231, 238, 247), outline=(130, 145, 160), width=2)
    draw.text((x + 34, y + 14), "K", font=label_font, fill=(31, 78, 121))
    y += header_height
    for row_index, row_label in enumerate(row_labels):
        cy = y + row_index * row_height
        draw.rectangle((x, cy, x + row_label_width, cy + row_height), fill=(244, 248, 252), outline=(130, 145, 160), width=2)
        label_box = draw.textbbox((0, 0), row_label, font=label_font)
        draw.text((x + (row_label_width - (label_box[2] - label_box[0])) / 2, cy + 28), row_label, font=label_font, fill=(31, 78, 121))
        for column_index, value in enumerate(values[row_index]):
            cx = x + row_label_width + column_index * cell_width
            draw.rectangle((cx, cy, cx + cell_width, cy + row_height), fill=heat_color(value), outline=(130, 145, 160), width=2)
            text = "未评估" if value is None else f"{value:.2f}%"
            font = note_font if value is None else value_font
            box = draw.textbbox((0, 0), text, font=font)
            draw.text((cx + (cell_width - (box[2] - box[0])) / 2, cy + 27), text, font=font, fill=(45, 52, 60))
    return y + len(row_labels) * row_height


def make_stage2c_heatmap(output: Path) -> None:
    image = Image.new("RGB", (1700, 920), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    note_font = get_font(21)
    draw.text((70, 28), "Stage2-C正式配置：旧／新类联合调和均值H_old,new", font=title_font, fill=(31, 78, 121))
    columns = ["new=1", "new=3", "new=5", "new=10", "new=20", "new=25"]
    rows = ["5", "10", "20"]
    csil = [
        [0.00, 0.00, None, None, 0.54, None],
        [0.00, 0.00, None, None, 3.17, None],
        [0.00, 0.14, None, None, 9.96, None],
    ]
    mopc = [
        [0.00, 0.04, 37.97, 32.30, None, 23.43],
        [0.00, 56.78, 48.05, 37.57, None, 29.32],
        [72.69, 57.23, 41.09, 35.48, None, 30.23],
    ]
    bottom = draw_matrix(draw, x=70, y=105, width=1560, row_labels=rows, column_labels=columns, values=csil, title="CSIL")
    bottom = draw_matrix(draw, x=70, y=bottom + 48, width=1560, row_labels=rows, column_labels=columns, values=mopc, title="MoPC-HR")
    draw.text((70, bottom + 24), "注：空白配置表示该方法未在该new-class规模上运行，不等于0；颜色越深表示联合性能越高。", font=note_font, fill=(80, 88, 98))
    image.save(output, dpi=(220, 220), optimize=True)


def make_common_heatmap(output: Path) -> None:
    image = Image.new("RGB", (1700, 650), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    note_font = get_font(21)
    draw.text((70, 28), "共同LEO切片：H_old,new方法对照", font=title_font, fill=(31, 78, 121))
    columns = ["K1/new20", "K5/new20", "K10/new5", "K10/new10", "K10/new20"]
    rows = ["ERTB-IDR", "CSIL", "MoPC-HR"]
    values = [
        [33.410, 60.955, 74.803, 69.106, 69.555],
        [0.000, 0.316, 1.264, 0.000, 2.979],
        [2.603, 14.309, 14.947, 13.770, 10.695],
    ]
    bottom = draw_matrix(draw, x=70, y=105, width=1560, row_labels=rows, column_labels=columns, values=values, title="三场景平均结果")
    draw.text((70, bottom + 24), "注：两组实验的seed集合并非严格一一配对，本图只作描述性比较，不表示paired显著性。", font=note_font, fill=(80, 88, 98))
    image.save(output, dpi=(220, 220), optimize=True)


def add_picture_before(doc, anchor, image_path: Path, *, width: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    anchor._p.addprevious(paragraph._p)
    caption_paragraph = append_paragraph_before(doc, anchor, caption, style="Body Text")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    for run in caption_paragraph.runs:
        style_run(run, size=9.0, bold=True, color=HEADING_BLUE)


def add_reference(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    style_run(run, size=9.0)


def update_fields_on_open(doc) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def revise(source: Path, output: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="cvs_phase2_opt_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        accepted_source = temp_dir / "accepted_source.docx"
        formal_heatmap = temp_dir / "stage2c_formal_heatmap.png"
        common_heatmap = temp_dir / "stage2c_common_heatmap.png"
        accept_tracked_changes(source, accepted_source)
        make_stage2c_heatmap(formal_heatmap)
        make_common_heatmap(common_heatmap)

        doc = Document(accepted_source)
        original_image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
        if len(original_image_paragraphs) != 1:
            raise RuntimeError(f"expected one original scenario figure, found {len(original_image_paragraphs)}")
        scenario_figure_xml = deepcopy(original_image_paragraphs[0]._p)

        formal_heading = find_paragraph(doc, "4.5.1按K-shot与新类规模拆分的正式结果")
        phenomenon_heading = find_paragraph(doc, "4.5.2主要现象")
        common_heading = find_paragraph(doc, "4.5.3ERTB-IDR与类增量方法的共同LEO切片对照")
        chapter5_heading = find_paragraph(doc, "5.CVS项目场景说明")
        chapter6_heading = find_paragraph(doc, "6.参考文献")

        remove_between(formal_heading, phenomenon_heading)
        replace_paragraph_text(formal_heading, "4.5.1正式配置矩阵与联合性能图")
        append_paragraph_before(
            doc,
            phenomenon_heading,
            (
                "原报告按K-shot和新类规模拆成18个配置段。为保留完整数值并减少重复，本节将24个方法结果"
                "合并到同一张配置矩阵。A_old、A_new、H_old,new和F_old沿用2.3节定义；所有query均在模型"
                "冻结后评价。CSIL与MoPC-HR的冻结new-class规模不同，未运行的组合不补零、不插值。"
            ),
            style="Body Text",
        )
        add_result_table_before(
            doc,
            phenomenon_heading,
            STAGE2C_FORMAL,
            fractions=[0.06, 0.09, 0.17, 0.19, 0.16, 0.18, 0.15],
            font_size=7.7,
            left_columns=(2,),
            group_column=0,
        )
        add_picture_before(
            doc,
            phenomenon_heading,
            formal_heatmap,
            width=6.25,
            caption="图1 Stage2-C正式配置的旧／新类联合调和均值（H_old,new）",
        )

        remove_between(phenomenon_heading, common_heading)
        append_paragraph_before(
            doc,
            common_heading,
            (
                "零注册集中在低K和小规模新类切片：CSIL在K=5／10且new=1或3时A_new为0，"
                "MoPC-HR在K=5／10、new=1时同样未形成有效新类边界。此时较高A_old不能被解释为联合成功。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            common_heading,
            (
                "MoPC-HR在K=20／new=1取得本矩阵最高H_old,new=72.69%，但A_old降至60.76%，"
                "F_old达到26.71pp；它体现较强可塑性，而不是无代价注册。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            common_heading,
            (
                "新类规模扩大后联合性能明显下降。MoPC-HR在K=20时，new从1增至25，A_new由96.53%"
                "降至27.86%，A_old由60.76%降至36.66%，H_old,new由72.69%降至30.23%。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            common_heading,
            (
                "CSIL呈现更尖锐的稳定性—可塑性冲突：部分切片几乎完全保留旧类却无法注册新类，"
                "另一些切片开始学习新类时旧类大幅下降。结论必须以同一row中的A_old、A_new、H_old,new"
                "和F_old共同判读，不能分别摘取边际最大值。"
            ),
            style="Body Text",
        )

        remove_between(common_heading, chapter5_heading)
        append_paragraph_before(
            doc,
            chapter5_heading,
            (
                "ERTB-IDR完成125／125个任务和375／375个LEO场景。下表把原来的5张共同切片表合并为"
                "一张15行对照表。CSIL和MoPC-HR使用seed 713101–713105，ERTB-IDR使用"
                "713102–713106；seed集合并非严格配对，因此只作描述性比较。"
            ),
            style="Body Text",
        )
        add_result_table_before(
            doc,
            chapter5_heading,
            COMMON_SLICES,
            fractions=[0.14, 0.20, 0.14, 0.14, 0.13, 0.13, 0.12],
            font_size=7.6,
            left_columns=(0, 1),
            group_column=0,
        )
        add_picture_before(
            doc,
            chapter5_heading,
            common_heatmap,
            width=6.25,
            caption="图2 ERTB-IDR、CSIL与MoPC-HR在共同LEO切片上的联合性能对照",
        )
        append_paragraph_before(
            doc,
            chapter5_heading,
            (
                "共同切片中，ERTB-IDR在K5／new20和三个K10切片上取得更高的H_old,new，同时保持较低"
                "遗忘；K1／new20仍只有33.410%，说明极低support预算下的目标域校准与新类注册仍未解决。"
                "该结论只适用于表内共同LEO切片，不外推为对所有FSCIL设置的普遍优势。"
            ),
            style="Body Text",
        )

        remove_between(chapter5_heading, chapter6_heading)
        replace_paragraph_text(chapter5_heading, "5.CVS项目场景研究")

        append_paragraph_before(doc, chapter6_heading, "5.1场景说明", style="Heading 2", keep_with_next=True)
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "CVS-RFFI面向3GPP非地面网络（Non-Terrestrial Network，NTN）中的地面终端上行识别。"
                "3GPP把NTN定义为利用卫星或高空平台承载中继节点或基站功能的网络或网络片段；Release 17"
                "首次形成规范性NTN支持，并把卫星接入纳入PLMN、核心网、移动性和业务连续性体系[6]。"
                "在此架构中，终端的逻辑身份由订阅和网络认证管理，而CVS识别的对象是终端内部具体的物理"
                "发射RF链。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "Iridium NTN Direct提供了具体的LEO直连终端背景：Iridium公开说明其服务依托66颗LEO"
                "卫星、全球协调的L频段和软件定义升级，面向3GPP Release 19的NB-IoT与Direct-to-Device"
                "接入[7]；2026年公开进展已包括双向在轨消息测试[8]。这些公开事实说明标准化、低功耗、大规模"
                "终端直接接入LEO星座的生态正在形成，但不表示Iridium已经部署RFFI或开放原始I/Q。"
            ),
            style="Body Text",
        )
        add_result_table_before(
            doc,
            chapter6_heading,
            SCENARIO_TABLE,
            fractions=[0.18, 0.38, 0.44],
            font_size=8.2,
            left_columns=(0, 1, 2),
            group_column=0,
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "研究顺序与运营顺序必须区分。研究上，Phase1先在地面source receiver学习稳定表征；部署后，"
                "Stage2-B用已登记旧终端的target-old support校准新的星载receiver domain，Stage2-C再用"
                "Y_old∪Y_new的合法K-shot support完成旧类适应与新类注册。运营上，未获得可信标签的信号不能"
                "直接成为Y_new；只有完成外部确权、registration_authorized=true并重新采集fresh独立support后，"
                "才能进入Stage2-C。"
            ),
            style="Body Text",
        )

        append_paragraph_before(doc, chapter6_heading, "5.2天基射频指纹识别的意义", style="Heading 2", keep_with_next=True)
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "网络认证回答“终端是否持有合法凭据”，射频指纹识别回答“当前无线信号是否来自登记的那一套"
                "物理RF链”。功率放大器非线性、I/Q不平衡、本振偏差、相位噪声和DAC误差会在发射波形中"
                "留下设备相关特征。当天基接收系统能够稳定提取这些特征时，运营方获得一条独立于账号和协议"
                "地址的物理层证据，用于发现凭据—硬件不一致、RF模组更换未登记或跨卫星观测身份不连续。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "对CVS而言，科学价值不只是提高一次闭集分类准确率，而是同时解决两个变化：同一发射机进入"
                "未见星载接收域时，身份应跨receiver保持；新终端获得授权后，又必须以少量样本加入统一类别空间。"
                "因此评价必须同一row报告旧类准确率、新类准确率、调和均值、遗忘和逐类floor，避免用“只保旧类”"
                "或“只学新类”的单侧结果代替注册成功。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "该能力是补充证据，不替代5G／NTN密码学认证、协议完整性、运营登记或人工处置。低质量、"
                "接收域超出适用范围或类别边界冲突时，系统应输出defer并保留审计信息，而不是强制给出身份。"
            ),
            style="Body Text",
        )

        append_paragraph_before(doc, chapter6_heading, "5.3NTN注册", style="Heading 2", keep_with_next=True)
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "NTN注册不是从普通业务流量中自动抓取若干query后自训练，而是运营方主动创建的受控物理RF"
                "注册事件。终端先完成正常网络认证；运营方核对逻辑身份、设备记录和授权范围，再下发带时间窗、"
                "频率、波束、receiver ID和challenge／nonce的registration episode。目标接收链采集每类K个"
                "独立物理发送事件，经过SNR、AGC、clipping、同步和重复样本检查后，才能成为support。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "已登记旧终端的support用于Stage2-B接收域校准；新授权终端或实质更换的RF模组使用fresh"
                " K-shot进入Stage2-C。适应和注册完成后，系统生成带版本、有效期和回滚点的状态，并冻结面向"
                "全部已注册类别的预测器。后续query只执行独立推理，不更新backbone、prototype、adapter、"
                "阈值或选择规则。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "CVS不是第二套NTN认证协议，也不是让卫星在普通业务期间持续在线学习。它复用既有认证形成的"
                "可信logical identity，只在明确授权的registration episode内执行一次事件触发式少样本校准或"
                "注册；事件结束后提交状态，正常业务阶段恢复为query zero-update的只推理模式。"
            ),
            style="Body Text",
        )
        enrollment_delta_table = add_result_table_before(
            doc,
            chapter6_heading,
            ENROLLMENT_DELTA_TABLE,
            fractions=[0.16, 0.34, 0.50],
            font_size=7.9,
            left_columns=(0, 1, 2),
            group_column=0,
        )
        for row in enrollment_delta_table.rows[1:]:
            for paragraph in row.cells[2].paragraphs:
                for run in paragraph.runs:
                    style_run(run, size=7.9, bold=True, color=TEXT_RED)
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "与常规NTN注册相比，RFFI的资源开销不是又一次大规模训练，而是一次有界事件的增量成本。"
                "其大小由K、特征提取器、更新头、每类持久状态和是否回传原始I/Q共同决定，必须与识别收益同表报告。"
            ),
            style="Body Text",
        )
        add_result_table_before(
            doc,
            chapter6_heading,
            RESOURCE_OVERHEAD_TABLE,
            fractions=[0.15, 0.29, 0.28, 0.28],
            font_size=7.7,
            left_columns=(0, 1, 2, 3),
            group_column=0,
        )
        chapter6_heading._p.addprevious(deepcopy(scenario_figure_xml))
        registration_caption = append_paragraph_before(
            doc,
            chapter6_heading,
            "图3 受控NTN终端物理RF链注册流程",
            style="Body Text",
        )
        registration_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        registration_caption.paragraph_format.space_before = Pt(0)
        registration_caption.paragraph_format.space_after = Pt(8)
        for run in registration_caption.runs:
            style_run(run, size=9.0, bold=True, color=HEADING_BLUE)
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "注册绑定必须带有效期。只修改账号或业务配置时，物理RF链通常不变，可更新逻辑别名；更换PA、"
                "本振或整套RF模组时，应建立新的physical_RF_chain_id并重新授权采集。历史unknown query保留为"
                "不可变检测证据，不能追溯改成新类support。"
            ),
            style="Body Text",
        )
        heading_54 = append_paragraph_before(
            doc,
            chapter6_heading,
            "5.4为什么要在星载目标域进行预适应／类注册，以及实施要求",
            style="Heading 2",
            keep_with_next=True,
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "所谓“在天上进行”，首先意味着适应和注册必须以目标卫星实际接收链产生的I/Q为数据源，而非"
                "只在地面source数据上预先假设目标分布。发射信号经过LEO传播后，还要通过星载LNA、滤波器、"
                "混频器、AGC、ADC和数字前端；这些receiver-specific效应在地面训练阶段不可完全观测。"
                "Stage2-B因此需要目标域旧类support估计接收域偏移，Stage2-C则在该目标域几何中加入新类边界。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "第一，只有目标接收链能够直接暴露该卫星前端的残余失真。地面多接收机域泛化可以提供稳健初始化，"
                "但不能穷尽未来星载本振、采样时钟、增益链、温度和数字前端组合。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "第二，星上或近星处理具有时效性。LEO可见窗口有限、卫星切换频繁，馈电链路也可能间歇；若每次都"
                "等待原始I/Q下传和地面回传状态，注册完成时间可能超过当前过境窗口。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "第三，原始I/Q的数据率远高于prototype、特征或模型delta。对受控K-shot在星上完成质量门和特征"
                "压缩，可以降低回传带宽与原始无线数据的集中存储需求。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "第四，接收域状态具有局部性。不同卫星或更换后的载荷可能需要不同的轻量校准状态；把所有星载"
                "receiver强行压成一个全局状态，容易把接收机差异重新混入发射机身份。"
            ),
            style="Body Text",
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "计算位置不必绝对限定在卫星内部。具备数字I/Q接口和余量的再生式或专用感知载荷，可以在星上"
                "执行冻结特征提取与轻量闭式更新，从而减少原始I/Q回传、降低告警时延，并在馈电链路间歇时保持"
                "本地能力。透明转发载荷、算力不足或原始I/Q允许安全下传时，应把ground-assisted网关方案作为"
                "基线。严谨表述应是“以星载目标接收域为依据，按接口和资源选择星上或星地协同实现”，而不是"
                "声称所有训练都必须在轨完成。"
            ),
            style="Body Text",
        )
        add_result_table_before(
            doc,
            chapter6_heading,
            REQUIREMENT_TABLE,
            fractions=[0.18, 0.45, 0.37],
            font_size=7.5,
            left_columns=(0, 1, 2),
            group_column=0,
        )
        append_paragraph_before(
            doc,
            chapter6_heading,
            (
                "当前报告中的地面OTA数据和LEO弱信道变换用于验证算法在目标域压力下的行为，尚不能证明真实"
                "Iridium或其他在轨接收链上的性能。若要形成星上部署结论，还需要目标频段多接收机采集、RF"
                "硬件在环、受控在轨I/Q、飞行软件资源测量以及运营方授权／回滚闭环。"
            ),
            style="Body Text",
        )

        add_reference(doc, "[6] 3GPP. Non-Terrestrial Networks (NTN) overview. https://www.3gpp.org/technologies/ntn-overview")
        add_reference(doc, "[7] Iridium. Iridium NTN Direct. https://www.iridium.com/services/iridium-ntn-direct")
        add_reference(doc, "[8] Iridium. On-Air Trials Underway: Iridium NTN Direct Prepares to Enter Beta as Testing Continues, 2026-01-21. https://investor.iridium.com/2026-01-21-On-Air-Trials-Underway-Iridium-NTN-Direct-Prepares-to-Enter-Beta-as-Testing-Continues")

        for table in doc.tables:
            if table.rows:
                set_repeat_table_header(table.rows[0])
            for row in table.rows:
                set_row_cant_split(row)
        normalize_all_visible_run_fonts(doc)
        update_fields_on_open(doc)
        doc.core_properties.title = "CVS-RFFI Phase2详细复现报告（图表与场景优化版）"
        doc.core_properties.subject = "Stage2-C合并结果矩阵与3GPP NTN／Iridium NTN Direct场景研究"
        doc.core_properties.comments = ""
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)


def structural_check(path: Path) -> None:
    doc = Document(path)
    visible = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    required = (
        "4.5.1正式配置矩阵与联合性能图",
        "5.1场景说明",
        "5.2天基射频指纹识别的意义",
        "5.3NTN注册",
        "5.4为什么要在星载目标域进行预适应／类注册，以及实施要求",
        "Iridium NTN Direct",
        "ground-assisted",
    )
    for text in required:
        if text not in visible:
            raise AssertionError(text)
    for removed in ("配置：K=5，新类数=1", "共同切片：K=1，新类数=20", "Starlink", "ICAO地址", "DCP地址"):
        if removed in visible:
            raise AssertionError(f"obsolete text remains: {removed}")
    if len(doc.tables) != 23:
        raise AssertionError(f"expected 23 consolidated tables, found {len(doc.tables)}")
    if len(doc.inline_shapes) != 3:
        raise AssertionError(f"expected 3 figures, found {len(doc.inline_shapes)}")
    with ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
        tracked = root.xpath(".//w:ins | .//w:del | .//w:moveTo | .//w:moveFrom", namespaces={"w": WORD_NS})
        if tracked:
            raise AssertionError(f"tracked change markers remain: {len(tracked)}")
        if b"\xee\x88\x80cite" in archive.read("word/document.xml"):
            raise AssertionError("internal citation token remains")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    revise(args.source, args.output)
    structural_check(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
