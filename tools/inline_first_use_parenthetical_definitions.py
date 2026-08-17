from __future__ import annotations

import argparse
import hashlib
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as _Document

from build_expanded_weekly_reports import add_text_with_inline_math
from revise_phase2_report_numbering_and_results import (
    clean_text,
    normalize_all_visible_run_fonts,
)


LABELS = ("符号说明：", "缩写说明：", "术语说明：")
REQUIRED_DEFINITION_PHRASES = (
    "Radio Frequency Fingerprint Identification",
    "发射机类别集合",
    "一次独立少样本任务",
    "任务开始前的冻结先验状态",
    "源域与目标域数据分布",
    "期望分类风险",
    "进入第",
    "ADV3B02是本项目Phase1地面域泛化基座的实验版本标识",
    "适配/注册前、后的旧类准确率",
    "旧类query集",
    "Prototypical Networks",
    "固定LEO接收IQ",
    "Mitigating Receiver Impact",
    "源域batch与目标域support batch",
    "source样本",
    "Domain Adaptation with Dynamic Distribution Alignment",
    "衡量两个特征向量",
    "归一化样本权重向量",
    "动态权重",
    "Channel Separation Enabled Incremental Learning",
    "当前新类support训练batch",
    "被采样旧类",
    "quantized K-nearest neighbors",
    "当前增量阶段的新类数",
    "Institute of Electrical and Electronics Engineers",
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_paragraph(doc: _Document, exact: str):
    matches = [p for p in doc.paragraphs if clean_text(p.text) == clean_text(exact)]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {exact!r}, found {len(matches)}")
    return matches[0]


def find_paragraph_prefix(doc: _Document, prefix: str):
    target = clean_text(prefix)
    matches = [p for p in doc.paragraphs if clean_text(p.text).startswith(target)]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph prefix {prefix!r}, found {len(matches)}")
    return matches[0]


def append_parenthetical(paragraph, latex_text: str) -> None:
    paragraph.add_run("（")
    add_text_with_inline_math(paragraph, latex_text)
    paragraph.add_run("）")


def insert_parenthetical_after_text(paragraph, anchor: str, latex_text: str) -> None:
    matches = [run for run in paragraph.runs if anchor in run.text]
    if len(matches) != 1:
        raise ValueError(
            f"expected one run containing {anchor!r}, found {len(matches)} "
            f"in paragraph {paragraph.text!r}"
        )

    run = matches[0]
    before, after = run.text.split(anchor, 1)
    run.text = before + anchor

    trailing_run = paragraph.add_run(after)
    if run._r.rPr is not None:
        trailing_run._r.insert(0, deepcopy(run._r.rPr))

    temporary = paragraph._parent.add_paragraph()
    append_parenthetical(temporary, latex_text)
    reference = run._r
    for child in list(temporary._p):
        if child.tag.endswith("}pPr"):
            continue
        reference.addnext(child)
        reference = child
    reference.addnext(trailing_run._r)
    temporary._element.getparent().remove(temporary._element)


def _math_count(doc: _Document) -> int:
    return sum(len(paragraph._p.xpath(".//m:oMath")) for paragraph in doc.paragraphs) + sum(
        len(cell._tc.xpath(".//m:oMath"))
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )


def _visible_paragraphs(doc: _Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def check_document(path: Path, reference: Path | None = None) -> dict[str, int]:
    doc = Document(str(path))
    paragraphs = list(_visible_paragraphs(doc))
    visible = "\n".join(paragraph.text for paragraph in paragraphs)
    standalone_labels = sum(
        clean_text(paragraph.text).startswith(LABELS) for paragraph in doc.paragraphs
    )
    rows = sum(len(table.rows) for table in doc.tables)
    math_count = _math_count(doc)

    if len(doc.tables) != 17:
        raise RuntimeError(f"expected 17 tables, found {len(doc.tables)}")
    if rows != 153:
        raise RuntimeError(f"expected 153 table rows, found {rows}")
    if standalone_labels != 0:
        raise RuntimeError(f"expected no standalone labels, found {standalone_labels}")
    if math_count < 300:
        raise RuntimeError(f"expected at least 300 OMML objects, found {math_count}")
    forbidden = [token for token in ("D92", "\\", "\ufffd") if token in visible]
    if forbidden:
        raise RuntimeError(f"forbidden visible tokens: {forbidden}")
    missing = [phrase for phrase in REQUIRED_DEFINITION_PHRASES if phrase not in visible]
    if missing:
        raise RuntimeError(f"missing definition phrases: {missing}")
    adjacency = (
        "RFFI（Radio Frequency Fingerprint Identification",
        "Prototypical Networks（ProtoNet",
        "MRIOR（Mitigating Receiver Impact",
        "DADDA（Domain Adaptation with Dynamic Distribution Alignment",
        "qKNN（quantized K-nearest neighbors",
    )
    missing_adjacency = [token for token in adjacency if token not in visible]
    if missing_adjacency:
        raise RuntimeError(f"definitions are not adjacent to first-use terms: {missing_adjacency}")
    headers = "".join(
        paragraph.text for section in doc.sections for paragraph in section.header.paragraphs
    )
    footers = "".join(
        paragraph.text for section in doc.sections for paragraph in section.footer.paragraphs
    )
    if headers.strip() or footers.strip():
        raise RuntimeError("header or footer is not empty")

    font_violations = []
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            rfonts = run._r.rPr.rFonts if run._r.rPr is not None else None
            if rfonts is None:
                font_violations.append(run.text[:30])
                continue
            ascii_font = rfonts.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
            )
            hansi_font = rfonts.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi"
            )
            east_asia_font = rfonts.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
            )
            if (
                ascii_font != "Times New Roman"
                or hansi_font != "Times New Roman"
                or east_asia_font != "宋体"
            ):
                font_violations.append(run.text[:30])
    if font_violations:
        raise RuntimeError(f"font violations: {font_violations[:10]}")

    if reference is not None:
        original = Document(str(reference))
        if len(original.paragraphs) != len(doc.paragraphs):
            raise RuntimeError("paragraph count changed")
        if [table._tbl.xml for table in original.tables] != [
            table._tbl.xml for table in doc.tables
        ]:
            raise RuntimeError("table content or formatting changed")

    return {
        "definitions": len(REQUIRED_DEFINITION_PHRASES),
        "tables": len(doc.tables),
        "rows": rows,
        "omml": math_count,
        "standalone_labels": standalone_labels,
    }


def revise(source: Path, output: Path) -> None:
    doc = Document(str(source))
    definition_count = 0

    def insert_after(paragraph, anchor: str, text: str) -> None:
        nonlocal definition_count
        insert_parenthetical_after_text(paragraph, anchor, text)
        definition_count += 1

    def append_to(paragraph, text: str) -> None:
        nonlocal definition_count
        append_parenthetical(paragraph, text)
        definition_count += 1

    taxonomy = find_paragraph_prefix(doc, "表1从三个相互独立的轴定位RFFI任务")
    insert_after(
        taxonomy,
        "RFFI",
        (
            "Radio Frequency Fingerprint Identification，即射频指纹识别；"
            "FSL是Few-Shot Learning，即少样本学习；DA是Domain Adaptation，即域适应；"
            "CIL是Class-Incremental Learning，即类增量学习；FSCIL是Few-Shot "
            "Class-Incremental Learning，即少样本类增量学习；UDA是Unsupervised "
            "Domain Adaptation，即无监督域适应。IQ是In-phase/Quadrature，即同相/正交"
            "基带采样；LEO是Low Earth Orbit，即低地球轨道；TX是transmitter，即发射机；"
            "CVS-RFFI是本项目与模型流程的名称，其中CVS不在本文中充当数学变量"
        ),
    )
    append_to(
        taxonomy,
        (
            r"\(\mathcal C\)表示发射机类别集合，物理上每个类别对应一台发射机身份；"
            r"下标\(B,N,S,T\)依次表示base（基类）、novel（新类）、source（源域）和"
            r"target（目标域）。\(d\)表示观测域，物理上由接收机链路、采集条件和信道共同"
            r"决定；因此\(d_S\ne d_T\)表示源接收机域与目标接收机域不同。\(t\)表示"
            r"增量session序号，\(\mathcal C^{(\le t)}\)表示截至session \(t\)已注册的"
            r"全部类别，\(\mathcal C_{\mathrm{unknown}}\)表示尚未注册的发射机。"
            r"\(K\)是每类互不重复的物理support观测数；\(\cap\)、\(\cup\)、"
            r"\(\varnothing\)、\(\subset\)分别表示交集、并集、空集和真子集"
        ),
    )

    nway = find_paragraph_prefix(doc, "一次-way -shot任务记为")
    append_to(
        nway,
        (
            r"\(\tau\)是一次独立少样本任务；\(S_\tau\)和\(Q_\tau\)分别是该任务的"
            r"support集合与query集合。support是允许建立或更新模型状态的带标签样本，query"
            r"只用于冻结后的测试。\(N\)是任务中的类别数，\(K\)是每类support数，"
            r"\(\mathcal C_\tau\)是任务类别集合；\(c\)是类别索引，\(k\)是类内样本"
            r"索引，\(x_{c,k}\)是类别\(c\)的第\(k\)个固定接收IQ观测。\(|\cdot|\)"
            r"表示集合基数，\(\bigcup\)表示把各类样本合并为support集合"
        ),
    )
    learner = find_paragraph_prefix(doc, "学习算法读取先验状态")
    append_to(
        learner,
        (
            r"\(\Omega_0\)是任务开始前的冻结先验状态，例如Phase1 deployment bundle；"
            r"\(\mathcal A\)是只读取先验状态和support的适配/注册算法；\(h_\tau\)是由它"
            r"生成的当前任务预测器。\(x_q\)是单个query观测，\(y_q\)是真实发射机标签，"
            r"\(\hat y_q\)是模型预测；\(h_\tau(x_q)_c\)是query属于类别\(c\)的分数。"
            r"\(\arg\max\)表示选择分数最大的类别。\(S_\tau\cap Q_\tau=\varnothing\)"
            r"表示support与query不能共享同一物理样本"
        ),
    )

    domain_intro = find_paragraph_prefix(doc, "source域与target域可写为")
    append_to(
        domain_intro,
        (
            r"\(\mathcal D_s\)和\(\mathcal D_t\)分别表示源域与目标域数据分布；"
            r"\(\mathcal X\)是接收IQ的输入空间，\(\mathcal Y\)是发射机标签空间。"
            r"\(P_s(X,Y)\)和\(P_t(X,Y)\)是两个域中的联合分布，\(P(X\mid Y)\)表示"
            r"给定同一发射机身份时观测到某种IQ波形的条件分布。CVS中的物理含义是发射机"
            r"身份不变，但目标接收机与LEO弱信道使波形分布改变"
        ),
    )
    risk_intro = find_paragraph_prefix(doc, "域适应的目标是利用source知识")
    append_to(
        risk_intro,
        (
            r"\(R_t(h)\)是预测器\(h\)在目标接收机域上的期望分类风险；"
            r"\(\ell(h(X),Y)\)是单个样本的分类损失；"
            r"\(\mathbb E_{(X,Y)\sim P_t}[\cdot]\)表示对目标域分布中的样本取数学期望。"
            r"物理上，降低\(R_t\)就是减少目标卫星接收机上旧发射机的平均识别错误"
        ),
    )
    incremental_intro = find_paragraph_prefix(doc, "类增量学习要求新类按阶段到达")
    append_to(
        incremental_intro,
        (
            r"\(\mathcal C^{(\le t-1)}\)是进入第\(t\)次增量前已经注册的类别，"
            r"\(\mathcal C_t^{\mathrm{new}}\)是本次到达并获得合法support的新发射机，"
            r"二者并集构成\(\mathcal C^{(\le t)}\)。\(h_t\)是完成第\(t\)次更新后"
            r"持续保存的预测器，\(h_t(x)_c\)是观测\(x\)对类别\(c\)的分类分数。"
            r"统一\(\arg\max\)意味着旧类和已注册新类在同一候选空间竞争"
        ),
    )

    phase1_sentence = find_paragraph_prefix(doc, "两轮实验统一使用ADV3B02")
    insert_after(
        phase1_sentence,
        "ADV3B02",
        (
            "ADV3B02是本项目Phase1地面域泛化基座的实验版本标识，不是通用算法缩写；"
            "checkpoint是训练后冻结的模型参数快照，deployment bundle是与checkpoint共同"
            "封存、可供Phase2只读使用的部署状态。backbone指把接收IQ映射为身份特征的主干"
            "网络，adapter是附加的小型可训练适配模块，prototype是某一发射机类别的特征中心"
        ),
    )

    metric_intro = find_paragraph_prefix(doc, "记和分别为旧类")
    append_to(
        metric_intro,
        (
            r"\(A_{old}^{pre}\)和\(A_{old}^{post}\)分别是适配/注册前、后的旧类准确率；"
            r"\(G_{old}\)是旧类适配收益；\(A_{new}\)是已注册新类准确率；"
            r"\(H_{old,new}\)是旧类与新类准确率的调和均值；\(F_{old}\)是旧类遗忘量；"
            r"\(A_{min,old}\)是所有旧发射机中最低的单类准确率。\(\Delta\)表示相对正式"
            r"LEO结果的变化量；pp是percentage points，即百分点；s是second，即秒"
        ),
    )
    append_to(
        metric_intro,
        (
            r"\(Q_{old}\)、\(Q_{new}\)和\(Q_c\)分别是旧类query集、新类query集和旧类"
            r"\(c\)的query子集；\(i\)是query样本索引。\(y_i\)是真值，"
            r"\(\hat y_i^{(0)}\)与\(\hat y_i^{(1)}\)分别是状态更新前后的预测。"
            r"\(\mathbb I[\cdot]\)是指示函数：条件成立取1，否则取0；求和后除以query数"
            r"即得到准确率。\(\min\)表示在旧类集合\(Y_{old}\)中取最低单类准确率"
        ),
    )

    protonet_intro = find_paragraph_prefix(doc, "Snell等人提出的Prototypical Networks")
    insert_after(
        protonet_intro,
        "Prototypical Networks",
        (
            "ProtoNet，即原型网络；CDA在本文中指Cross-Domain Adaptation，即跨域适配。"
            "该基线不更新backbone，只用目标接收机旧类support重新估计类别中心"
        ),
    )
    protonet_formula = find_paragraph_prefix(doc, "CVS数据与更新： 加载ADV3B02")
    append_to(
        protonet_formula,
        (
            r"\(x\)是固定LEO接收IQ，\(f_\theta\)是参数为\(\theta\)的冻结身份编码器，"
            r"\(z_i=f_\theta(x_i)\)是第\(i\)个support的160维身份特征。\(S_c\)是类别"
            r"\(c\)的support集合，\(p\)是待求中心，\(p_c^*\)是使类内平方欧氏距离最小"
            r"的类别原型；\(\|\cdot\|_2^2\)表示平方L2范数。\(\arg\min\)表示选择距离"
            r"最小的中心或类别"
        ),
    )

    mrior_intro = find_paragraph_prefix(doc, "Yang等人的MRIOR原本属于")
    insert_after(
        mrior_intro,
        "MRIOR",
        (
            "Mitigating Receiver Impact on Radio Frequency Fingerprint Identification via "
            "Domain Adaptation；SDA是Supervised Domain Adaptation，即监督域适应。CE是"
            "Cross-Entropy，即交叉熵；DV-KL是Donsker-Varadhan形式的Kullback-Leibler"
            "域差异估计，其中KL表示相对熵；Adam是Adaptive Moment Estimation优化器"
        ),
    )
    mrior_batch = find_paragraph_prefix(doc, "记为source和target-support batch")
    append_to(
        mrior_batch,
        (
            r"\(B_s\)和\(B_t\)分别是源域batch与目标域support batch，\(|B_D|\)是域"
            r"\(D\)的batch样本数；\(D\in\{s,t\}\)表示公式同时用于source与target。"
            r"\(p_\theta(c\mid x)\)是模型把观测\(x\)判为旧类\(c\)的概率；\(n_c\)是"
            r"目标support中类别\(c\)的样本数，\(C\)是旧类总数，\(w_c\)是逆频率类别权重，"
            r"\(\epsilon\)防止样本数为0时除零；\(\log\)是自然对数"
        ),
    )
    mrior_dv = find_paragraph_prefix(doc, "令为DV估计网络")
    append_to(
        mrior_dv,
        (
            r"\(z_i^s\)和\(z_j^t\)是source样本\(i\)与target-support样本\(j\)的特征；"
            r"\(T_\phi\)是参数为\(\phi\)的DV域统计网络。\(\exp\)是指数函数，"
            r"\(\mathcal L_{DV-KL}\)越大表示两个接收机域越可区分；训练采用内层最大化域"
            r"差异估计、外层最小化分类与对齐目标。总损失中的0.5、0.5和0.005分别是source"
            r"分类、target-support分类和域差异项的固定权重"
        ),
    )

    dadda_intro = find_paragraph_prefix(doc, "Feng等人的DADDA属于")
    insert_after(
        dadda_intro,
        "DADDA",
        (
            "Domain Adaptation with Dynamic Distribution Alignment，即动态分布对齐域适应；"
            "SDA仍表示监督域适应。MMD是Maximum Mean Discrepancy，即最大均值差异；LMMD是"
            "Local Maximum Mean Discrepancy，即按类别局部计算的最大均值差异；RBF是Radial "
            "Basis Function，即径向基函数核；SGD是Stochastic Gradient Descent，即随机梯度下降"
        ),
    )
    rbf_intro = find_paragraph_prefix(doc, "采用RBF核")
    append_to(
        rbf_intro,
        (
            r"\(k(u,v)\)是衡量两个特征向量\(u\)与\(v\)相似度的RBF核，距离越近核值"
            r"越大；\(\sigma\)是核带宽，控制多大特征距离仍被视为相似。\(n_s\)和\(n_t\)"
            r"分别是source与target-support样本数；\(i,i',j,j'\)是两域中的样本索引。"
            r"\(\mathcal L_{MMD}\)比较两域整体特征分布，数值越小表示全局域差异越小"
        ),
    )
    lmmd_intro = find_paragraph_prefix(doc, "对类别，令为")
    append_to(
        lmmd_intro,
        (
            r"\(w_s^c\)和\(w_t^c\)是类别\(c\)在source/target-support中的归一化样本"
            r"权重向量；\(K_{ss}\)、\(K_{tt}\)和\(K_{st}\)分别是source内部、target内部"
            r"以及source-target之间的核矩阵；上标\(\top\)表示转置。"
            r"\(\sum_{c=1}^{C}\)把所有旧发射机的类条件差异汇总为\(\mathcal L_{LMMD}\)"
        ),
    )
    dadda_total = find_paragraph_prefix(doc, "动态系数和总损失分别为")
    append_to(
        dadda_total,
        (
            r"\(\alpha\in[0,1]\)是由当前batch自动计算的动态权重：全局MMD较大时更强调"
            r"LMMD，反之更强调MMD；分母中的\(\epsilon\)保证数值稳定。"
            r"\(\mathcal L_{CE}^{source}\)和\(\mathcal L_{CE}^{target-support}\)分别监督"
            r"源样本与目标support的发射机分类；\(\mathcal L_{LMMD-sum}\)是各类LMMD项"
            r"的求和；momentum是动量系数，weight decay是权重衰减"
        ),
    )

    stage2c_intro = find_paragraph_prefix(doc, "Stage2-C在同一target receiver中注册新类")
    append_to(
        stage2c_intro,
        (
            "CSIL是Channel Separation Enabled Incremental Learning，即通道隔离型增量"
            "学习；MoPC-HR是Momentum-based Prototype Correction and Hierarchical "
            "Regularization，即基于动量的原型校正与分层正则；HR表示Hierarchical "
            "Regularization；IoT是Internet of Things，即物联网。exemplar是保存供后续"
            "回放的历史原始样本；无exemplar表示不保存旧类原始IQ，但仍可保留旧模型、"
            "Fisher信息或类别prototype等聚合历史状态"
        ),
    )
    csil_intro = find_paragraph_prefix(doc, "记为新类support训练batch")
    append_to(
        csil_intro,
        (
            r"\(B_{new}\)是当前新类support训练batch，\(q_\theta(c\mid x)\)是当前模型"
            r"对全部已注册类别的概率。\(\theta^*\)是增量前冻结参数，\(\theta_j\)是当前"
            r"第\(j\)个参数，\(F_j\)是该参数对旧任务的重要性。EWC是Elastic Weight "
            r"Consolidation（弹性权重固化），通过\(\mathcal L_{EWC}\)限制重要参数漂移；"
            r"KD是Knowledge Distillation（知识蒸馏），\(r_{old}^*(x)\)和\(r_{old}(x)\)"
            r"分别是旧模型与当前模型对旧类的fingerprint响应，系数\(1/32\)按32维响应取均值。"
            r"\(\mathcal L_{CSIL}\)是新类CE、EWC和KD的加权和；SGDM是带动量的SGD；"
            r"epoch是完整遍历训练集一次，batch size是一次参数更新读取的样本数"
        ),
    )
    mopc_intro = find_paragraph_prefix(doc, "CVS数据与更新： 进入Phase2前用全部8400")
    append_to(
        mopc_intro,
        (
            r"\(p_{c_r}\)是被采样旧类\(c_r\)的prototype，\(\epsilon_r\sim"
            r"\mathcal N(0,0.05^2I)\)是零均值各向同性高斯扰动，\(I\)为单位矩阵，"
            r"\(\tilde z_r\)是增强后的伪旧类特征。\(B_p\)是每次采样的旧prototype数，"
            r"\(g_\theta\)是当前分类器，\(\tau\)是softmax温度；softmax把logit转换为"
            r"类别概率。\(\ell\)是参数组索引，\(L\)是参数组总数，\(a_\ell\)是逐层"
            r"正则权重，\(\theta_\ell^*\)与\(\theta_\ell\)分别是旧/当前参数组。"
            r"\(P_{old}\)和\(P_{new}\)是旧类与新类prototype矩阵，\(P_{new}^*\)是更新前"
            r"的新类prototype；0.97/0.03是保留旧prototype与吸收校正量的固定动量系数"
        ),
    )

    qknn_intro = find_paragraph_prefix(doc, "qKNN的Stage2-B状态读取")
    insert_after(
        qknn_intro,
        "qKNN",
        (
            "quantized K-nearest neighbors，即量化K近邻；其中q表示quantized，KNN表示"
            "K-nearest neighbors。本文按照项目命名约定，用qKNN统称Phase2轻量适配/注册路线。"
            "所列实验版本在量化近邻/原型记忆基础上加入support稳健状态和共享判别头，因此不能"
            "理解为最原始的裸KNN多数投票；S2B-old表示完成旧类target support域适配后、注册"
            "新类前的Stage2-B旧类准确率"
        ),
    )
    lowk_intro = find_paragraph_prefix(doc, "若增量阶段只有个样本")
    append_to(
        lowk_intro,
        (
            r"\(N_{new}\)是当前增量阶段的新类数，\(K\)是每个新类的support数，因此"
            r"\(N_{new}K\)是真实增量训练样本总数；\(B\)是batch size。"
            r"\(\operatorname{floor}(N_{new}K/B)\)表示向下取整后的完整batch数量。"
            r"当其为0且DataLoader设置drop_last=True时，不足一个batch的样本被丢弃，"
            r"optimizer step为0，模型没有发生任何梯度更新"
        ),
    )
    first_reference = find_paragraph_prefix(doc, "[1] SNELL J")
    append_to(
        first_reference,
        (
            "IEEE是Institute of Electrical and Electronics Engineers，即电气电子工程师"
            "学会；DOI是Digital Object Identifier，即数字对象唯一标识符；参考文献类型标识"
            "[J]表示期刊论文，[C]表示会议论文"
        ),
    )

    if definition_count != 25:
        raise RuntimeError(f"expected 25 definition groups, found {definition_count}")
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    normalize_all_visible_run_fonts(doc)
    doc.core_properties.title = "CVS-RFFI Phase2详细复现报告（首次出现括号定义版）"
    doc.core_properties.subject = "Phase2任务定义、方法公式、缩写全称、物理意义与实验结果"
    doc.core_properties.comments = (
        "definitions=25；全部定义位于首次正文使用处的中文全角括号内；"
        "实验数据、表格和参考文献范围保持不变。"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    results = check_document(output, source)
    if results["definitions"] != definition_count:
        raise RuntimeError("definition validation count mismatch")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Inline Phase2 symbols and abbreviations in full-width parentheses."
    )
    parser.add_argument("--check", action="store_true")
    parser.add_argument("paths", nargs="+", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.check:
        if len(args.paths) != 1:
            raise ValueError("--check requires exactly one DOCX path")
        path = args.paths[0].resolve()
        results = check_document(path)
        print("STRUCTURE_QA=PASS")
        print(" ".join(f"{key}={value}" for key, value in results.items()))
        return
    if len(args.paths) != 2:
        raise ValueError("generation requires SOURCE and OUTPUT paths")
    source, output = (path.resolve() for path in args.paths)
    if source == output:
        raise ValueError("source and output must be different files")
    if not source.is_file():
        raise FileNotFoundError(source)
    revise(source, output)
    print(f"source_sha256={sha256(source)}")
    print(f"output_sha256={sha256(output)}")
    print("definitions=25 tables=17 standalone_labels=0")
    print(f"output={output}")


if __name__ == "__main__":
    main()
