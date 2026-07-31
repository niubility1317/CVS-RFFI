from __future__ import annotations

import argparse
import hashlib
from copy import deepcopy
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from revise_phase2_report_numbering_and_results import (
    HEADING_BLUE,
    add_result_table_before,
    append_paragraph_before,
    apply_captured_cell_content,
    capture_cell_content,
    clean_text,
    find_formula_header_table,
    find_paragraph,
    normalize_all_visible_run_fonts,
    replace_paragraph_text,
    set_cell_shading,
    set_repeat_table_header,
    set_row_cant_split,
    style_run,
    table_matrix,
)


STAGE2B_D92_TABLE = [
    [
        "K",
        "直接ADV3B02",
        "MRIOR-SDA",
        "DADDA-SDA",
        "ProtoNet CDA",
        "qKNN Stage2-B旧类准确率",
    ],
    ["1", "75.21%", "69.88%", "72.58%", "59.47%", "68.144%"],
    ["5", "75.21%", "79.17%", "76.74%", "70.28%", "81.267%"],
    ["10", "75.21%", "84.50%", "79.36%", "70.86%", "86.111%"],
]

STAGE2B_D92_DETAIL_TABLE = [
    ["K", "target receiver", "clear", "low-elev", "rain", "三场景均值"],
    ["1", "20-1", "67.000%", "67.000%", "65.500%", "66.500%"],
    ["1", "3-19", "52.833%", "43.333%", "52.667%", "49.611%"],
    ["1", "7-14", "80.000%", "80.833%", "76.667%", "79.167%"],
    ["1", "7-7", "81.667%", "74.167%", "71.000%", "75.611%"],
    ["1", "8-8", "75.167%", "70.500%", "63.833%", "69.833%"],
    ["1", "5接收机总体", "71.333%", "67.167%", "65.933%", "68.144%"],
    ["5", "20-1", "88.667%", "83.667%", "79.333%", "83.889%"],
    ["5", "3-19", "69.833%", "65.667%", "65.167%", "66.889%"],
    ["5", "7-14", "87.667%", "80.500%", "77.667%", "81.944%"],
    ["5", "7-7", "93.500%", "89.667%", "90.333%", "91.167%"],
    ["5", "8-8", "85.500%", "82.500%", "79.333%", "82.444%"],
    ["5", "5接收机总体", "85.033%", "80.400%", "78.367%", "81.267%"],
    ["10", "20-1", "93.000%", "88.000%", "86.500%", "89.167%"],
    ["10", "3-19", "76.000%", "72.500%", "71.500%", "73.333%"],
    ["10", "7-14", "91.333%", "83.000%", "84.500%", "86.278%"],
    ["10", "7-7", "96.833%", "96.333%", "94.167%", "95.778%"],
    ["10", "8-8", "91.167%", "83.833%", "83.000%", "86.000%"],
    ["10", "5接收机总体", "89.667%", "84.733%", "83.933%", "86.111%"],
]


STAGE2C_D92_COMPARISON_TABLE = [
    ["切片", "方法", "适应前", "", "", "", ""],
    ["K1/new20", "qKNN", "68.144%", "44.033%", "27.150%", "33.410%", "24.111pp"],
    [
        "K1/new20",
        "CSIL官方流程",
        "42.833%",
        "42.833%",
        "0.000%",
        "0.000%",
        "0.000pp",
    ],
    [
        "K1/new20",
        "MoPC-HR官方流程",
        "45.322%",
        "40.722%",
        "1.363%",
        "2.603%",
        "4.600pp",
    ],
    ["K5/new20", "qKNN", "81.267%", "63.711%", "58.883%", "60.955%", "17.556pp"],
    [
        "K5/new20",
        "CSIL官方流程",
        "42.833%",
        "0.200%",
        "5.557%",
        "0.316%",
        "42.633pp",
    ],
    [
        "K5/new20",
        "MoPC-HR官方流程",
        "45.322%",
        "13.511%",
        "17.433%",
        "14.309%",
        "31.811pp",
    ],
    ["K10/new5", "qKNN", "86.111%", "76.189%", "74.133%", "74.803%", "9.922pp"],
    [
        "K10/new5",
        "CSIL官方流程",
        "42.833%",
        "0.689%",
        "20.413%",
        "1.264%",
        "42.144pp",
    ],
    [
        "K10/new5",
        "MoPC-HR官方流程",
        "45.322%",
        "9.322%",
        "49.547%",
        "14.947%",
        "36.000pp",
    ],
    ["K10/new10", "qKNN", "86.111%", "72.533%", "66.353%", "69.106%", "13.578pp"],
    [
        "K10/new10",
        "CSIL官方流程",
        "42.833%",
        "0.000%",
        "10.460%",
        "0.000%",
        "42.833pp",
    ],
    [
        "K10/new10",
        "MoPC-HR官方流程",
        "45.322%",
        "9.500%",
        "32.900%",
        "13.770%",
        "35.822pp",
    ],
    ["K10/new20", "qKNN", "86.111%", "71.333%", "68.150%", "69.555%", "14.778pp"],
    [
        "K10/new20",
        "CSIL官方流程",
        "42.833%",
        "38.222%",
        "1.660%",
        "2.979%",
        "4.611pp",
    ],
    [
        "K10/new20",
        "MoPC-HR官方流程",
        "45.322%",
        "7.611%",
        "25.187%",
        "10.695%",
        "37.711pp",
    ],
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_table_by_math_signature(
    doc: _Document,
    signature: Sequence[str],
) -> object:
    matches = []
    for table in doc.tables:
        if len(table.columns) != len(signature):
            continue
        math_values = [
            "".join(cell._tc.xpath(".//m:t/text()"))
            for cell in table.rows[0].cells
        ]
        if math_values == list(signature):
            matches.append(table)
    if len(matches) != 1:
        raise ValueError(
            f"expected one table with math signature {signature!r}, found {len(matches)}"
        )
    return matches[0]


def style_inserted_formula_headers(table, formula_contents: Sequence[tuple[int, Sequence]]) -> None:
    for column_index, captured in formula_contents:
        apply_captured_cell_content(
            table.rows[0].cells[column_index],
            captured,
            font_size=7.8,
        )
        set_cell_shading(table.rows[0].cells[column_index], "E7E6E6")
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)


def set_table_header_keep_with_next(table) -> None:
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                style_run(run, bold=True, color=HEADING_BLUE)


def revise(source: Path, output: Path) -> None:
    doc = Document(str(source))
    original_tables = [table_matrix(table) for table in doc.tables]
    if len(original_tables) != 14:
        raise RuntimeError(f"expected 14 source tables, found {len(original_tables)}")

    formal_table = find_formula_header_table(
        doc,
        column_count=7,
        math_signature=["", "", "", "Aoldpost", "Anew", "Hold,new", "Fold"],
        word_signature=["K-shot", "方法", "新类数", "", "", "", ""],
    )
    stage2b_overall = find_table_by_math_signature(
        doc,
        ["", "Aold", "Aold", "Gold", "", "", ""],
    )
    pre_header = capture_cell_content(stage2b_overall.rows[0].cells[1])
    post_header = capture_cell_content(formal_table.rows[0].cells[3])
    new_header = capture_cell_content(formal_table.rows[0].cells[4])
    h_header = capture_cell_content(formal_table.rows[0].cells[5])
    forgetting_header = capture_cell_content(formal_table.rows[0].cells[6])

    boundary_heading = find_paragraph(doc, "3.6.4结果边界")
    replace_paragraph_text(boundary_heading, "3.6.6结果边界")
    d92_stage2b_heading = append_paragraph_before(
        doc,
        boundary_heading,
        "3.6.4qKNN Stage2-B域适应实验结果",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        boundary_heading,
        (
            "qKNN的Stage2-B状态读取6个旧类的K-shot target support，在注册任何"
            "新类之前完成旧类目标域适配。原始artifact将该指标记为B-old；为避免把“注册前”"
            "误解成“尚未域适应”，本节统一记为S2B-old。下表按K、target receiver和LEO弱"
            "信道场景列出S2B-old，每个单元格为5个seed（713102–713106）的均值。"
        ),
        style="Body Text",
    )
    stage2b_detail_table = add_result_table_before(
        doc,
        boundary_heading,
        STAGE2B_D92_DETAIL_TABLE,
        fractions=[0.07, 0.20, 0.17, 0.19, 0.17, 0.20],
        font_size=8.0,
        left_columns=(1,),
        group_column=0,
    )
    set_table_header_keep_with_next(stage2b_detail_table)
    append_paragraph_before(
        doc,
        boundary_heading,
        (
            "qKNN的Stage2-B域适应准确率随support由K=1增加到K=5、K=10而提升：5个接收机"
            "三场景总体均值依次为68.144%、81.267%和86.111%。同一K下，clear场景总体"
            "最高，low-elev与rain场景更低；接收机3-19始终最弱，K=10时三场景均值仍仅"
            "73.333%，说明接收机域偏移和弱信道扰动尚未被完全消除。"
        ),
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        boundary_heading,
        "3.6.5qKNN与域适应方法的共同LEO场景对照",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        boundary_heading,
        (
            "下表将qKNN的S2B-old与相同三类LEO弱信道场景下的域适应对比方法并列。域适应"
            "矩阵使用5个receiver、seed 713101–713105，qKNN使用5个receiver和seed "
            "713102–713106；两套矩阵的seed集合及artifact没有严格一一配对，因此只作"
            "描述性比较，不计算paired显著性。"
        ),
        style="Body Text",
    )
    stage2b_table = add_result_table_before(
        doc,
        boundary_heading,
        STAGE2B_D92_TABLE,
        fractions=[0.07, 0.17, 0.17, 0.17, 0.17, 0.25],
        font_size=8.2,
        left_columns=(1, 2, 3, 4, 5),
    )
    set_table_header_keep_with_next(stage2b_table)
    append_paragraph_before(
        doc,
        boundary_heading,
        (
            "在这一描述性口径下，qKNN的Stage2-B旧类准确率在K=5和K=10时分别为81.267%和"
            "86.111%，高于同表中的三个论文适配头；K=1时MRIOR-SDA为69.88%，高于qKNN的"
            "68.144%。这里比较的是完成旧类support域适应后的S2B-old，而不是Phase1模型"
            "直接推理结果。此时尚未加入新类，qKNN的旧/新任务均衡协方差也尚未启用；数值"
            "反映qKNN的旧类support稳健状态构造，不能归因于Stage2-C均衡模块。"
        ),
        style="Body Text",
    )

    boundary_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph._p.getprevious() is not None
        and paragraph._p.getprevious() is boundary_heading._p
    ]
    if boundary_paragraphs:
        boundary_text = clean_text(boundary_paragraphs[0].text)
        if boundary_text:
            replace_paragraph_text(
                boundary_paragraphs[0],
                boundary_text
                + "qKNN补充表只报告完成旧类域适配后的Stage2-B状态，不应被解释为新类注册结果。",
            )

    stage2c_anchor = find_paragraph(doc, "4.6matched无LEO新类归因诊断")
    append_paragraph_before(
        doc,
        stage2c_anchor,
        "4.5.3qKNN与类增量方法的共同LEO切片对照",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        stage2c_anchor,
        (
            "qKNN完成125/125个任务和375/375个LEO场景，失败数为0。下表在"
            "K1/new20、K5/new20、K10/new5、K10/new10和K10/new20五个共同切片上，"
            "对照qKNN与CSIL、MoPC-HR官方流程的三场景平均结果。CSIL和MoPC-HR使用seed "
            "713101–713105，qKNN使用713102–713106；三者的base训练、状态构造和训练权限"
            "也不同，因此该表是同LEO场景的描述性对照，不是严格paired显著性比较。"
        ),
        style="Body Text",
    )
    stage2c_table = add_result_table_before(
        doc,
        stage2c_anchor,
        STAGE2C_D92_COMPARISON_TABLE,
        fractions=[0.14, 0.23, 0.13, 0.13, 0.12, 0.12, 0.13],
        font_size=7.8,
        left_columns=(0, 1),
        group_column=0,
    )
    style_inserted_formula_headers(
        stage2c_table,
        [
            (2, pre_header),
            (3, post_header),
            (4, new_header),
            (5, h_header),
            (6, forgetting_header),
        ],
    )
    set_table_header_keep_with_next(stage2c_table)
    append_paragraph_before(
        doc,
        stage2c_anchor,
        (
            "qKNN在五个共同切片上的旧新调和均值均高于两种论文流程，但不能把差距全部"
            "归因于qKNN算法：qKNN的Stage2-B旧类状态明显更强，方法生命周期和允许使用的历史"
            "状态也不同。qKNN在K10/new20得到注册后旧类71.333%、新类68.150%、调和均值"
            "69.555%，遗忘14.778个百分点；相对其matched D81控制，它改善旧类和遗忘，"
            "但新类下降且全部绝对性能门仍失败。正式结论保持"
            "COMPLETED_DIAGNOSTIC_NEGATIVE_NOT_PROMOTABLE，不能表述为qKNN已完成晋级。"
        ),
        style="Body Text",
    )

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                style_run(run)
    normalize_all_visible_run_fonts(doc)

    doc.core_properties.title = "CVS-RFFI Phase2详细复现报告（qKNN域适应与类增量结果补充版）"
    doc.core_properties.subject = "Phase2对比方法、qKNN Stage2-B域适应结果、Stage2-C类增量结果与证据边界"
    doc.core_properties.comments = (
        "在Stage2-B补充qKNN按接收机和场景展开的域适应结果；"
        "在Stage2-C保留qKNN共同LEO切片类增量对照。"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    check = Document(str(output))
    headings = [
        clean_text(paragraph.text)
        for paragraph in check.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    required = {
        "3.6.4qKNN Stage2-B域适应实验结果",
        "3.6.5qKNN与域适应方法的共同LEO场景对照",
        "3.6.6结果边界",
        "4.5.3qKNN与类增量方法的共同LEO切片对照",
    }
    missing = required.difference(headings)
    if missing:
        raise RuntimeError(f"missing inserted headings: {sorted(missing)}")
    if "3.6.4结果边界" in headings:
        raise RuntimeError("obsolete Stage2-B boundary heading remains")
    if len(check.tables) != 17:
        raise RuntimeError(f"expected 17 final tables, found {len(check.tables)}")

    final_tables = [table_matrix(table) for table in check.tables]
    for original in original_tables:
        if original not in final_tables:
            raise RuntimeError(f"source table was not preserved: {original[0]}")
    if STAGE2B_D92_TABLE not in final_tables:
        raise RuntimeError("Stage2-B qKNN table was not preserved exactly")
    if STAGE2B_D92_DETAIL_TABLE not in final_tables:
        raise RuntimeError("detailed Stage2-B qKNN table was not preserved exactly")
    if not any(
        len(matrix) == len(STAGE2C_D92_COMPARISON_TABLE)
        and matrix[1:] == STAGE2C_D92_COMPARISON_TABLE[1:]
        for matrix in final_tables
    ):
        raise RuntimeError("Stage2-C qKNN table body was not preserved exactly")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Insert formal qKNN same-LEO-scene descriptive results into the "
            "Stage2-B and Stage2-C comparison sections of an existing DOCX."
        )
    )
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source = args.source.resolve()
    output = args.output.resolve()
    if source == output:
        raise ValueError("source and output must be different files")
    if not source.is_file():
        raise FileNotFoundError(source)
    revise(source, output)
    print(f"source_sha256={sha256(source)}")
    print(f"output_sha256={sha256(output)}")
    print(f"output={output}")


if __name__ == "__main__":
    main()
