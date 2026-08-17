from __future__ import annotations

import argparse
import re
from collections import OrderedDict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ACCENT_BLUE = RGBColor(31, 78, 121)
TABLE_WIDTH_DXA = 9360


def visible_text(element) -> str:
    nodes = element.xpath(".//w:t | .//m:t")
    return "".join(node.text or "" for node in nodes).strip()


def cell_text(cell) -> str:
    return visible_text(cell._tc)


def set_run_fonts(run, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")


def append_inline_math(paragraph, expression: str) -> None:
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    run_properties = OxmlElement("m:rPr")
    style = OxmlElement("m:sty")
    style.set(qn("m:val"), "i")
    run_properties.append(style)
    text = OxmlElement("m:t")
    text.text = expression
    math_run.append(run_properties)
    math_run.append(text)
    math.append(math_run)
    paragraph._p.append(math)


def create_configuration_caption(
    document: Document,
    prefix: str,
    k: int,
    new_count: int,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True

    prefix_run = paragraph.add_run(prefix)
    set_run_fonts(prefix_run, bold=True, color=ACCENT_BLUE)
    append_inline_math(paragraph, f"K={k}")
    suffix_run = paragraph.add_run(f"，新类数={new_count}")
    set_run_fonts(suffix_run, bold=True, color=ACCENT_BLUE)
    return paragraph


def set_cell_width(cell_element, width: int) -> None:
    properties = cell_element.find(qn("w:tcPr"))
    if properties is None:
        properties = OxmlElement("w:tcPr")
        cell_element.insert(0, properties)
    tc_width = properties.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.insert(0, tc_width)
    tc_width.set(qn("w:w"), str(width))
    tc_width.set(qn("w:type"), "dxa")


def scaled_widths(source_table, selected_columns: tuple[int, ...]) -> list[int]:
    source_widths = [
        int(column.get(qn("w:w")))
        for column in source_table._tbl.tblGrid.gridCol_lst
    ]
    selected = [source_widths[index] for index in selected_columns]
    total = sum(selected)
    widths = [round(width * TABLE_WIDTH_DXA / total) for width in selected]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def replace_row_cells(target_row, source_row, selected_columns: tuple[int, ...], widths: list[int]) -> None:
    target_cells = list(target_row._tr.tc_lst)
    for target_cell, source_index, width in zip(target_cells, selected_columns, widths):
        copied_cell = deepcopy(source_row._tr.tc_lst[source_index])
        set_cell_width(copied_cell, width)
        target_cell.getparent().replace(target_cell, copied_cell)


def create_compact_table(
    document: Document,
    source_table,
    source_rows,
    selected_columns: tuple[int, ...],
):
    table = document.add_table(rows=0, cols=len(selected_columns))
    table._tbl.replace(table._tbl.tblPr, deepcopy(source_table._tbl.tblPr))
    widths = scaled_widths(source_table, selected_columns)
    for column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        column.set(qn("w:w"), str(width))

    header = table.add_row()
    replace_row_cells(header, source_table.rows[0], selected_columns, widths)
    for source_row in source_rows:
        target_row = table.add_row()
        replace_row_cells(target_row, source_row, selected_columns, widths)
        row_properties = target_row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
    return table


def insert_grouped_tables(
    document: Document,
    source_table,
    groups,
    caption_prefix: str,
    selected_columns: tuple[int, ...],
) -> None:
    anchor = source_table._tbl
    for k, new_count, rows in groups:
        caption = create_configuration_caption(document, caption_prefix, k, new_count)
        compact_table = create_compact_table(document, source_table, rows, selected_columns)
        anchor.addprevious(caption._p)
        anchor.addprevious(compact_table._tbl)
    anchor.getparent().remove(anchor)


def find_result_tables(document: Document):
    formal = common = diagnostic = None
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        first = tuple(cell_text(cell) for cell in table.rows[1].cells)
        if len(table.rows) == 25 and first[:3] == ("5", "CSIL", "1"):
            formal = table
        elif len(table.rows) == 16 and first[:2] == ("K1/new20", "qKNN"):
            common = table
        elif len(table.rows) == 19 and first[:3] == ("5", "CSIL", "3"):
            diagnostic = table
    if formal is None or common is None or diagnostic is None:
        raise RuntimeError("unable to locate all three Phase2 result tables")
    return formal, common, diagnostic


def group_k_new_rows(table, k_values: tuple[int, ...], new_values: tuple[int, ...]):
    rows_by_key: OrderedDict[tuple[int, int], list] = OrderedDict(
        ((k, new_count), []) for k in k_values for new_count in new_values
    )
    for row in table.rows[1:]:
        k = int(cell_text(row.cells[0]))
        new_count = int(cell_text(row.cells[2]))
        rows_by_key[(k, new_count)].append(row)
    empty = [key for key, rows in rows_by_key.items() if not rows]
    if empty:
        raise RuntimeError(f"expected configuration rows are missing: {empty}")
    return [(k, new_count, rows) for (k, new_count), rows in rows_by_key.items()]


def group_common_rows(table):
    rows_by_slice: OrderedDict[str, list] = OrderedDict()
    for row in table.rows[1:]:
        slice_name = cell_text(row.cells[0])
        rows_by_slice.setdefault(slice_name, []).append(row)
    groups = []
    for slice_name, rows in rows_by_slice.items():
        match = re.fullmatch(r"K(\d+)/new(\d+)", slice_name)
        if match is None:
            raise RuntimeError(f"unexpected common-slice name: {slice_name}")
        if len(rows) != 3:
            raise RuntimeError(f"common slice must contain three methods: {slice_name}")
        groups.append((int(match.group(1)), int(match.group(2)), rows))
    return groups


def replace_visible_text(paragraph, old: str, new: str) -> None:
    count = 0
    for node in paragraph._p.xpath(".//w:t"):
        if node.text and old in node.text:
            count += node.text.count(old)
            node.text = node.text.replace(old, new)
    if count != 1:
        raise RuntimeError(f"expected one text replacement {old!r}, found {count}")


def regroup_report(source: Path | str, output: Path | str) -> None:
    source_path = Path(source).resolve()
    output_path = Path(output).resolve()
    if source_path == output_path:
        raise ValueError("source and output must differ")
    if not source_path.is_file():
        raise FileNotFoundError(source_path)

    document = Document(str(source_path))
    formal, common, diagnostic = find_result_tables(document)

    formal_groups = group_k_new_rows(formal, (5, 10, 20), (1, 3, 5, 10, 20, 25))
    common_groups = group_common_rows(common)
    diagnostic_groups = group_k_new_rows(diagnostic, (5, 10, 20), (3, 5, 10, 20, 25))

    for paragraph in document.paragraphs:
        if paragraph.text == "4.5.1逐配置结果":
            replace_visible_text(paragraph, "4.5.1逐配置结果", "4.5.1按K-shot与新类规模拆分的正式结果")
        elif paragraph.text == "4.6.1逐配置结果":
            replace_visible_text(paragraph, "4.6.1逐配置结果", "4.6.1按K-shot与新类规模拆分的诊断结果")
        elif paragraph.text.startswith("下表直接采用《学习进展情况+7.24》"):
            replace_visible_text(paragraph, "下表直接采用", "以下配置表直接采用")
            replace_visible_text(paragraph, "按K-shot、方法和新类数逐项展开", "按K-shot与新类数拆分，并在每个配置内按方法展开")
        elif paragraph.text.startswith("qKNN完成125/125个任务"):
            replace_visible_text(paragraph, "下表在", "以下五张配置表在")

    insert_grouped_tables(document, formal, formal_groups, "配置：", (1, 3, 4, 5, 6))
    insert_grouped_tables(document, common, common_groups, "共同切片：", (1, 2, 3, 4, 5, 6))
    insert_grouped_tables(document, diagnostic, diagnostic_groups, "无LEO诊断：", (1, 3, 4, 5, 6))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Regroup Phase2 result tables by exact K/new-class configuration")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    regroup_report(args.source, args.output)


if __name__ == "__main__":
    main()
