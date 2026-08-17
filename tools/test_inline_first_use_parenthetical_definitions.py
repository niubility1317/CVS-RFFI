from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from inline_first_use_parenthetical_definitions import (
    append_parenthetical,
    insert_parenthetical_after_text,
    revise,
)


class InlineParentheticalDefinitionTests(unittest.TestCase):
    def test_revise_inlines_all_definition_groups_without_changing_tables(self) -> None:
        repository = Path(__file__).resolve().parents[1]
        source = repository / "docs" / "weekly_reports" / (
            "CVS-RFFI_Phase2详细复现报告1_qKNN域适应与类增量结果补充版_"
            "截至20260731.docx"
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "output.docx"
            revise(source, output)

            original = Document(source)
            revised = Document(output)
            self.assertEqual(len(original.tables), 17)
            self.assertEqual(len(revised.tables), 17)
            self.assertEqual(len(original.paragraphs), len(revised.paragraphs))
            self.assertEqual(
                [table._tbl.xml for table in original.tables],
                [table._tbl.xml for table in revised.tables],
            )

            visible = "\n".join(paragraph.text for paragraph in revised.paragraphs)
            self.assertNotIn("符号说明：", visible)
            self.assertNotIn("缩写说明：", visible)
            self.assertNotIn("术语说明：", visible)
            self.assertNotIn("D92", visible)
            self.assertNotIn("\\", visible)

            definition_phrases = [
                "Radio Frequency Fingerprint Identification",
                "发射机类别集合",
                "一次独立少样本任务",
                "任务开始前的冻结先验状态",
                "源域与目标域数据分布",
                "期望分类风险",
                "进入第",
                "ADV3B02是本项目Phase1地面域泛化基座的实验版本标识",
                "适配/注册前、后的旧类准确率",
                "旧类query集",
                "Prototypical Networks",
                "固定LEO接收IQ",
                "Mitigating Receiver Impact",
                "源域batch与目标域support batch",
                "source样本",
                "Domain Adaptation with Dynamic Distribution Alignment",
                "衡量两个特征向量",
                "归一化样本权重向量",
                "动态权重",
                "Channel Separation Enabled Incremental Learning",
                "当前新类support训练batch",
                "被采样旧类",
                "quantized K-nearest neighbors",
                "当前增量阶段的新类数",
                "Institute of Electrical and Electronics Engineers",
            ]
            for phrase in definition_phrases:
                self.assertIn(phrase, visible)

            math_count = sum(
                len(paragraph._p.xpath(".//m:oMath"))
                for paragraph in revised.paragraphs
            ) + sum(
                len(cell._tc.xpath(".//m:oMath"))
                for table in revised.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertGreaterEqual(math_count, 300)

    def test_insert_parenthetical_places_definition_immediately_after_term(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("RFFI任务")

        insert_parenthetical_after_text(
            paragraph,
            "RFFI",
            "Radio Frequency Fingerprint Identification，即射频指纹识别",
        )

        self.assertEqual(
            paragraph.text,
            "RFFI（Radio Frequency Fingerprint Identification，即射频指纹识别）任务",
        )

    def test_append_parenthetical_adds_omml_and_preserves_table(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "fixture.docx"
            document = Document()
            paragraph = document.add_paragraph("一次N-way K-shot任务记为")
            document.add_table(rows=1, cols=1)
            document.save(source)

            loaded = Document(source)
            target = loaded.paragraphs[0]
            before_tables = len(loaded.tables)
            append_parenthetical(
                target,
                r"\(N\)表示类别数；\(K\)表示每类support数。",
            )

            self.assertTrue(target.text.startswith("一次N-way K-shot任务记为（"))
            self.assertTrue(target.text.endswith("）"))
            self.assertEqual(len(target._p.xpath(".//m:oMath")), 2)
            self.assertEqual(len(loaded.tables), before_tables)
            self.assertNotIn("符号说明：", target.text)


if __name__ == "__main__":
    unittest.main()
