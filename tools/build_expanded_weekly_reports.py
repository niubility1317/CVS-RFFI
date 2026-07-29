from __future__ import annotations

import argparse
from functools import lru_cache
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = REPO_ROOT / "docs" / "weekly_reports"
REPORTS = (
    (
        SOURCE_DIR / "CVS_RFFI_Phase2阶段工作详细报告_20260724.md",
        "CVS-RFFI_Phase2阶段工作详细报告_截至20260724_导师批注修订版.docx",
        "Phase2阶段综合报告",
    ),
    (
        SOURCE_DIR / "学习进展情况_20260716_详细扩展版.md",
        "学习进展情况+7.16_详细扩展版.docx",
        "跨接收机域适应",
    ),
    (
        SOURCE_DIR / "学习进展情况_20260724_详细扩展版.md",
        "学习进展情况+7.24_详细扩展版.docx",
        "类增量与新类注册",
    ),
)

INK = "1F1F1F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


@lru_cache(maxsize=128)
def latex_to_omml_xml(latex: str) -> bytes | None:
    """Convert one display equation to professional Word OMML via Pandoc."""

    pandoc = shutil.which("pandoc")
    if pandoc is None:
        return None
    with tempfile.TemporaryDirectory(prefix="cvs_report_math_") as temp_dir:
        output = Path(temp_dir) / "equation.docx"
        completed = subprocess.run(
            [
                pandoc,
                "--from=markdown+tex_math_dollars",
                "--to=docx",
                "--output",
                str(output),
            ],
            input=f"$$\n{latex}\n$$\n",
            text=True,
            encoding="utf-8",
            capture_output=True,
            check=False,
        )
        if completed.returncode != 0 or not output.exists():
            return None
        with zipfile.ZipFile(output) as archive:
            document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    equation = root.find(f".//{{{MATH_NS}}}oMathPara")
    return None if equation is None else etree.tostring(equation)


@lru_cache(maxsize=256)
def latex_to_inline_omml_xml(latex: str) -> bytes | None:
    """Convert inline LaTeX to an inline Word OMML equation."""

    equation_paragraph_xml = latex_to_omml_xml(latex)
    if equation_paragraph_xml is None:
        return None
    equation_paragraph = etree.fromstring(equation_paragraph_xml)
    equation = equation_paragraph.find(f".//{{{MATH_NS}}}oMath")
    return None if equation is None else etree.tostring(equation)


def set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        if edge not in kwargs:
            continue
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(kwargs[edge]))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color: str = "C7CDD4", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell, **CELL_MARGINS)


def add_page_field(paragraph) -> None:
    paragraph.add_run("第")
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run("1")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    paragraph.add_run("页")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    if "Equation" in styles:
        equation = styles["Equation"]
    else:
        equation = styles.add_style("Equation", 1)
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    equation._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    equation.font.size = Pt(10.5)
    equation.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(5)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.keep_together = True


def configure_page(doc: Document, topic: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"CVS-RFFI近期学习进展｜{topic}")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)
    for run in p.runs:
        set_run_font(run, size=8.5, color=MUTED)


def strip_inline_math(text: str) -> str:
    text = text.replace("\\left", "").replace("\\right", "")
    text = text.replace("\\bigl", "").replace("\\bigr", "")
    text = text.replace("\\qquad", "    ").replace("\\quad", "  ")
    replacements = {
        r"\varnothing": "∅",
        r"\cap": "∩",
        r"\cup": "∪",
        r"\in": "∈",
        r"\le": "≤",
        r"\ge": "≥",
        r"\neq": "≠",
        r"\approx": "≈",
        r"\propto": "∝",
        r"\sim": "∼",
        r"\cdot": "·",
        r"\times": "×",
        r"\sum": "Σ",
        r"\arg\min": "arg min",
        r"\arg\max": "arg max",
        r"\min": "min",
        r"\max": "max",
        r"\Delta": "Δ",
        r"\epsilon": "ε",
        r"\lambda": "λ",
        r"\sigma": "σ",
        r"\theta": "θ",
        r"\rho": "ρ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\tilde": "̃",
        r"\mathbf": "",
        r"\mathcal": "",
        r"\|": "‖",
        r"\,": "",
        r"\;": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\(?:mathrm|text)\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\mathbb\{?([^{}\s]+)\}?", r"\1", text)
    for _ in range(4):
        new_text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
        if new_text == text:
            break
        text = new_text
    text = text.replace("\\", "")
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


RICH_TOKEN_RE = re.compile(r"(\*\*.*?\*\*|`.*?`)")
INLINE_MATH_RE = re.compile(r"\\\((.*?)\\\)")


def add_text_with_inline_math(
    paragraph,
    text: str,
    *,
    size: float | None = None,
    bold: bool = False,
) -> None:
    pos = 0
    for match in INLINE_MATH_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, bold=bold)
        latex = match.group(1).strip()
        omml_xml = latex_to_inline_omml_xml(latex)
        if omml_xml is None:
            run = paragraph.add_run(strip_inline_math(latex))
            set_run_font(
                run,
                latin="Cambria Math",
                east_asia="Microsoft YaHei",
                size=size,
                bold=bold,
            )
        else:
            paragraph._p.append(etree.fromstring(omml_xml))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=bold)


def add_rich_text(paragraph, text: str, *, size: float | None = None) -> None:
    pos = 0
    for match in RICH_TOKEN_RE.finditer(text):
        if match.start() > pos:
            add_text_with_inline_math(
                paragraph,
                text[pos : match.start()],
                size=size,
            )
        token = match.group(0)
        if token.startswith("**"):
            add_text_with_inline_math(
                paragraph,
                token[2:-2],
                size=size,
                bold=True,
            )
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=(size or 11) - 0.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        add_text_with_inline_math(paragraph, text[pos:], size=size)


def parse_table(lines: Sequence[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def is_numeric_column(values: Iterable[str]) -> bool:
    samples = [v.strip() for v in values if v.strip()]
    if not samples:
        return False
    numeric = 0
    for value in samples:
        if re.fullmatch(r"[+\-−]?\d+(?:\.\d+)?(?:%|pp|s)?", value):
            numeric += 1
    return numeric / len(samples) >= 0.65


def compute_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    columns = list(zip(*rows))
    raw: list[float] = []
    for col in columns:
        max_len = max(
            len(
                re.sub(
                    r"\\\((.*?)\\\)",
                    lambda item: strip_inline_math(item.group(1)),
                    re.sub(r"[*`]", "", cell),
                )
            )
            for cell in col
        )
        if is_numeric_column(col[1:]):
            raw.append(min(max(max_len, 7), 10))
        else:
            raw.append(min(max(max_len, 8), 28))
    total = sum(raw)
    widths = [max(660, round(TABLE_WIDTH_DXA * weight / total)) for weight in raw]
    delta = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += delta
    if widths[-1] < 660:
        need = 660 - widths[-1]
        widths[-1] = 660
        largest = max(range(len(widths) - 1), key=widths.__getitem__)
        widths[largest] -= need
    return widths


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = compute_widths(rows)
    font_size = 8.2 if len(rows[0]) >= 7 else 8.8 if len(rows[0]) == 6 else 9.2

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if is_numeric_column([r[c_idx] for r in rows[1:]]) or c_idx == 0 and len(value) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, value, size=font_size)
            if r_idx == 0:
                set_cell_fill(cell, LIGHT_FILL)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            elif r_idx % 2 == 0 and len(rows) > 12:
                set_cell_fill(cell, "FAFBFC")

    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.space_before = Pt(0)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    add_rich_text(p, text)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="B7C9DD", size="8")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("每周研究进展｜导师汇报详细版")
    set_run_font(run, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(title)
    set_run_font(run, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=13.5, color=DARK_BLUE)


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.append(level)
    number_properties.append(number)
    paragraph_properties.append(number_properties)


def markdown_to_docx(source: Path, destination: Path, topic: str) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError(f"missing report title: {source}")
    title = lines[0][2:].strip()
    subtitle = ""
    body_start = 1
    while body_start < len(lines) and not lines[body_start].strip():
        body_start += 1
    if body_start < len(lines) and lines[body_start].startswith("## "):
        subtitle = lines[body_start][3:].strip()
        body_start += 1

    doc = Document()
    configure_styles(doc)
    configure_page(doc, topic)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.keywords = "RFFI, CVS-RFFI, 域适应, 类增量, 周报"
    add_title_block(doc, title, subtitle)

    i = body_start
    paragraph_buffer: list[str] = []
    active_number_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_rich_text(p, text)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            active_number_id = None
            i += 1
            continue

        numbered_item = re.match(r"^\d+\.\s+", stripped)
        if numbered_item is None:
            active_number_id = None

        if stripped.startswith("$$"):
            flush_paragraph()
            equation_lines: list[str] = []
            if stripped != "$$":
                equation_lines.append(stripped[2:])
            i += 1
            while i < len(lines):
                current = lines[i].strip()
                if current.endswith("$$"):
                    if current != "$$":
                        equation_lines.append(current[:-2])
                    break
                equation_lines.append(current)
                i += 1
            p = doc.add_paragraph(style="Equation")
            latex = " ".join(equation_lines)
            omml_xml = latex_to_omml_xml(latex)
            if omml_xml is None:
                p.add_run(strip_inline_math(latex))
                for run in p.runs:
                    set_run_font(
                        run,
                        latin="Cambria Math",
                        east_asia="Microsoft YaHei",
                        size=10.5,
                        color=DARK_BLUE,
                    )
            else:
                p._p.append(etree.fromstring(omml_xml))
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            heading = stripped[3:].strip()
            if heading in {
                "2.Phase2任务定义与三个Stage",
                "7.五种方法的统一横向比较",
                "附录A：非类增量对比方法复现实验结果",
            }:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, heading)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            heading = stripped[4:].strip()
            if heading in {
                "3.6本报告涉及的数据与实验矩阵",
                "6.7matched无LEO新类归因诊断",
            }:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 2")
            add_rich_text(p, heading)
            i += 1
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 3")
            add_rich_text(p, stripped[5:].strip())
            i += 1
            continue

        if numbered_item:
            flush_paragraph()
            if active_number_id is None:
                active_number_id = create_decimal_numbering(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_number_id)
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:].strip())
            i += 1
            continue

        paragraph_buffer.append(stripped)
        if line.endswith("  "):
            flush_paragraph()
        i += 1

    flush_paragraph()
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path.home() / "Desktop" / "周报",
    )
    args = parser.parse_args()
    for source, filename, topic in REPORTS:
        destination = args.output_dir / filename
        markdown_to_docx(source, destination, topic)
        print(destination)


if __name__ == "__main__":
    main()
