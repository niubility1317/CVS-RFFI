from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from inline_first_use_parenthetical_definitions import normalize_all_visible_run_fonts


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def all_paragraphs(doc: Document) -> Iterable:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def math_nodes(paragraph) -> list:
    return list(paragraph._p.xpath(".//m:t"))


def math_text(paragraph) -> str:
    return "".join(node.text or "" for node in math_nodes(paragraph))


def replace_math_token(paragraph, old: str, new: str, expected: int | None = None) -> int:
    count = 0
    for node in math_nodes(paragraph):
        if node.text == old:
            node.text = new
            count += 1
    if expected is not None and count != expected:
        raise RuntimeError(
            f"math token {old!r}->{new!r}: expected {expected}, found {count}"
        )
    return count


def replace_math_sequence(
    paragraph,
    old: tuple[str, ...],
    new: str,
    expected: int | None = None,
) -> int:
    nodes = math_nodes(paragraph)
    values = [node.text or "" for node in nodes]
    starts: list[int] = []
    index = 0
    while index <= len(values) - len(old):
        if tuple(values[index : index + len(old)]) == old:
            starts.append(index)
            index += len(old)
        else:
            index += 1
    for start in starts:
        nodes[start].text = new
        for offset in range(1, len(old)):
            nodes[start + offset].text = ""
    if expected is not None and len(starts) != expected:
        raise RuntimeError(
            f"math sequence {old!r}->{new!r}: expected {expected}, found {len(starts)}"
        )
    return len(starts)


def replace_math_token_after(
    paragraph,
    before: tuple[str, ...],
    old: str,
    new: str,
    expected: int | None = None,
) -> int:
    nodes = math_nodes(paragraph)
    values = [node.text or "" for node in nodes]
    count = 0
    width = len(before)
    for index in range(width, len(values)):
        if tuple(values[index - width : index]) == before and values[index] == old:
            nodes[index].text = new
            count += 1
    if expected is not None and count != expected:
        raise RuntimeError(
            f"math token after {before!r}, {old!r}->{new!r}: "
            f"expected {expected}, found {count}"
        )
    return count


def set_math_script(node, value: str = "script") -> None:
    run = node.getparent()
    if run.tag != qn("m:r"):
        return
    rpr = run.find(qn("m:rPr"))
    if rpr is None:
        rpr = OxmlElement("m:rPr")
        run.insert(0, rpr)
    script = rpr.find(qn("m:scr"))
    if script is None:
        script = OxmlElement("m:scr")
        rpr.append(script)
    script.set(qn("m:val"), value)


def replace_class_set_token(paragraph, old: str = "C") -> int:
    count = 0
    for node in math_nodes(paragraph):
        if node.text == old:
            node.text = "Y"
            set_math_script(node)
            count += 1
    return count


def script_existing_y(paragraph) -> int:
    count = 0
    for node in math_nodes(paragraph):
        if node.text == "Y":
            set_math_script(node)
            count += 1
    return count


def replace_plain(paragraph, old: str, new: str, expected: int | None = None) -> int:
    count = 0
    for node in paragraph._p.xpath(".//w:t"):
        if node.text and old in node.text:
            occurrences = node.text.count(old)
            node.text = node.text.replace(old, new)
            count += occurrences
    if expected is not None and count != expected:
        raise RuntimeError(
            f"plain text {old!r}->{new!r}: expected {expected}, found {count}"
        )
    return count


def append_plain(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"paragraph not found: {prefix}")


def normalize_domain_tokens(paragraph) -> None:
    replace_math_token(paragraph, "s", "src")
    replace_math_token(paragraph, "t", "tgt")
    replace_math_token(paragraph, "ss", "src,src")
    replace_math_token(paragraph, "tt", "tgt,tgt")
    replace_math_token(paragraph, "st", "src,tgt")
    replace_math_token(paragraph, "source", "src")
    replace_math_token(paragraph, "target-support", "tgt-support")
    replace_math_sequence(paragraph, tuple("source"), "src")
    replace_math_sequence(paragraph, tuple("target-support"), "tgt-support")


def normalize(source: Path, output: Path) -> None:
    source = source.resolve()
    output = output.resolve()
    if source == output:
        raise ValueError("source and output must differ")
    if not source.is_file():
        raise FileNotFoundError(source)

    doc = Document(str(source))

    # Category sets: one symbol family, Y, throughout task definitions and result metrics.
    for index in (2, 7, 8, 10, 19, 24, 25):
        replace_class_set_token(doc.paragraphs[index])
    for paragraph in doc.tables[0]._tbl.xpath(".//w:p"):
        class Wrapper:
            _p = paragraph

        replace_class_set_token(Wrapper())
    for index in (16, 17, 20, 21, 43, 65, 77):
        script_existing_y(doc.paragraphs[index])

    # Domain labels use src/tgt; t is reserved for the incremental session index.
    replace_math_token(doc.paragraphs[2], "S", "src", expected=2)
    replace_math_token(doc.paragraphs[2], "T", "tgt", expected=2)
    for index in (16, 17, 19, 20, 21, 81, 84, 85, 86, 88, 93, 94, 97, 98, 99, 100, 102):
        normalize_domain_tokens(doc.paragraphs[index])

    # The pointwise loss has a semantic subscript; the parameter-group index is g.
    replace_math_token(doc.paragraphs[20], "ℓ", "ℓcls", expected=1)
    replace_math_token(doc.paragraphs[21], "ℓ", "ℓcls", expected=1)

    # Explicit adaptation/registration states replace overloaded 0/1 and pre/post.
    p = doc.paragraphs[43]
    replace_math_token(p, "pre", "DA0_REG0", expected=1)
    replace_math_sequence(p, tuple("pre"), "DA0_REG0", expected=1)
    replace_math_token(p, "post", "DA1_REG0", expected=1)
    replace_math_sequence(p, tuple("post"), "DA1_REG0", expected=1)
    replace_math_token(p, "0", "DA0_REG0", expected=2)
    replace_math_token(p, "1", "DA1_REG0", expected=2)
    replace_plain(
        p,
        "分别表示适应/注册前后的预测",
        "分别表示Stage2-B域适应前后的预测",
        expected=1,
    )
    replace_plain(
        p,
        "分别是适配/注册前、后的旧类准确率",
        "分别是DA0_REG0与DA1_REG0状态的旧类准确率",
        expected=1,
    )
    replace_plain(
        p,
        "分别是状态更新前后的预测",
        "分别是DA0_REG0与DA1_REG0状态的预测",
        expected=1,
    )
    append_plain(
        p,
        " 四个完整状态记为DA0_REG0（域适应前/注册前）、DA1_REG0（域适应后/注册前）、"
        "DA0_REG1（域适应前/注册后）和DA1_REG1（域适应后/注册后）。",
    )

    replace_math_token(doc.paragraphs[46], "pre", "DA0_REG0", expected=1)
    replace_math_token_after(doc.paragraphs[46], ("y", "i"), "0", "DA0_REG0", expected=1)
    replace_math_token(doc.paragraphs[48], "post", "DA1_REG0", expected=1)
    replace_math_token_after(doc.paragraphs[48], ("y", "i"), "1", "DA1_REG0", expected=1)
    replace_math_token(doc.paragraphs[50], "post", "DA1_REG0", expected=1)
    replace_math_token(doc.paragraphs[50], "pre", "DA0_REG0", expected=1)
    replace_math_token_after(doc.paragraphs[54], ("y", "i"), "1", "DA1_REG1", expected=1)
    replace_math_token(doc.paragraphs[58], "post", "DA1_REG1", expected=2)
    replace_math_token(doc.paragraphs[61], "pre", "DA1_REG0", expected=1)
    replace_math_token(doc.paragraphs[61], "post", "DA1_REG1", expected=1)
    replace_math_token_after(doc.paragraphs[64], ("y", "i"), "1", "DA1_REG1", expected=1)

    # Prototype uses mu; pi is reserved for classification probability.
    for index in (74, 75, 77):
        replace_math_token(doc.paragraphs[index], "p", "μ")
    for index, old in ((81, "p"), (84, "p"), (93, "p"), (128, "q"), (129, "q"), (141, "q")):
        replace_math_token(doc.paragraphs[index], old, "π")
    append_plain(doc.paragraphs[81], " 本文各方法的分类概率统一记为上述符号。")

    # Epsilon is split into numerical stabilization and stochastic perturbation.
    for index in (81, 82, 100, 101):
        replace_math_token(doc.paragraphs[index], "ϵ", "ε0")
    for index in (139, 142):
        replace_math_token(doc.paragraphs[index], "ϵ", "ξ")
    replace_plain(doc.paragraphs[139], "是零均值各向同性高斯扰动", "是零均值各向同性高斯扰动向量")

    # MoPC-HR: tau remains the few-shot task symbol; temperature and group index are unique.
    for index in (139, 142, 143):
        replace_math_token(doc.paragraphs[index], "τ", "Ttemp")
    for index in (139, 144, 145):
        replace_math_token(doc.paragraphs[index], "ℓ", "g")
        replace_math_token(doc.paragraphs[index], "ℓ−", "g−")
        replace_math_token(doc.paragraphs[index], "ℓ=", "g=")
    replace_math_token(doc.paragraphs[136], "0.01t", "0.01u", expected=1)
    replace_plain(doc.paragraphs[136], "学习率为", "学习率为（其中u表示optimizer step）", expected=1)
    for index in (139, 142):
        replace_math_token(doc.paragraphs[index], "p", "μ")

    # qKNN acronym K is not the shot budget K.
    qknn = find_paragraph(doc, "qKNN（")
    replace_plain(
        qknn,
        "其中q表示quantized，KNN表示K-nearest neighbors。",
        "其中q表示quantized，KNN表示K-nearest neighbors；qKNN名称中的K不表示K-shot，"
        "本文只有单独出现的K才表示每类support物理样本数。",
        expected=1,
    )

    # Stage-specific table labels remove ambiguous standalone before/after notation.
    for cell, old, new in (
        (doc.tables[3].cell(1, 4), "pre", "DA0_REG0"),
        (doc.tables[3].cell(1, 4), "post", "DA1_REG0"),
        (doc.tables[3].cell(2, 4), "pre", "DA1_REG0"),
        (doc.tables[3].cell(2, 4), "post", "DA1_REG1"),
        (doc.tables[4].cell(4, 3), "post", "DA1_REG1"),
        (doc.tables[7].cell(0, 1), "pre", "DA0_REG0"),
        (doc.tables[7].cell(0, 2), "post", "DA1_REG0"),
        (doc.tables[14].cell(0, 3), "post", "DA0_REG1"),
        (doc.tables[16].cell(0, 5), "post", "DA0_REG1"),
    ):
        for paragraph in cell.paragraphs:
            replace_math_token(paragraph, old, new)
    replace_plain(doc.tables[7].cell(0, 1).paragraphs[0], "适应前", "DA0_REG0")
    replace_plain(doc.tables[7].cell(0, 2).paragraphs[0], "适应后", "DA1_REG0")
    replace_plain(doc.tables[15].cell(0, 2).paragraphs[0], "适应前", "初始状态")
    for paragraph in doc.tables[15].cell(0, 3).paragraphs:
        replace_math_token(paragraph, "post", "REG1")

    shared = find_paragraph(doc, "qKNN完成125/125")
    append_plain(
        shared,
        " 状态路径明确为：qKNN采用DA1_REG0→DA1_REG1；CSIL与MoPC-HR官方流程采用"
        "DA0_REG0→DA0_REG1。",
    )

    normalize_all_visible_run_fonts(doc)
    doc.core_properties.title = "CVS-RFFI Phase2详细复现报告（符号统一版）"
    doc.core_properties.subject = "统一集合、域、状态、概率、prototype和损失记号"
    doc.core_properties.comments = (
        "符号统一；实验数值、17张表、153行记录和5条对比方法参考文献保持不变。"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    audit(output, source)


def result_number_tokens(doc: Document) -> list[str]:
    pattern = re.compile(r"[+-]?\d+(?:\.\d+)?(?:%|pp|s)")
    return [
        token
        for table in doc.tables
        for row in table.rows[1:]
        for cell in row.cells
        for token in pattern.findall(cell.text)
    ]


def audit(path: Path, reference: Path | None = None) -> dict[str, int]:
    doc = Document(str(path))
    rows = sum(len(table.rows) for table in doc.tables)
    references = [p.text for p in doc.paragraphs if p.text.startswith("[")]
    omml = len(doc.element.xpath(".//m:oMath"))
    body = "\n".join("".join(p._p.itertext()) for p in all_paragraphs(doc))
    math = "\n".join(math_text(p) for p in all_paragraphs(doc))

    required = ("src", "tgt", "π", "Ttemp", "ξ", "DA0_REG0", "DA1_REG0", "DA1_REG1")
    missing = [token for token in required if token not in math and token not in body]
    if missing:
        raise RuntimeError(f"missing normalized symbols: {missing}")
    if "分别表示适应/注册前后的预测" in body:
        raise RuntimeError("ambiguous state phrase remains")
    if "qKNN名称中的K不表示K-shot" not in body:
        raise RuntimeError("qKNN/K-shot distinction missing")
    if len(doc.tables) != 17 or rows != 153 or len(references) != 5 or omml < 309:
        raise RuntimeError(
            f"document scale changed: tables={len(doc.tables)} rows={rows} "
            f"references={len(references)} omml={omml}"
        )
    if "�" in body or "\\mathcal" in body:
        raise RuntimeError("damaged or visible source notation found")

    if reference is not None:
        original = Document(str(reference))
        original_references = [
            p.text for p in original.paragraphs if p.text.startswith("[")
        ]
        if references != original_references:
            raise RuntimeError("reference entries changed")
        if result_number_tokens(doc) != result_number_tokens(original):
            raise RuntimeError("experiment result numbers changed")
        if len(original.paragraphs) != len(doc.paragraphs):
            raise RuntimeError("body paragraph count changed")

    return {
        "tables": len(doc.tables),
        "rows": rows,
        "references": len(references),
        "omml": omml,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Normalize Phase2 report symbols.")
    parser.add_argument("--check", action="store_true")
    parser.add_argument("--reference", type=Path)
    parser.add_argument("paths", nargs="+", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.check:
        if len(args.paths) != 1:
            raise ValueError("--check requires one DOCX path")
        result = audit(args.paths[0].resolve(), args.reference.resolve() if args.reference else None)
        print("SYMBOL_QA=PASS")
        print(" ".join(f"{key}={value}" for key, value in result.items()))
        return
    if len(args.paths) != 2:
        raise ValueError("generation requires SOURCE and OUTPUT")
    source, output = (path.resolve() for path in args.paths)
    normalize(source, output)
    print(f"source_sha256={sha256(source)}")
    print(f"output_sha256={sha256(output)}")
    print(f"output={output}")


if __name__ == "__main__":
    main()
