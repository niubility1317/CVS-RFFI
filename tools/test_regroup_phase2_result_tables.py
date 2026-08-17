from __future__ import annotations

import tempfile
import unittest
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from regroup_phase2_result_tables import regroup_report


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/weekly_reports/CVS-RFFI_Phase2详细复现报告1_qKNN符号统一版_截至20260817.docx"


def normalized_cell_text(cell) -> str:
    nodes = cell._tc.xpath(".//w:t | .//m:t")
    return "".join(node.text or "" for node in nodes).strip()


def row_values(row) -> tuple[str, ...]:
    return tuple(normalized_cell_text(cell) for cell in row.cells)


def next_table_after_caption(document: Document, caption: str):
    body = document._element.body
    for index, element in enumerate(body.iterchildren()):
        if element.tag != qn("w:p"):
            continue
        nodes = element.xpath(".//w:t | .//m:t")
        text = "".join(node.text or "" for node in nodes).strip()
        if text != caption:
            continue
        next_element = body[index + 1]
        if next_element.tag != qn("w:tbl"):
            raise AssertionError(f"caption is not followed by a table: {caption}")
        for table in document.tables:
            if table._tbl is next_element:
                return table
    raise AssertionError(f"caption not found: {caption}")


class RegroupPhase2ResultTablesTest(unittest.TestCase):
    def test_result_rows_are_preserved_and_grouped_by_exact_configuration(self):
        source = Document(SOURCE)
        formal_source = Counter(row_values(row) for row in source.tables[14].rows[1:])
        common_source = Counter(row_values(row) for row in source.tables[15].rows[1:])
        diagnostic_source = Counter(row_values(row) for row in source.tables[16].rows[1:])

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "regrouped.docx"
            regroup_report(SOURCE, output)
            result = Document(output)

            formal_captions = [
                f"配置：K={k}，新类数={new_count}"
                for k in (5, 10, 20)
                for new_count in (1, 3, 5, 10, 20, 25)
            ]
            common_configs = ((1, 20), (5, 20), (10, 5), (10, 10), (10, 20))
            common_captions = [
                f"共同切片：K={k}，新类数={new_count}"
                for k, new_count in common_configs
            ]
            diagnostic_captions = [
                f"无LEO诊断：K={k}，新类数={new_count}"
                for k in (5, 10, 20)
                for new_count in (3, 5, 10, 20, 25)
            ]

            formal_result = Counter()
            for caption in formal_captions:
                table = next_table_after_caption(result, caption)
                self.assertEqual(len(table.columns), 5)
                _, config = caption.split("：", 1)
                k_text, new_text = config.split("，")
                k = k_text.split("=")[1]
                new_count = new_text.split("=")[1]
                for row in table.rows[1:]:
                    values = row_values(row)
                    formal_result[(k, values[0], new_count, *values[1:])] += 1

            common_result = Counter()
            for caption, (k, new_count) in zip(common_captions, common_configs):
                table = next_table_after_caption(result, caption)
                self.assertEqual(len(table.columns), 6)
                slice_name = f"K{k}/new{new_count}"
                self.assertEqual(len(table.rows), 4)
                for row in table.rows[1:]:
                    values = row_values(row)
                    common_result[(slice_name, *values)] += 1

            diagnostic_result = Counter()
            for caption in diagnostic_captions:
                table = next_table_after_caption(result, caption)
                self.assertEqual(len(table.columns), 5)
                _, config = caption.split("：", 1)
                k_text, new_text = config.split("，")
                k = k_text.split("=")[1]
                new_count = new_text.split("=")[1]
                for row in table.rows[1:]:
                    values = row_values(row)
                    diagnostic_result[(k, values[0], new_count, *values[1:])] += 1

            self.assertEqual(formal_result, formal_source)
            self.assertEqual(common_result, common_source)
            self.assertEqual(diagnostic_result, diagnostic_source)
            self.assertEqual(len(result.tables), 52)

    def test_metric_headers_remain_word_equations(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "regrouped.docx"
            regroup_report(SOURCE, output)
            result = Document(output)
            for caption in (
                "配置：K=5，新类数=5",
                "共同切片：K=10，新类数=20",
                "无LEO诊断：K=10，新类数=20",
            ):
                table = next_table_after_caption(result, caption)
                math_count = len(table.rows[0]._tr.xpath(".//m:oMath"))
                self.assertGreaterEqual(math_count, 3, caption)


if __name__ == "__main__":
    unittest.main()
