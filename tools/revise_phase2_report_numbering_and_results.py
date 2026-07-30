from __future__ import annotations

import argparse
import hashlib
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


LATIN_FONT = "Times New Roman"
EAST_ASIA_FONT = "宋体"
HEADING_BLUE = RGBColor(31, 78, 121)
TEXT_RED = RGBColor(192, 0, 0)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_rfonts(r_pr, *, latin: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def style_run(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    set_rfonts(run._element.get_or_add_rPr())
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clean_text(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(doc: _Document, exact_text: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if clean_text(p.text) == clean_text(exact_text)]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph {exact_text!r}, found {len(matches)}")
    return matches[0]


def find_paragraph_startswith(doc: _Document, prefix: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if clean_text(p.text).startswith(clean_text(prefix))]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph starting {prefix!r}, found {len(matches)}")
    return matches[0]


def replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    style_run(run, bold=bold, color=color)


def append_paragraph_before(
    doc: _Document,
    anchor: Paragraph,
    text: str,
    *,
    style: str = "Normal",
    keep_with_next: bool = False,
    segments: Sequence[tuple[str, bool, RGBColor | None]] | None = None,
) -> Paragraph:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if segments is None:
        run = paragraph.add_run(text)
        style_run(run)
    else:
        for segment_text, bold, color in segments:
            run = paragraph.add_run(segment_text)
            style_run(run, bold=bold, color=color)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        tr_pr.append(marker)


def set_cell_margins(
    cell,
    *,
    top: int = 55,
    start: int = 70,
    bottom: int = 55,
    end: int = 70,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Table, widths_twips: Sequence[int]) -> None:
    if len(widths_twips) != len(table.columns):
        raise ValueError("column width count does not match table")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    total_width = sum(widths_twips)

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(total_width))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")

    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def table_matrix(table: Table) -> list[list[str]]:
    return [[clean_text(cell.text) for cell in row.cells] for row in table.rows]


def math_text(cell) -> str:
    return "".join(cell._tc.xpath(".//m:t/text()"))


def word_text(cell) -> str:
    return "".join(cell._tc.xpath(".//w:t/text()"))


def find_formula_header_table(
    doc: _Document,
    *,
    column_count: int,
    math_signature: Sequence[str],
    word_signature: Sequence[str] | None = None,
) -> Table:
    matches: list[Table] = []
    for table in doc.tables:
        if not table.rows or len(table.columns) != column_count:
            continue
        math_values = [math_text(cell) for cell in table.rows[0].cells]
        if math_values != list(math_signature):
            continue
        if word_signature is not None:
            word_values = [clean_text(word_text(cell)) for cell in table.rows[0].cells]
            if word_values != [clean_text(value) for value in word_signature]:
                continue
        matches.append(table)
    if len(matches) != 1:
        raise ValueError(
            "expected one formula-header table with "
            f"math signature {math_signature!r}, found {len(matches)}"
        )
    return matches[0]


def capture_cell_content(cell) -> list:
    return [
        deepcopy(child)
        for child in cell._tc
        if child.tag != qn("w:tcPr")
    ]


def apply_captured_cell_content(
    cell,
    captured_content: Sequence,
    *,
    font_size: float,
) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)
    for child in captured_content:
        cell._tc.append(deepcopy(child))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            style_run(run, size=font_size, bold=True, color=HEADING_BLUE)


def source_table(
    doc: _Document,
    *,
    header: Sequence[str],
    expected_rows: int,
) -> list[list[str]]:
    normalized_header = [clean_text(value) for value in header]
    matches: list[list[list[str]]] = []
    for table in doc.tables:
        values = table_matrix(table)
        if values and values[0] == normalized_header:
            matches.append(values)
    if len(matches) != 1:
        raise ValueError(f"expected one source table with header {header!r}, found {len(matches)}")
    values = matches[0]
    if len(values) != expected_rows:
        raise ValueError(
            f"source table {header!r} expected {expected_rows} rows, found {len(values)}"
        )
    return values


def add_result_table_before(
    doc: _Document,
    anchor: Paragraph,
    data: Sequence[Sequence[str]],
    *,
    fractions: Sequence[float],
    font_size: float,
    left_columns: Iterable[int] = (),
    group_column: int | None = None,
) -> Table:
    if not data:
        raise ValueError("table data is empty")
    column_count = len(data[0])
    if any(len(row) != column_count for row in data):
        raise ValueError("table data has inconsistent column counts")
    if len(fractions) != column_count:
        raise ValueError("table fraction count does not match data")

    table = doc.add_table(rows=len(data), cols=column_count)
    if "Table Grid" in [style.name for style in doc.styles]:
        table.style = "Table Grid"

    left_columns_set = set(left_columns)
    previous_group: str | None = None
    group_band = False
    for row_index, values in enumerate(data):
        if row_index > 0 and group_column is not None:
            group_value = values[group_column]
            if group_value != previous_group:
                group_band = not group_band
                previous_group = group_value
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.keep_with_next = row_index == 0
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index in left_columns_set
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in paragraph.runs:
                style_run(
                    run,
                    size=font_size,
                    bold=(row_index == 0),
                    color=HEADING_BLUE if row_index == 0 else None,
                )
            if row_index == 0:
                set_cell_shading(cell, "E7E6E6")
            elif group_column is not None and group_band:
                set_cell_shading(cell, "F4F8FC")

    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)

    section = doc.sections[0]
    usable_width = int(
        section.page_width.twips
        - section.left_margin.twips
        - section.right_margin.twips
    )
    widths = [int(usable_width * fraction) for fraction in fractions[:-1]]
    widths.append(usable_width - sum(widths))
    set_table_geometry(table, widths)

    anchor._p.addprevious(table._tbl)
    return table


def remove_between(start: Paragraph, end: Paragraph) -> None:
    if start._p.getparent() is not end._p.getparent():
        raise ValueError("section anchors do not share the same parent")
    parent = start._p.getparent()
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        parent.remove(node)
        node = next_node
    if node is None:
        raise ValueError(f"end anchor {end.text!r} not found after {start.text!r}")


def remove_from_heading_to_end(heading: Paragraph) -> None:
    parent = heading._p.getparent()
    node = heading._p
    while node is not None:
        next_node = node.getnext()
        if node.tag == qn("w:sectPr"):
            break
        parent.remove(node)
        node = next_node


def find_table(doc: _Document, first_cell: str, second_cell: str | None = None) -> Table:
    matches: list[Table] = []
    for table in doc.tables:
        if not table.rows:
            continue
        header = [clean_text(cell.text) for cell in table.rows[0].cells]
        if not header or header[0] != clean_text(first_cell):
            continue
        if second_cell is not None and (len(header) < 2 or header[1] != clean_text(second_cell)):
            continue
        matches.append(table)
    if len(matches) != 1:
        raise ValueError(
            f"expected one table starting with {first_cell!r}, {second_cell!r}; "
            f"found {len(matches)}"
        )
    return matches[0]


def set_cell_text(cell, text: str, *, size: float = 9.0) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            style_run(run, size=size)


def normalize_all_visible_run_fonts(doc: _Document) -> None:
    """Apply the requested Chinese/Latin font pair, including hyperlink runs."""
    for run_element in doc.element.xpath(".//w:r"):
        set_rfonts(run_element.get_or_add_rPr())


def add_stage2b_results(
    doc: _Document,
    anchor: Paragraph,
    overall: Sequence[Sequence[str]],
    by_k: Sequence[Sequence[str]],
    by_receiver: Sequence[Sequence[str]],
    overall_header_templates: Sequence[Sequence],
) -> None:
    append_paragraph_before(
        doc,
        anchor,
        "3.6对比方法复现实验结果",
        style="Heading 2",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "以下三张表直接采用《学习进展情况+7.16》中的原始结果表，依次给出总体表现、"
            "不同K-shot和不同target receiver的结果。三种方法各运行125个任务，共375个任务；"
            "表内数值不再压缩为单一结论。"
        ),
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        anchor,
        "3.6.1总体性能",
        style="Heading 3",
        keep_with_next=True,
    )
    overall_table = add_result_table_before(
        doc,
        anchor,
        overall,
        fractions=(0.18, 0.13, 0.13, 0.12, 0.18, 0.12, 0.14),
        font_size=8.4,
        left_columns=(0,),
    )
    for target_column, captured_content in zip(
        (1, 2, 3),
        overall_header_templates,
    ):
        apply_captured_cell_content(
            overall_table.rows[0].cells[target_column],
            captured_content,
            font_size=8.4,
        )
    append_paragraph_before(
        doc,
        anchor,
        (
            "MRIOR-SDA的适应后old_acc最高，平均收益为+8.98个百分点；"
            "DADDA-SDA平均收益为+4.75个百分点；ProtoNet CDA平均下降6.75个百分点。"
            "这说明target-old support含有可利用的域校准信息，但固定embedding上的单prototype"
            "不足以修正系统性跨接收机偏移。"
        ),
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        anchor,
        "3.6.2不同K-shot下的结果",
        style="Heading 3",
        keep_with_next=True,
    )
    add_result_table_before(
        doc,
        anchor,
        by_k,
        fractions=(0.10, 0.225, 0.225, 0.22, 0.22),
        font_size=8.8,
        left_columns=(),
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "K从1增加到20时，MRIOR-SDA与DADDA-SDA持续受益，说明更多target support"
            "提高了梯度估计和类条件分布估计的稳定性。ProtoNet CDA也随K增加而改善，"
            "但K=20时仍低于直接ADV3B02。"
        ),
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        anchor,
        "3.6.3不同target receiver下的结果",
        style="Heading 3",
        keep_with_next=True,
    )
    add_result_table_before(
        doc,
        anchor,
        by_receiver,
        fractions=(0.18, 0.205, 0.205, 0.205, 0.205),
        font_size=8.8,
        left_columns=(0,),
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "不同receiver的direct基线差异明显：3-19最困难，7-14的direct基线最高。"
            "MRIOR-SDA在20-1、7-7和8-8上获得明显正收益，但在7-14上略低于direct，"
            "说明高基线receiver继续更新时存在负迁移风险。"
        ),
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        anchor,
        "3.6.4结果边界",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "ProtoNet CDA只读取target-old support并闭式更新prototype；MRIOR-SDA和DADDA-SDA"
            "还读取封存source LEO弱信道标签缓存并更新完整backbone。因此三者可以比较机制与"
            "适应趋势，但权限和计算代价并不完全相同；这些结果不能直接证明support-only的"
            "Phase2主方法已经解决旧类域适应。"
        ),
        style="Body Text",
    )


def add_stage2c_formal_results(
    doc: _Document,
    anchor: Paragraph,
    formal: Sequence[Sequence[str]],
    header_templates: Sequence[Sequence],
) -> None:
    append_paragraph_before(
        doc,
        anchor,
        "4.5.1逐配置结果",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "下表直接采用《学习进展情况+7.24》中的正式LEO结果，按K-shot、方法和新类数"
            "逐项展开。旧类与新类query只用于冻结后评价；表内数值保持周报原口径，"
            "不再汇总为跨条件总均值，也不跨行拼接“最佳指标”。"
        ),
        style="Body Text",
    )
    formal_table = add_result_table_before(
        doc,
        anchor,
        formal,
        fractions=(0.09, 0.14, 0.10, 0.17, 0.17, 0.17, 0.16),
        font_size=8.2,
        left_columns=(1,),
        group_column=0,
    )
    for target_column, captured_content in zip(
        (3, 4, 5, 6),
        header_templates,
    ):
        apply_captured_cell_content(
            formal_table.rows[0].cells[target_column],
            captured_content,
            font_size=8.2,
        )

    append_paragraph_before(
        doc,
        anchor,
        "4.5.2主要现象",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        "",
        style="Body Text",
        segments=(
            ("低K、少新类时存在零注册。", True, TEXT_RED),
            (
                "CSIL在K=5/10且新类数为1或3时seen_new_acc为0；"
                "MoPC-HR在K=5或10、新类数为1时也为0。",
                False,
                None,
            ),
        ),
    )
    append_paragraph_before(
        doc,
        anchor,
        "",
        style="Body Text",
        segments=(
            ("MoPC-HR的可塑性更强，但会牺牲旧类。", True, TEXT_RED),
            (
                "K=20、新类数1时seen_new_acc为96.53%，H_old_new为72.69%；"
                "与此同时，old_acc_after降至60.76%，forgetting为26.71个百分点。",
                False,
                None,
            ),
        ),
    )
    append_paragraph_before(
        doc,
        anchor,
        "",
        style="Body Text",
        segments=(
            ("新类规模扩大后联合性能下降。", True, TEXT_RED),
            (
                "MoPC-HR在K=20时，新类数从1增加到25，seen_new_acc从96.53%降至27.86%，"
                "old_acc_after也从60.76%降至36.66%。",
                False,
                None,
            ),
        ),
    )
    append_paragraph_before(
        doc,
        anchor,
        "",
        style="Body Text",
        segments=(
            ("CSIL呈现明显的稳定性—可塑性极端。", True, TEXT_RED),
            (
                "小规模时旧类几乎完全保留但新类不注册；部分较大规模或高K切片中，"
                "新类开始被学习，却伴随旧类显著下降。",
                False,
                None,
            ),
        ),
    )


def add_stage2c_no_leo_results(
    doc: _Document,
    anchor: Paragraph,
    no_leo: Sequence[Sequence[str]],
    header_templates: Sequence[Sequence],
) -> None:
    append_paragraph_before(
        doc,
        anchor,
        "4.6.1逐配置结果",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "该诊断保持方法、物理样本ID、support/query划分、K-shot和旧类评测条件一致，"
            "只把新类support/query替换为未叠加LEO的同一物理记录。"
            "结果永久标记为DIAGNOSTIC_NEW_CLASS_NO_LEO_NON_FORMAL。"
        ),
        style="Body Text",
    )
    no_leo_table = add_result_table_before(
        doc,
        anchor,
        no_leo,
        fractions=(0.09, 0.14, 0.10, 0.20, 0.16, 0.16, 0.15),
        font_size=8.2,
        left_columns=(1,),
        group_column=0,
    )
    for target_column, captured_content in zip(
        (3, 4, 5, 6),
        header_templates,
    ):
        apply_captured_cell_content(
            no_leo_table.rows[0].cells[target_column],
            captured_content,
            font_size=8.2,
        )
    append_paragraph_before(
        doc,
        anchor,
        "Δ表示无LEO诊断相对正式LEO结果的变化。",
        style="Body Text",
    )

    append_paragraph_before(
        doc,
        anchor,
        "4.6.2诊断解读",
        style="Heading 3",
        keep_with_next=True,
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "移除新类LEO扰动后，MoPC-HR的新类准确率在多数切片显著提高，说明信道失真"
            "确实破坏了新类特征与旧类prototype之间的几何关系。但新类学得更好不一定提高"
            "H_old_new：新类梯度更强时，旧类准确率可能进一步下降。"
        ),
        style="Body Text",
    )
    append_paragraph_before(
        doc,
        anchor,
        (
            "例如K=10、新类数10时，无LEO使seen_new_acc提高42.48个百分点，"
            "但old_acc_after下降29.12个百分点，H_old_new反而下降12.28个百分点。"
            "该诊断只能用于信道归因，不能替代正式卫星场景结果或用于方法晋级。"
        ),
        style="Body Text",
    )


def revise_document(
    input_path: Path,
    weekly_716_path: Path,
    weekly_724_path: Path,
    output_path: Path,
) -> None:
    doc = Document(str(input_path))
    weekly_716 = Document(str(weekly_716_path))
    weekly_724 = Document(str(weekly_724_path))

    stage2b_overall = source_table(
        weekly_716,
        header=(
            "方法",
            "适应前old_acc",
            "适应后old_acc",
            "平均收益",
            "正/负迁移任务",
            "平均时延",
            "backbone更新",
        ),
        expected_rows=4,
    )
    stage2b_by_k = source_table(
        weekly_716,
        header=("K", "直接ADV3B02", "MRIOR-SDA", "DADDA-SDA", "ProtoNet CDA"),
        expected_rows=6,
    )
    stage2b_by_receiver = source_table(
        weekly_716,
        header=(
            "target receiver",
            "直接ADV3B02",
            "MRIOR-SDA",
            "DADDA-SDA",
            "ProtoNet CDA",
        ),
        expected_rows=6,
    )
    stage2c_formal = source_table(
        weekly_724,
        header=(
            "K-shot",
            "方法",
            "新类数",
            "old_acc_after",
            "seen_new_acc",
            "H_old_new",
            "forgetting",
        ),
        expected_rows=25,
    )
    stage2c_no_leo = source_table(
        weekly_724,
        header=(
            "K-shot",
            "方法",
            "新类数",
            "无LEO seen_new",
            "Δnew",
            "Δold",
            "ΔH",
        ),
        expected_rows=19,
    )

    aggregate_formal_table = find_formula_header_table(
        doc,
        column_count=7,
        math_signature=(
            "",
            "Aoldpre",
            "Aoldpost",
            "Anew",
            "Hold,new",
            "Fold",
            "Amin,old",
        ),
    )
    aggregate_no_leo_table = find_formula_header_table(
        doc,
        column_count=8,
        math_signature=(
            "",
            "Aoldpost",
            "Anew",
            "Hold,new",
            "ΔAoldpost",
            "ΔAnew",
            "ΔHold,new",
            "ΔFold",
        ),
    )
    appendix_stage2b_table = find_formula_header_table(
        doc,
        column_count=7,
        math_signature=("", "Aold", "Aold", "Gold", "", "", ""),
        word_signature=(
            "方法",
            "适应前",
            "适应后",
            "平均",
            "正/负迁移任务",
            "平均时延",
            "3场景backbone更新",
        ),
    )
    stage2b_header_templates = [
        capture_cell_content(appendix_stage2b_table.rows[0].cells[index])
        for index in (1, 2, 3)
    ]
    stage2c_formal_header_templates = [
        capture_cell_content(aggregate_formal_table.rows[0].cells[index])
        for index in (2, 3, 4, 5)
    ]
    stage2c_no_leo_header_templates = [
        capture_cell_content(aggregate_no_leo_table.rows[0].cells[index])
        for index in (2, 5, 4, 6)
    ]

    first_paragraph = doc.paragraphs[0]
    if clean_text(first_paragraph.text) != "1.核心概念与Phase2阶段定位":
        new_heading = doc.add_paragraph(style="Heading 1")
        run = new_heading.add_run("1.核心概念与Phase2阶段定位")
        style_run(run, color=HEADING_BLUE)
        new_heading.paragraph_format.keep_with_next = True
        first_paragraph._p.addprevious(new_heading._p)

    heading_replacements = {
        "1.2RFFI任务的三轴谱系": "1.1RFFI任务的三轴谱系",
        "1.4少样本学习的严格任务定义": "1.2少样本学习的严格任务定义",
        "1.5域适应及Stage2-B定位": "1.3域适应及Stage2-B定位",
        "1.6类增量、FSCIL与新类注册": "1.4类增量、FSCIL与新类注册",
        "1.7与CVS-RFFI Phase2阶段的对应": "1.5与CVS-RFFI Phase2阶段的对应",
        "1.8样本角色与成功条件": "1.6样本角色与成功条件",
        "2.3实验矩阵": "2.2实验矩阵",
        "2.4评价指标": "2.3评价指标",
        "2.5统一输入、状态更新与输出": "2.4统一输入、状态更新与输出",
        "旧类适应": "2.3.1旧类适应",
        "新类注册": "2.3.2新类注册",
        "旧新联合评价": "2.3.3旧新联合评价",
        "3.5机制、权限与结果位置": "3.5对比方法机制与权限对比",
        "3.2 ProtoNet CDA：原型式K-shot目标域校准基线": (
            "3.2ProtoNet CDA：原型式K-shot目标域校准基线"
        ),
        "4.5正式LEO弱信道结果": "4.5正式LEO弱信道逐配置结果",
        "4.6matched无LEO新类归因诊断": "4.6matched无LEO新类归因诊断",
        "参考文献": "5.参考文献",
    }
    for old_text, new_text in heading_replacements.items():
        paragraph = find_paragraph(doc, old_text)
        replace_paragraph_text(paragraph, new_text)

    experiment_table = find_table(doc, "工作包", "基座与数据")
    set_cell_text(
        experiment_table.rows[2].cells[2],
        "K∈{5,10,20}；新类数按两种方法的周报实验配置展开",
    )
    set_cell_text(
        experiment_table.rows[2].cells[3],
        "24个正式LEO逐配置结果",
    )
    set_cell_text(
        experiment_table.rows[3].cells[2],
        "与周报正式结果按配置配对",
    )
    set_cell_text(
        experiment_table.rows[3].cells[3],
        "18个matched无LEO逐配置结果",
    )
    stage2c_setup_table = find_table(doc, "环节", "CSIL")
    set_cell_text(
        stage2c_setup_table.rows[4].cells[1],
        "K=5、10、20；新类规模1、3、20（周报列出9个正式配置）",
    )
    set_cell_text(
        stage2c_setup_table.rows[4].cells[2],
        "K=5、10、20；新类规模1、3、5、10、25（周报列出15个正式配置）",
    )

    stage2c_heading = find_paragraph(doc, "4.Stage2-C：少样本类增量仿真实验")
    add_stage2b_results(
        doc,
        stage2c_heading,
        stage2b_overall,
        stage2b_by_k,
        stage2b_by_receiver,
        stage2b_header_templates,
    )

    formal_heading = find_paragraph(doc, "4.5正式LEO弱信道逐配置结果")
    no_leo_heading = find_paragraph(doc, "4.6matched无LEO新类归因诊断")
    remove_between(formal_heading, no_leo_heading)
    add_stage2c_formal_results(
        doc,
        no_leo_heading,
        stage2c_formal,
        stage2c_formal_header_templates,
    )

    low_k_heading = find_paragraph(doc, "4.7低K条件下的零注册问题")
    remove_between(no_leo_heading, low_k_heading)
    add_stage2c_no_leo_results(
        doc,
        low_k_heading,
        stage2c_no_leo,
        stage2c_no_leo_header_templates,
    )

    zero_step_paragraph = find_paragraph_startswith(
        doc,
        "CSIL还要先执行约60%的官方训练切分",
    )
    replace_paragraph_text(
        zero_step_paragraph,
        (
            "结果表中，CSIL在K=5/10且新类数为1或3时seen_new_acc为0；"
            "MoPC-HR在K=5或10、新类数为1时也为0。官方流程审计表明，"
            "固定batch和drop_last=True会使样本不足一个完整batch的切片不产生optimizer step。"
            "因此这些零准确率首先反映训练是否真正发生，不能解释成“充分训练后方法仍然无效”。"
        ),
    )
    next_paragraph = find_paragraph_startswith(
        doc,
        "低新类数并不必然意味着任务更简单",
    )
    replace_paragraph_text(
        next_paragraph,
        (
            "即使能够产生更新，低新类数也不意味着边界更容易学习：新增权重缺少新类间对比，"
            "旧类logit已经充分训练，而新类logit接近初始化，统一单头竞争仍可能偏向旧类。"
            "LEO弱信道还会放大support估计误差。"
        ),
    )

    appendix_heading = find_paragraph(doc, "附录A：非类增量对比方法复现实验结果")
    remove_from_heading_to_end(appendix_heading)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            if paragraph.style.name == "Heading 1":
                paragraph.paragraph_format.page_break_before = False
            for run in paragraph.runs:
                style_run(run)

    normalize_all_visible_run_fonts(doc)
    doc.core_properties.title = "CVS-RFFI Phase2阶段工作详细报告（标题与分阶段结果修订版）"
    doc.core_properties.subject = "Phase2概念、对比方法、分阶段仿真实验与逐配置结果"
    doc.core_properties.comments = (
        "标题序号已连续化；Stage2-B与Stage2-C结果表分别采用7.16和7.24周报原表。"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    check = Document(str(output_path))
    heading_texts = [
        clean_text(paragraph.text)
        for paragraph in check.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    required_headings = {
        "1.核心概念与Phase2阶段定位",
        "1.1RFFI任务的三轴谱系",
        "2.数据、权限与评价框架",
        "3.Stage2-B：跨接收机旧类域适应仿真实验",
        "3.6对比方法复现实验结果",
        "4.Stage2-C：少样本类增量仿真实验",
        "4.5正式LEO弱信道逐配置结果",
        "4.6matched无LEO新类归因诊断",
        "5.参考文献",
    }
    missing = sorted(required_headings.difference(heading_texts))
    if missing:
        raise RuntimeError(f"final document is missing headings: {missing}")
    if any(text.startswith("附录A") for text in heading_texts):
        raise RuntimeError("obsolete appendix remains in final document")

    final_matrices = [table_matrix(table) for table in check.tables]
    for expected in (
        stage2b_by_k,
        stage2b_by_receiver,
    ):
        if expected not in final_matrices:
            raise RuntimeError(f"source table not preserved exactly: {expected[0]}")
    for expected in (stage2b_overall, stage2c_formal, stage2c_no_leo):
        if not any(
            len(matrix) == len(expected)
            and matrix[1:] == expected[1:]
            for matrix in final_matrices
        ):
            raise RuntimeError(f"source table body not preserved exactly: {expected[0]}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Surgically revise the user-edited CVS-RFFI Phase2 DOCX: "
            "fix heading numbering and place detailed weekly-result tables "
            "after the corresponding Stage2-B/Stage2-C method comparisons."
        )
    )
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--weekly-716", type=Path, required=True)
    parser.add_argument("--weekly-724", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--repo-output", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    for source in (args.input, args.weekly_716, args.weekly_724):
        if not source.is_file():
            raise FileNotFoundError(source)

    revise_document(
        args.input,
        args.weekly_716,
        args.weekly_724,
        args.output,
    )
    if args.repo_output:
        args.repo_output.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.output, args.repo_output)

    print(f"input_sha256={sha256(args.input)}")
    print(f"weekly_716_sha256={sha256(args.weekly_716)}")
    print(f"weekly_724_sha256={sha256(args.weekly_724)}")
    print(f"output={args.output}")
    print(f"output_sha256={sha256(args.output)}")
    if args.repo_output:
        print(f"repo_output={args.repo_output}")
        print(f"repo_output_sha256={sha256(args.repo_output)}")


if __name__ == "__main__":
    main()
