from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from build_cvs_ntn_scenario_docx import (
    BLUE,
    add_body,
    add_bullet,
    add_callout,
    add_custom_numbering,
    add_heading,
    add_hyperlink,
    add_numbered,
    add_table,
    configure_styles,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "CVS天基射频指纹识别_NTN与核心场景_精简版_20260817.docx"
OUT = ROOT / "docs" / "CVS天基射频指纹识别_3GPP_NTN与Iridium_NTN_Direct场景_20260817.docx"


def add_reference(doc: Document, label: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    add_hyperlink(paragraph, label, url, color=BLUE, underline=True, size=9)


def build_document() -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    # Rebuild a clean Word package. The user's revision is kept untouched and
    # remains the content baseline, while the approved compact style is reapplied.
    doc = Document()
    configure_styles(doc)
    bullet_id, decimal_id = add_custom_numbering(doc)

    add_heading(doc, "1. 核心概念与研究对象", 1)
    add_heading(doc, "1.1 NTN与3GPP NTN", 2)
    add_body(
        doc,
        "NTN是Non-Terrestrial Network的缩写，中文通常译为“非地面网络”。它指利用卫星或高空平台等非地面节点承载中继、接入或基站功能的网络或网络片段。NTN不是某一颗卫星、某一个频段或某一种业务，而是一类网络体系。它把无线覆盖延伸到海洋、极地、沙漠、山区、航空航路和灾害区域等地面网络难以持续覆盖的位置。[1]",
    )
    add_body(
        doc,
        "3GPP NTN是由3GPP移动通信标准体系定义的NTN技术路线。它强调卫星接入与移动运营商的PLMN、5G核心网、移动性管理、认证、漫游和业务平台协同。Release 17首次形成规范性NTN支持，包含面向5G宽带接入的NR-NTN，以及面向低功耗物联网的IoT-NTN；后续版本继续增强覆盖、移动性、载荷和终端能力。[1][2]",
    )
    add_table(
        doc,
        ["组成", "定义", "与射频指纹场景的关系"],
        [
            ["UE／NTN终端", "地面、海上或移动平台上的用户设备，经过服务链路向卫星发送上行信号。", "终端中的具体物理发射RF链是待识别和注册的对象。"],
            ["服务链路", "UE与卫星载荷之间的无线链路。", "轨道运动、多普勒、仰角、衰落和噪声会改变接收I/Q的分布。"],
            ["星载载荷", "位于卫星上的转发或数字处理节点。", "其接收前端形成目标receiver domain；具备数字I/Q接口时可运行轻量推理。"],
            ["馈电链路与网关", "卫星与地面网关、基站功能和核心网之间的连接。", "可承载地面辅助推理、模型下发、审计和安全管理。"],
            ["PLMN与核心网", "完成订阅认证、接入控制、移动性和业务授权。", "为RFF注册提供可信逻辑身份和受控采集窗口，但不直接证明物理RF链身份。"],
        ],
        [1700, 3500, 4160],
        font_size=8.45,
    )
    add_body(
        doc,
        "透明转发载荷主要完成滤波、频率变换、放大和转发，数字基带处理通常位于地面；再生式载荷可在星上执行解调、译码、交换或基站功能。因而，“天基射频指纹识别”不等于所有计算都必须在卫星上完成：透明转发架构可在网关处理，再生式或专用数字感知载荷才适合直接在星上读取I/Q并执行推理。",
    )

    add_heading(doc, "1.2 Iridium NTN Direct", 2)
    add_body(
        doc,
        "Iridium NTN Direct是Iridium面向标准化直连终端推出的NTN服务路线。Iridium公开资料显示，该服务依托现有66颗LEO卫星组成的全球星座、全球协调的L频段和软件定义升级，面向基于3GPP Release 19的NB-IoT及Direct-to-Device（D2D）接入，并与移动运营商网络衔接。[3]",
    )
    add_body(
        doc,
        "截至2026年公开进展，Iridium已经完成NTN Direct双向在轨消息测试，并继续推进芯片合作伙伴的live on-air testing，目标是推进beta测试和商业服务准备。[4][5]这些事实说明，标准化低功耗终端经LEO星座直接接入的工程生态正在形成；它们并不表示Iridium已经部署射频指纹识别，也不表示其对外提供原始I/Q。本文把Iridium NTN Direct作为现实部署背景和研究对象示例，而不是既有RFFI产品案例。",
    )
    add_table(
        doc,
        ["概念", "本文采用的含义"],
        [
            ["LEO", "Low Earth Orbit，即低地球轨道。卫星相对地面快速运动，具有有限可见窗口和频繁服务切换。"],
            ["NB-IoT", "Narrowband Internet of Things，面向低功耗、低速率和大规模物联网终端的蜂窝技术。"],
            ["D2D", "Direct-to-Device，普通或标准化终端不依赖专用大型地面站，直接与卫星建立业务连接。"],
            ["L频段", "Iridium公开说明其NTN Direct使用全球协调的L频段；该频段是具体系统的无线资源，不等同于NTN本身。"],
        ],
        [1900, 7460],
        font_size=8.8,
    )

    add_heading(doc, "1.3 天基射频指纹识别", 2)
    add_body(
        doc,
        "射频指纹识别（Radio-Frequency Fingerprint Identification，RFFI）利用功率放大器、本振、I/Q调制器、DAC、滤波器等器件制造偏差和非线性在无线信号中留下的细微、相对稳定特征，识别具体物理发射RF链。它识别的是设备实例的物理实现，不是IMSI、账号、地址或应用身份等逻辑标识。",
    )
    add_body(
        doc,
        "天基射频指纹识别是指由卫星或星地融合接收系统采集地面终端上行信号，并据此判断其物理发射RF链是否与登记对象一致。观测到的I/Q同时包含发射机特征、传播信道、卫星接收机前端、轨道运动和时间状态，因此核心科学问题不是简单分类，而是从强接收域和链路扰动中提取可迁移的发射机特征。",
    )
    add_callout(
        doc,
        "安全定位",
        "射频指纹是对订阅认证和密码学认证的补充证据，不替代SIM／密钥、5G认证、协议完整性保护或运营审计。合理目标是发现“逻辑凭据正确但物理RF链不一致”、辅助受控注册，并在证据不足时输出defer或复核请求。",
    )
    add_table(
        doc,
        ["术语", "严格定义"],
        [
            ["逻辑身份", "由运营商或业务系统管理的订阅、账号、设备标识或授权记录。"],
            ["物理RF链", "一台具体终端中的发射硬件组合，是RFFI的类别对象。"],
            ["receiver domain", "由卫星接收前端、采样链、处理配置及其时间状态共同形成的接收分布。"],
            ["K-shot support", "在合法注册或校准窗口内，由每个类别提供的K个相互独立物理发送样本；不是从一次接收记录复制出的多个视图。"],
            ["query", "注册完成后的待识别观测，只用于独立推理，不用于更新模型、选择超参数或推断真实类别配额。"],
        ],
        [2150, 7210],
        font_size=8.75,
    )

    add_heading(doc, "2. 场景一：3GPP NTN中的可信物理RF链绑定", 1)
    add_heading(doc, "2.1 场景定义", 2)
    add_body(
        doc,
        "该场景面向接入3GPP NTN的UE。网络已经能够通过PLMN认证确认“谁持有合法订阅与密钥”，但仍需要研究“当前上行无线信号是否来自登记的那一套物理RF发射链”。CVS在正常网络认证之后增加一个受控、可审计的RFF注册或校准环节，把逻辑身份与物理RF链建立带版本和有效期的绑定。",
    )
    add_table(
        doc,
        ["要素", "详细定义"],
        [
            ["参与对象", "NTN UE、卫星接收载荷、网关／gNB、5G核心网、授权应用功能和RFF身份库。"],
            ["信号方向", "地面UE上行→服务链路→卫星接收机；必要时只把特征或判断结果经馈电链路送往地面。"],
            ["可信标签", "来自订阅认证、设备登记和运营方授权的registration episode，而不是由模型从query真值反推。"],
            ["注册样本", "在限定时间、频率、波束、挑战序列和receiver ID下采集的K个独立上行burst。"],
            ["识别输出", "已注册物理RF链、物理链不一致、低质量／证据不足defer，以及审计所需置信度和版本信息。"],
            ["部署位置", "具备数字I/Q和算力的星载载荷可本地运行；否则在可信网关运行，并把星载接收机标识纳入receiver domain。"],
        ],
        [2100, 7260],
        font_size=8.55,
    )

    add_heading(doc, "2.2 运行流程", 2)
    add_numbered(doc, "UE先完成3GPP网络规定的订阅认证和接入流程。", decimal_id)
    add_numbered(doc, "运营方确认设备登记、业务目的和授权范围，创建一次不可复用的RFF registration episode。", decimal_id)
    add_numbered(doc, "网络下发受控挑战和采集参数，卫星接收机记录K个独立物理发送样本及receiver、时间和链路质量元数据。", decimal_id)
    add_numbered(doc, "冻结的地面学习表征先处理目标接收域差异，再根据合法support建立或更新该物理RF链的紧凑注册状态。", decimal_id)
    add_numbered(doc, "后续每个query都独立面对全部已注册类别，不能利用query真值、真实old／new角色、类别配额或跨query全局重排。", decimal_id)
    add_numbered(doc, "当证据不足、质量门失败或物理链显著不一致时，系统输出defer或告警，并交由既有认证与运营流程复核。", decimal_id)

    add_heading(doc, "2.3 场景中的两类合法事件", 2)
    add_table(
        doc,
        ["事件", "变化", "CVS处理", "必须保持的边界"],
        [
            ["新星载接收域启用", "UE物理RF链不变，卫星接收前端或其状态改变。", "利用已登记旧终端的K-shot support校准receiver domain，保持旧类身份。", "只能声明接收域适配，不能把接收机变化误判成新终端。"],
            ["新UE或RF模组接入", "出现新的物理发射RF链，或原设备完成RF模组更换。", "在授权窗口采集fresh K-shot并注册新类，同时验证旧类保持。", "逻辑账号不自动等于新物理类别；历史query不能转作support。"],
        ],
        [2050, 2450, 2750, 2110],
        font_size=8.3,
    )

    add_heading(doc, "2.4 对3GPP NTN场景的意义", 2)
    for text in (
        "把网络层的“合法订阅者”与物理层的“具体发射RF链”区分开，并提供可审计的绑定证据。",
        "在凭据复制、设备替换、RF模组更换或维护记录遗漏时，为运营系统增加一条独立于协议标识的异常线索。",
        "利用3GPP已有的认证、授权和应用安全体系约束标签来源，降低RFF训练样本被错误标注、重放或中继污染的风险。",
        "把跨卫星接收域适配和新类注册纳入同一生命周期，而不是把一次闭集准确率当作长期身份能力。",
        "为星上本地、网关处理和ground-assisted三种部署方式提供统一任务定义，便于比较资源、时延和风险。",
    ):
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "3. 场景二：Iridium NTN Direct中的跨星适配与终端注册", 1)
    add_heading(doc, "3.1 场景定义", 2)
    add_body(
        doc,
        "该场景把Iridium NTN Direct作为3GPP标准化直连终端落地的具体部署背景。低功耗NB-IoT／D2D终端直接向Iridium LEO卫星发送上行信号；随着卫星运动，同一终端可能在不同时间由不同卫星或接收链观测。CVS要同时解决两个问题：同一终端跨receiver domain仍应保持身份一致，新加入或更换RF模组的终端又必须能用少量可信样本注册。",
    )
    add_table(
        doc,
        ["要素", "Iridium NTN Direct背景", "研究性CVS定义"],
        [
            ["网络基础", "66颗LEO卫星、L频段、软件定义升级和全球覆盖能力。[3]", "将不同卫星／接收链视为不同receiver domain，不假定分布天然一致。"],
            ["终端形态", "基于3GPP Release 19的NB-IoT和D2D设备生态。[3]", "把具体终端RF链作为类别，把运营登记作为授权来源。"],
            ["运行状态", "2026年已公开双向在轨测试，并继续开展合作芯片live testing。[4][5]", "仅据此论证场景现实性，不声称现网已提供I/Q或RFFI。"],
            ["链路动态", "LEO运动带来有限可见窗口、较强多普勒和卫星切换。", "要求身份特征对接收机与链路变化具有稳定性，并允许低质量defer。"],
            ["部署约束", "低功耗、大规模终端和星地协同网络。", "比较紧凑本地状态、网关辅助更新、时延、内存和回滚能力。"],
        ],
        [1700, 3700, 3960],
        font_size=8.35,
    )

    add_heading(doc, "3.2 子场景A：同一终端跨卫星接收域迁移", 2)
    add_body(
        doc,
        "终端物理RF链没有变化，但服务卫星、接收前端、观测仰角、瞬时多普勒和链路质量发生变化。系统应利用受控旧终端support建立少量receiver-specific校准状态，使同一终端在不同卫星上仍映射到同一物理身份。研究重点不是强迫所有卫星生成完全相同的embedding，而是在共享身份表示与局部接收域状态之间取得平衡。",
    )
    for text in (
        "同一终端跨卫星识别准确率及逐类最低表现；",
        "接收域适配前后旧类增益、负迁移比例和遗忘量；",
        "support数量K、卫星接收链差异和链路质量对结果的影响；",
        "receiver-specific状态字节、推理时延、更新时延和失败回滚。",
    ):
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "3.3 子场景B：新NTN Direct终端或RF模组注册", 2)
    add_body(
        doc,
        "新NB-IoT／D2D终端完成运营登记和正常网络认证后，在受控窗口执行K次独立上行发送。CVS利用这些fresh support注册新的物理RF链，并在不访问历史source IQ、不使用query真值的条件下，同时验证新类可识别性和旧类保持。如果终端只更新账号或业务配置而物理RF链未变，应更新逻辑别名；如果PA、本振或完整RF模组发生实质更换，则创建新的物理链版本。",
    )
    add_callout(
        doc,
        "关键区别",
        "跨卫星迁移解决“同一发射机、不同接收域”；新终端注册解决“同一接收域中加入新的物理类别”。两者必须分别评价，不能用新类收益掩盖旧类退化，也不能把接收域变化错误解释为新身份。",
    )

    add_heading(doc, "3.4 对Iridium NTN Direct场景的意义", 2)
    for text in (
        "为CVS提供一个具有公开星座、频段、终端类型和在轨测试进展的现实部署锚点，使研究问题不再停留于抽象卫星链路。",
        "66颗LEO卫星和持续移动的接收关系，使跨receiver domain泛化成为系统性需求，而不是一次性的实验扰动。",
        "标准化NB-IoT／D2D终端生态意味着设备数量和供应链会扩展，少样本新类注册及旧类保持具有明确的生命周期意义。",
        "软件定义星座适合讨论算法更新和星地协同，但任何星上部署结论仍需以实际I/Q接口、算力、内存、功耗和运营授权为前提。",
        "RFF不匹配可作为物理链异常证据，帮助发现设备替换或注册状态失配；最终处置仍由Iridium／MNO的认证、审计和安全策略决定。",
    ):
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "4. 天基射频指纹识别的研究意义", 1)
    add_heading(doc, "4.1 一般意义", 2)
    add_table(
        doc,
        ["研究价值", "具体含义"],
        [
            ["补充逻辑身份", "在账号、密钥和协议标识之外，验证当前无线发射是否与登记的物理RF链一致。"],
            ["研究跨接收域不变性", "把发射机、信道和星载接收前端的混合效应显式分离，推动真正跨卫星可迁移的表征学习。"],
            ["支持设备生命周期", "覆盖首次接入、跨星迁移、接收载荷变化、RF模组维修和新终端持续加入。"],
            ["降低原始数据搬运", "在具备条件时只在星上保留紧凑特征或注册状态，减少持续下传原始I/Q的需求。"],
            ["形成可审计决策", "把标签来源、support权限、模型更新、defer和回滚记录纳入运营流程。"],
        ],
        [2500, 6860],
        font_size=8.8,
    )

    add_heading(doc, "4.2 对当前CVS项目的意义", 2)
    add_body(
        doc,
        "这两类场景共同给出CVS的完整研究闭环：3GPP NTN提供标准化网络、认证与授权语境，回答可信标签和物理链绑定从哪里来；Iridium NTN Direct提供具体LEO星座、L频段、NB-IoT／D2D终端和跨星移动背景，回答为什么必须处理接收域迁移、少样本注册和资源约束。",
    )
    add_table(
        doc,
        ["CVS环节", "场景中的现实含义", "需要观察的结果"],
        [
            ["Phase1地面学习", "利用多个地面source receiver学习尽量保留发射机差异、抑制接收机和链路扰动的冻结表征。", "跨接收机稳定性、身份可分性、receiver leakage和可部署资源。"],
            ["Stage2-A直接部署", "模型首次面对训练阶段未见的目标星载receiver domain。", "无目标标签时的直接迁移风险；不把结果表述为未知确权。"],
            ["Stage2-B旧类校准", "3GPP NTN接收载荷启用或Iridium场景中的跨卫星接收域迁移。", "旧类准确率、逐类floor、负迁移、状态字节和时延。"],
            ["Stage2-C新类注册", "受认证的新3GPP NTN UE或新的Iridium NTN Direct终端／RF模组接入。", "seen-new准确率、旧／新调和指标、旧类遗忘和注册成本。"],
        ],
        [1900, 4350, 3110],
        font_size=8.25,
    )
    add_body(
        doc,
        "当前项目数据仍属于地面采集与LEO弱信道压力代理，能够研究receiver shift、链路扰动、K-shot适配和类增量注册，但不能直接证明真实Iridium卫星接收链上的性能。下一步若要提升场景证据等级，应依次补充目标频段多接收机采集、RF硬件在环、受控卫星接收数据和运营级授权流程，而不是仅扩大数字仿真规模。",
    )

    add_heading(doc, "5. 评价要求与声明边界", 1)
    add_heading(doc, "5.1 最小评价集合", 2)
    add_table(
        doc,
        ["维度", "必要指标或检查"],
        [
            ["识别能力", "旧类准确率、seen-new准确率、旧／新调和指标、逐类旧类floor。"],
            ["跨域效果", "适配前后同一注册集合上的差值、负迁移类别比例、跨receiver稳定性。"],
            ["遗忘与注册", "新类注册后旧类下降、注册成功率、K敏感性和物理链版本一致性。"],
            ["可信协议", "support与query物理样本隔离、query零更新、全注册类独立竞争、无真值／角色／配额泄漏。"],
            ["部署资源", "模型大小、峰值RAM、紧凑状态字节、推理与更新时间、通信量和回滚能力。"],
            ["安全输出", "accept、物理链不一致、defer的定义，以及与网络认证和人工复核的交接方式。"],
        ],
        [2500, 6860],
        font_size=8.65,
    )

    add_heading(doc, "5.2 可以与不可以声明的内容", 2)
    add_table(
        doc,
        ["可以声明", "不可以声明"],
        [
            ["围绕3GPP NTN架构研究逻辑身份与物理RF链的补充绑定机制。", "RFFI已经成为3GPP NTN标准规定的认证功能。"],
            ["以Iridium NTN Direct的公开架构与进展说明LEO直连终端场景具有现实性。", "Iridium已经部署CVS／RFFI，或会向研究系统开放原始I/Q。"],
            ["在地面代理和LEO弱信道压力条件下评价跨接收域适配与K-shot注册。", "已经完成真实Iridium星座或其他在轨系统验证。"],
            ["把RFF结果作为异常线索、注册证据或复核触发条件。", "单凭一次RFF判断即可替代密码学认证、停用终端或完成法律归因。"],
        ],
        [4680, 4680],
        font_size=8.55,
    )
    add_callout(
        doc,
        "核心结论",
        "3GPP NTN给出可信接入与授权框架，Iridium NTN Direct给出真实LEO直连终端生态。CVS的研究意义是在两者交汇处维护“逻辑身份—具体物理RF链”的可信绑定，并使该绑定能够经受跨卫星接收域变化、少样本新终端注册和长期设备演化。",
    )

    add_heading(doc, "参考资料", 2)
    add_reference(doc, "[1] 3GPP，Non-Terrestrial Networks (NTN) overview。", "https://www.3gpp.org/technologies/ntn-overview")
    add_reference(doc, "[2] 3GPP，Release 17。", "https://www.3gpp.org/specifications-technologies/releases/release-17")
    add_reference(doc, "[3] Iridium，Iridium NTN Direct。", "https://www.iridium.com/services/iridium-ntn-direct")
    add_reference(doc, "[4] Iridium，On-Air Trials Underway: Iridium NTN Direct Prepares to Enter Beta，2026-01-21。", "https://investor.iridium.com/2026-01-21-On-Air-Trials-Underway-Iridium-NTN-Direct-Prepares-to-Enter-Beta-as-Testing-Continues")
    add_reference(doc, "[5] Iridium，Iridium NTN Direct Begins Live Testing with Mlink，2026-06-24。", "https://investor.iridium.com/2026-06-24-Iridium-NTN-Direct-Begins-Live-Testing-with-Mlink")

    doc.core_properties.title = "CVS天基射频指纹识别：3GPP NTN与Iridium NTN Direct场景"
    doc.core_properties.subject = "两类核心NTN场景、定义、研究意义与证据边界"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        from docx.oxml import OxmlElement

        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def structural_check(path: Path) -> None:
    if not path.exists() or path.stat().st_size < 35_000:
        raise AssertionError("DOCX output is missing or unexpectedly small")
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        xml = archive.read("word/document.xml").decode("utf-8")
        for required in (
            "3GPP NTN中的可信物理RF链绑定",
            "Iridium NTN Direct中的跨星适配与终端注册",
            "天基射频指纹识别的研究意义",
            "NTN是Non-Terrestrial Network",
            "可以与不可以声明的内容",
        ):
            if required not in xml:
                raise AssertionError(required)
        for forbidden in (
            "TT&amp;C",
            "DCP",
            "GOES",
            "HawkEye",
            "ADS-B",
            "AIS",
            "PhiSat",
            "PAST-AI",
            "SatIQ",
            "WiSig",
            "ManySig",
            "Phase3",
            "九类现实研究场景",
        ):
            if forbidden in xml:
                raise AssertionError(f"unrelated content remains: {forbidden}")
        if any(name.startswith("word/header") for name in names):
            raise AssertionError("header remains")
        if any(name.startswith("word/footer") for name in names):
            raise AssertionError("footer remains")
        if "cite" in xml:
            raise AssertionError("internal web citation marker remains")


if __name__ == "__main__":
    output = build_document()
    structural_check(output)
    print(output)
