from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document

from revise_phase2_report_numbering_and_results import (
    append_paragraph_before,
    add_result_table_before,
    find_paragraph,
    find_paragraph_startswith,
    normalize_all_visible_run_fonts,
    replace_paragraph_text,
)


REQUIREMENT_TABLE = [
    ["要求", "最低条件", "不满足时的处理"],
    [
        "目标域可观测",
        "能够从目标卫星接收链读取固定数字I/Q，并记录receiver ID、时间、频率和链路质量。",
        "透明转发载荷或无数字I/Q接口时，转到可信网关处理，不能声称星上适应。",
    ],
    [
        "可信support",
        "每类K个独立物理发送事件，来自认证、授权且带challenge／nonce的registration episode。",
        "标签来源、独立性或质量门失败时拒绝注册；历史query不得转作support。",
    ],
    [
        "算法权限",
        "更新只读取冻结bundle和合法support；query零更新，并对全部已注册类逐样本统一竞争。",
        "禁止使用query真值、old／new角色、类别配额、跨query重排或测试反馈调参。",
    ],
    [
        "资源与可靠性",
        "报告模型大小、峰值RAM、状态字节、推理／更新时间、能耗、原子提交和回滚点。",
        "超出星载预算时采用ground-assisted方案；不得用资源不可行的方法支撑星上部署结论。",
    ],
    [
        "安全与审计",
        "模型、注册凭证和状态更新可签名、可追溯、带有效期，并能抵抗重放和错误标签。",
        "发现冲突时保持旧状态、输出defer并转入运营复核。",
    ],
    [
        "证据等级",
        "地面代理、硬件在环和受控在轨数据分层报告，明确接收链与信道来源。",
        "当前地面数据与LEO压力代理不能表述为真实Iridium在轨验证。",
    ],
]


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_percent_units(doc) -> None:
    definition = find_paragraph_startswith(doc, "为避免符号、状态和结果字段混在一个段落中")
    replace_paragraph_text(
        definition,
        definition.text.replace(
            "差值以百分点（percentage points，pp）报告",
            "差值也以%报告",
        ),
    )

    for paragraph in iter_paragraphs(doc):
        if paragraph.text == "":
            continue
        changed = False
        for run in paragraph.runs:
            original = run.text
            updated = re.sub(r"(?<=\d)pp\b", "%", original)
            updated = re.sub(r"([+−-]?\d+(?:\.\d+)?)个百分点", r"\1%", updated)
            if updated != original:
                run.text = updated
                changed = True
        if changed:
            continue


def remove_empty_paragraphs_between(start, end) -> None:
    parent = start._p.getparent()
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        if node.tag.endswith("}p"):
            text = "".join(node.itertext()).strip()
            if not text:
                parent.remove(node)
        node = next_node


def repair(source: Path, output: Path) -> None:
    doc = Document(source)

    replace_percent_units(doc)

    p150 = find_paragraph_startswith(doc, "以上系统扩大了无线信号的观测范围")
    replace_paragraph_text(
        p150,
        (
            "以上系统扩大了无线信号的观测范围，却没有自动解决当前信号究竟由哪一部物理射频设备发出。"
            "密码学认证验证的是密钥或凭据，协议身份表示的是逻辑或行政身份，地理定位回答的是发射机位于何处，"
            "调制识别判断的是信号类型。逻辑凭据和网络地址均不必然与当前射频前端一一对应：凭据可能被复制，"
            "网络身份可能被错误配置、重放或冒用，合法设备也可能在账号不变的情况下更换功放、本振、应答机或整套射频模组。"
        ),
    )

    p153 = find_paragraph_startswith(doc, "NTN注册运营方主动创建")
    replace_paragraph_text(
        p153,
        (
            "NTN注册是运营方主动创建的受控物理RF注册事件。终端先完成正常网络认证；运营方核对逻辑身份、"
            "设备记录和授权范围，再下发带时间窗、频率、波束、receiver ID和challenge／nonce的registration episode。"
            "目标接收链采集每类K个独立物理发送事件，经过SNR、AGC、clipping、同步和重复样本检查后，才能成为support。"
        ),
    )

    p165 = find_paragraph_startswith(doc, "计算位置不必绝对限定在卫星内部")
    if p165.text.endswith("能力。。"):
        replace_paragraph_text(p165, p165.text[:-1])

    heading6 = find_paragraph(doc, "6.参考文献")
    remove_empty_paragraphs_between(p165, heading6)
    append_paragraph_before(
        doc,
        heading6,
        "星载目标域预适应／类注册的最低实施要求如下；不满足时应转为可信网关处理、拒绝注册或输出defer，不能把代理实验写成星上部署结论。",
        style="Body Text",
    )
    add_result_table_before(
        doc,
        heading6,
        REQUIREMENT_TABLE,
        fractions=[0.18, 0.45, 0.37],
        font_size=7.7,
        left_columns=(0, 1, 2),
        group_column=0,
    )
    append_paragraph_before(
        doc,
        heading6,
        (
            "当前报告中的地面OTA数据和LEO弱信道变换只用于验证算法在目标域压力下的行为，尚不能证明真实"
            "Iridium或其他在轨接收链上的性能。形成星上部署结论还需要目标频段多接收机采集、RF硬件在环、"
            "受控在轨I/Q、飞行软件资源测量以及运营方授权／回滚闭环。"
        ),
        style="Body Text",
    )

    normalize_all_visible_run_fonts(doc)
    doc.core_properties.title = "CVS-RFFI Phase2详细复现报告（图表与场景优化版·修复）"
    doc.core_properties.comments = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    repair(args.source, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
