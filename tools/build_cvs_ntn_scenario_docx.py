from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CVS天基射频指纹识别_NTN与场景设计_20260817.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
GOLD = "A67C00"
GRAY = "5A6470"
LIGHT_BLUE = "EAF1F8"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E1E8"
BLACK = "111111"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    assert sum(widths_dxa) == 9360, widths_dxa
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, latin="Calibri", east_asia="Microsoft YaHei", size=11,
                 color=BLACK, bold=None, italic=None) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(cell, text: str, bold=False, color=BLACK, size=9.2,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第")
    set_run_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GRAY)
    r_pr.append(color)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    run = paragraph.add_run("页")
    set_run_font(run, size=9, color=GRAY)


def add_custom_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    used_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    used_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(used_abs or [0]) + 20
    bullet_num_id = max(used_num or [0]) + 20
    decimal_abstract_id = abstract_id + 1
    decimal_num_id = bullet_num_id + 1

    def abstract(num_id: int, fmt: str, text: str, left: int, hanging: int):
        abs_el = OxmlElement("w:abstractNum")
        abs_el.set(qn("w:abstractNumId"), str(num_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_el.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)
        if fmt == "bullet":
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), "Symbol")
            r_fonts.set(qn("w:hAnsi"), "Symbol")
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abs_el.append(lvl)
        numbering.append(abs_el)

    def num(num_id: int, abstract_num_id: int):
        num_el = OxmlElement("w:num")
        num_el.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_num_id))
        num_el.append(abstract_ref)
        numbering.append(num_el)

    abstract(abstract_id, "bullet", "", 540, 280)
    num(bullet_num_id, abstract_id)
    abstract(decimal_abstract_id, "decimal", "%1.", 540, 280)
    num(decimal_num_id, decimal_abstract_id)
    return bullet_num_id, decimal_num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_body(doc: Document, text: str, bold_prefix: str | None = None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str, bullet_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, bullet_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)


def add_numbered(doc: Document, text: str, decimal_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, decimal_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(label + "：")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int = 1, page_break=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              font_size=8.7) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        set_cell_shading(hdr.cells[i], DARK_BLUE)
        set_cell_text(hdr.cells[i], text, bold=True, color=WHITE, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for i, text in enumerate(row_data):
            if len(table.rows) % 2 == 1:
                set_cell_shading(row.cells[i], LIGHT_GRAY)
            set_cell_text(row.cells[i], text, size=font_size,
                          align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document() -> Path:
    doc = Document()
    configure_styles(doc)
    bullet_id, decimal_id = add_custom_numbering(doc)
    section = doc.sections[0]

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("CVS-RFFI／CV-SincNet研究场景设计")
    set_run_font(r, size=9, color=GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("研究场景白皮书")
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("CVS天基射频指纹识别")
    set_run_font(r, size=29, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("NTN基础、现实机制与分阶段研究场景设计")
    set_run_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(78)
    r = p.add_run("地面弱标注跨接收机学习·目标星载接收域适配·授权新类注册·开放世界扩展")
    set_run_font(r, size=10.5, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("适用对象：论文读者、项目评审者与后续研发人员")
    set_run_font(r, size=10.5, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("版本日期：2026年8月17日")
    set_run_font(r, size=10.5, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("定位：研究场景与证据边界说明，不替代项目协议或实验报告")
    set_run_font(r, size=9.5, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_body(doc, "CVS研究天基射频指纹识别中的地面弱标注跨接收机域泛化、目标星载接收域少样本适配与授权新类注册，并把未知拒识、跨卫星匿名关联和可信确权作为后续开放世界扩展。项目关注的不是固定接收机上的闭集分类，而是当发射机身份、接收机前端、星地传播、时间状态和注册类别库共同变化时，如何维持可迁移、可扩展且可审计的物理RF身份空间。")
    add_body(doc, "本报告首先说明非地面网络（Non-Terrestrial Network，NTN）的定义、组成、轨道和载荷类型，再给出受控NTN终端注册、星载接收机commissioning、TT&C上行、环境DCP、跨卫星切换、维修重注册、未知干扰调查、ADS-B/AIS辅助核验及卫星下行指纹等九类场景。所有场景均予以保留，但按照当前Phase1/Phase2主线、Phase3计划和外部旁证分层，避免把不同信号方向、标签来源和完成状态混为同一任务。")
    add_callout(doc, "核心定位", "CVS不是从太空自动给未知设备贴标签，而是把网络认证和运营登记形成的逻辑身份信任延伸到具体物理RF发射链，并在接收域变化和身份库扩展时维持可验证的一致性。")

    add_heading(doc, "目录", 1)
    toc = [
        "1. NTN是什么",
        "2. NTN为什么形成新的RFFI问题",
        "3. CVS总体研究场景与对象定义",
        "4. Phase1、Stage2-A/B/C与Phase3关系",
        "5. 九类现实研究场景",
        "6. 为什么需要星上计算",
        "7. 统一科学问题与可观测效果",
        "8. 数据证据阶梯与声明边界",
        "9. 结论",
        "参考资料",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=NAVY)

    add_heading(doc, "1. NTN是什么", 1)
    add_heading(doc, "1.1 基本定义", 2)
    add_body(doc, "NTN是Non-Terrestrial Network的缩写，中文通常译为“非地面网络”。它不是某一种卫星、某一个频段或某一套具体通信协议，而是使用卫星、高空平台等非地面节点提供无线接入、传输或网络功能的一类网络体系。3GPP将NTN定义为利用空基或天基载体搭载中继节点或基站功能的网络或网络片段。其目的，是把移动通信和物联网覆盖扩展到海洋、沙漠、极地、山区、航空和灾害区域等地面基础设施难以覆盖的空间。[1]")
    add_body(doc, "与传统卫星专网不同，3GPP NTN强调与5G／移动运营网络的结合。终端可以通过标准化协议接入卫星网络，并与地面PLMN、核心网、应用功能、运营管理和安全体系协同。因此，NTN不仅是“卫星链路”，还包括终端、服务链路、星载载荷、馈电链路、网关、基站功能、核心网和运营控制等完整系统。")

    add_heading(doc, "1.2 NTN的基本组成", 2)
    add_table(doc,
              ["组成", "作用", "与CVS的关系"],
              [
                  ["NTN终端／UE", "地面、海上、空中或移动平台上的通信设备，向卫星发送上行信号并接收下行信号。", "在CVS上行主场景中，具体物理发射RF链是身份类别。"],
                  ["服务链路", "终端与卫星载荷之间的无线链路。", "引入多普勒、时延、仰角、衰落和信噪比变化。"],
                  ["NTN载荷", "卫星上执行RF转发或数字基带处理的载荷。", "决定CVS模型是否能在星上直接读取I/Q并运行。"],
                  ["馈电链路", "卫星与地面网关之间的链路。", "影响原始I/Q能否下传，以及地面辅助适配是否可行。"],
                  ["NTN网关", "连接星载载荷与地面基站／核心网。", "可承载地面辅助推理、模型签名、审计和更新服务。"],
                  ["基站与核心网", "完成接入控制、认证、移动性、业务和安全管理。", "可提供可信逻辑身份、授权窗口和注册凭证。"],
              ],
              [1600, 3340, 4420], font_size=8.2)

    add_heading(doc, "1.3 LEO、MEO和GEO", 2)
    add_table(doc,
              ["轨道", "典型特征", "对CVS的影响"],
              [
                  ["LEO低地球轨道", "轨道较低、传播时延较小、卫星相对地面高速运动，需要星座和频繁切换。", "多普勒变化强、可见窗口有限、终端会被不同接收机先后观测，是CVS当前主场景。"],
                  ["MEO中地球轨道", "高度和传播时延介于LEO与GEO之间，常用于导航或部分通信系统。", "接收域变化仍存在，但切换频率和链路动态与LEO不同。"],
                  ["GEO地球静止轨道", "相对地面位置近似固定，覆盖范围大、传播时延较高。", "更适合计划式DCP或固定终端场景，动态切换压力弱于LEO。"],
              ],
              [1700, 3530, 4130], font_size=8.9)

    add_heading(doc, "1.4 透明转发载荷与再生式载荷", 2)
    add_body(doc, "3GPP明确区分两类NTN载荷。[1]")
    add_bullet(doc, "透明转发载荷（transparent／bent-pipe payload）主要对上行RF信号进行滤波、频率变换、放大和下行转发，基站和主要数字处理功能位于地面。它不能被默认认为具有运行CVS模型的星上数字I/Q接口。", bullet_id)
    add_bullet(doc, "再生式载荷（regenerative payload）除RF处理外，还具有解调、译码、交换、路由、重新编码或调制等星上数字处理能力。CVS若要在卫星本地读取I/Q、提取embedding或执行轻量适配，应明确限定为再生式通信载荷或专用数字频谱感知载荷。", bullet_id)
    add_callout(doc, "场景前提", "CVS不假设所有NTN卫星都能运行射频指纹模型。透明转发系统可把模型放在网关或地面处理中心；“星上CVS”只对具备数字I/Q和边缘计算能力的载荷成立。")

    add_heading(doc, "1.5 NTN与传统地面网络的主要差异", 2)
    add_table(doc,
              ["差异", "NTN表现", "产生的RFFI问题"],
              [
                  ["高速相对运动", "LEO卫星快速经过服务区域。", "大多普勒、频偏变化率和有限观测窗口会掩盖硬件特征。"],
                  ["频繁服务切换", "同一终端可能由不同卫星或波束先后服务。", "接收机改变，但发射机身份不变，需要跨接收域一致性。"],
                  ["链路时延与中断", "馈电链路可能高时延、间歇或暂时不可用。", "推动星上本地推理、缓存和受约束更新。"],
                  ["覆盖范围大", "单个卫星接收节点可观测广域终端。", "类别规模、开放世界未知和隐私治理压力增加。"],
                  ["载荷资源受限", "算力、内存、功耗、热设计和可靠性受约束。", "模型必须报告状态字节、时延、能耗和回滚能力。"],
              ],
              [1800, 3200, 4360], font_size=8.9)

    add_heading(doc, "2. NTN为什么形成新的RFFI问题", 1)
    add_body(doc, "传统射频指纹识别常假设训练和测试来自相同或相近的接收机、信道和采集时间。NTN部署打破了这一假设：地面训练数据无法穷尽未来星载接收前端、实际轨道几何、温度状态和链路处理。目标卫星接收到的I/Q不是纯粹的发射机指纹，而是发射机、传播、接收机和处理链共同作用的结果。")
    add_callout(doc, "观测模型", "x=Pq[Rr,tau(Hell,tau(Ty(s)))]+n。Ty是身份来源；Rr、Hell、tau和Pq均可能让同一发射机在不同卫星上呈现不同统计分布。")
    add_body(doc, "WiSig包含174个发射机、41台USRP接收机、约1000万数据包和跨一个月的4次采集。其公开结果表明，仅改变接收机或采集日期就可能显著降低分类性能[5]。这意味着固定接收机上的高闭集准确率不能直接外推为天基跨接收机部署能力。")
    add_body(doc, "NTN同时带来类别库演化。新IoT终端、地面站、RF模组和维修后设备不断加入网络，模型不仅要适应新的接收域，还要注册新的物理发射链并保持旧类能力。因此，CVS把跨域泛化、部署期少样本适配和受控类增量注册组织为同一生命周期中的不同阶段。")

    add_heading(doc, "3. CVS总体研究场景与对象定义", 1)
    add_heading(doc, "3.1 总体场景", 2)
    add_body(doc, "CVS采用“地面训练、目标域部署”的分层机制。Phase1利用多个source receiver的有限标签和大量无标签信号，学习尽量保留发射机硬件差异、抑制接收机和信道扰动的开放世界就绪表征，并封存不可变deployment bundle。Phase2把该bundle部署到训练阶段未见的目标星载接收域，只读取固定LEO弱信道接收IQ、合法K-shot support和注册类别表，完成旧类校准与新类注册。Phase3处理标签尚不存在时的未知拒识、匿名关联和可信确权。")
    add_heading(doc, "3.2 身份对象", 2)
    add_bullet(doc, "逻辑或运营身份a：订阅、账号、呼号、ICAO地址、MMSI、DCP地址或任务登记项。", bullet_id)
    add_bullet(doc, "物理发射RF链y：具体设备实例中的PA、本振、I/Q调制器、DAC及相关硬件组成，是RFFI类别。", bullet_id)
    add_bullet(doc, "接收机前端r：目标卫星、地面站或网关中的LNA、滤波器、混频器、振荡器、AGC和ADC链路。", bullet_id)
    add_bullet(doc, "传播与轨道状态ell：仰角、多普勒、衰落、SNR、遮挡和多径。", bullet_id)
    add_bullet(doc, "时间与硬件状态tau：温度、电压、功率循环、老化、维修和器件更换。", bullet_id)
    add_body(doc, "逻辑身份与物理RF链不是永久一一对应关系。账号保持不变时，RF模组可能已经更换；同一物理设备也可能更新逻辑别名。因此，CVS应维护带有效期的映射，而不是把账号直接当作永恒的射频类别。")

    add_heading(doc, "4. Phase1、Stage2-A/B/C与Phase3关系", 1)
    add_table(doc,
              ["阶段", "现实事件", "可用信息", "任务与边界"],
              [
                  ["Phase1", "地面模型开发与bundle冻结", "多source receiver的少量有标签和大量无标签数据", "学习跨接收机表征；不执行部署期新类注册或真实多星协同。"],
                  ["Stage2-A", "模型首次进入目标星载接收域", "无target发射机标签，可有固定接收IQ", "测量直接迁移风险；不是完整未知拒识。"],
                  ["Stage2-B", "新接收机commissioning、载荷更换或跨星切换", "已登记旧类合法K-shot", "校准残余目标域偏移；只声明旧类适配。"],
                  ["Stage2-C", "授权新终端、RF模组更换或新地面站接入", "旧类和新类合法K-shot", "注册新物理RF链并保持旧类；标签是输入前提。"],
                  ["Phase3", "无标签未知信号首次出现", "冻结本地证据和合法外部上下文", "未知拒识、匿名关联和可信确权；属于后续开放世界研究。"],
                  ["Phase3交接", "外部确权并授权注册", "注册凭证和重新采集的fresh K-shot", "生成新注册事件；历史unknown query不能转为support。"],
              ],
              [1300, 2250, 2580, 3230], font_size=8.5)
    add_body(doc, "Phase2正式评价必须保持query零更新：每个query独立面对全部已注册类别，不能利用query真值、真实old/new角色、类别配额或跨query全局重排。数据协议中的LEO弱信道观测是物理启发压力代理，而不是对真实在轨接收链的复制。")

    add_heading(doc, "5. 九类现实研究场景", 1, page_break=True)
    add_heading(doc, "5.1 新星载接收机启用后的旧终端校准", 2)
    add_body(doc, "卫星运营方部署新卫星、更换接收载荷，或将已有终端切换到另一颗卫星。终端物理RF链没有改变，变化的是接收机前端、轨道几何和信号处理链。运营方在commissioning或计划维护窗口安排已登记旧终端执行受控参考发送，目标卫星采集每类K个独立burst，Stage2-B据此校准残余接收域偏移。")
    for text in [
        "Phase1表征能否减少未见接收域偏移；",
        "少量旧类support能否提升全部旧类而非少数类别；",
        "目标域适配能否在不访问source IQ和query真值的条件下完成；",
        "更新是否满足星上状态字节、时延和回滚约束。",
    ]:
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "5.2 受认证NTN新终端的物理RF链注册", 2)
    add_body(doc, "这是Stage2-C最具有规模化和产业说服力的场景。Iridium NTN Direct公开说明其标准化NB-IoT／D2D服务利用66颗LEO卫星、L-band频谱和软件定义升级，并在2026年开展on-air试验[3][4]。这些材料证明运营方管理的大规模NTN终端生态正在形成，但不证明Iridium正在运行RFFI或开放原始I/Q。", after=12)
    add_numbered(doc, "新终端完成订阅登记和5G主认证。", decimal_id)
    add_numbered(doc, "运营方或应用功能授权研究性的RFF registration episode。", decimal_id)
    add_numbered(doc, "系统分配严格的时间、频率、波束、挑战序列和receiver ID。", decimal_id)
    add_numbered(doc, "目标卫星采集K个相互独立的上行burst并执行质量门控。", decimal_id)
    add_numbered(doc, "审计系统建立逻辑身份与具体物理RF链的绑定凭证。", decimal_id)
    add_numbered(doc, "Stage2-C注册新类，并在全部旧类和新类上执行独立query推理。", decimal_id)
    add_body(doc, "3GPP AKMA允许应用功能复用PLMN已经完成的UE认证和安全密钥体系[2]。CVS可以据此设计RFF Enrollment Application Function，但必须明确：AKMA提供安全架构参考，RFF注册窗口和K-shot采集是CVS提出的研究扩展，不是现有3GPP标准功能。")

    add_heading(doc, "5.3 TT&C地面站及任务上行RF链管理", 2)
    add_body(doc, "TT&C类别规模较小，但标签和发送过程最容易审计。ESA Estrack在通信pass前按计划配置地面站，负责telecommand uplink、telemetry reception和测距[6]。CVS可用于新地面站RF链投入运行、HPA或本振更换、维修后物理一致性核验，以及认证凭据正常但RF链发生变化时的增强检查。")
    add_callout(doc, "边界", "RFFI不能替代telecommand密码认证。它提供的是“当前无线发射是否仍来自登记RF链”的补充证据。")

    add_heading(doc, "5.4 环境监测DCP与远程IoT平台注册", 2)
    add_body(doc, "NOAA GOES Data Collection System允许地面平台按照预定义频率、时间表、事件触发或查询命令发送环境数据[7]。类似系统中的平台地址、发送计划、设备认证和维护记录能够提供Stage2-C标签来源。CVS可研究新DCP接入、RF模组更换、逻辑地址冒用和无人值守设备长期漂移。DCP地址只证明逻辑标识，RFFI负责补充物理RF链一致性。")

    add_heading(doc, "5.5 跨卫星切换中的持续物理身份一致性", 2)
    add_body(doc, "LEO运动使同一终端先后被不同卫星观测。终端身份保持不变，但接收域发生变化。CVS应研究共享身份空间和receiver-specific局部状态的组合，而不是要求所有卫星产生完全相同的embedding。系统还需区分接收域变化、终端温度变化、RF模组更换和低质量观测；证据不足时应输出defer，而不是强制硬分类。")

    add_heading(doc, "5.6 维修、更换和长期漂移下的重新注册", 2)
    add_table(doc,
              ["现实事件", "发射类变化", "接收域变化", "处理阶段"],
              [
                  ["同一终端切换到另一颗卫星", "否", "是", "Stage2-B"],
                  ["目标卫星更换接收载荷", "否", "是", "新的Stage2-B校准事件"],
                  ["新增一台物理终端", "是", "否", "Stage2-C"],
                  ["终端更换PA、本振或RF模组", "是", "否", "Stage2-C重新注册"],
                  ["只修改账号、呼号或协议地址", "通常否", "否", "更新逻辑别名"],
                  ["温度变化、功率循环或器件老化", "通常不立即改变", "否", "漂移监测、校准或重新注册"],
              ],
              [3400, 1600, 1600, 2760], font_size=8.6)

    add_heading(doc, "5.7 未知干扰源发现与跨时空匿名关联", 2)
    add_body(doc, "该场景属于Phase3，不是当前Stage2-C的标签来源。当卫星首次捕获无法由注册身份库解释的信号时，系统只能输出unknown、anonymous entity或defer。合法生命周期为：本地拒识→跨卫星／跨过境匿名关联→TDOA／FDOA或其他几何定位→结合协议、登记、轨迹和现场调查→形成可信身份与注册授权→重新采集fresh K-shot→Stage2-C正式注册。")
    add_body(doc, "ITU-R SM.2355-2讨论了卫星频谱监测中的TDOA和FDOA定位[8]。HawkEye 360公开的NOAA GOES干扰案例展示了多次天基采集、逐步缩小区域和地面现场处置的业务链[9]。该案例支持“天基发现与定位—地面确权”的现实性，但不能被写成RFFI自动输出真实身份的算法证据。")

    add_heading(doc, "5.8 ADS-B、AIS和安全关键无线系统的辅助核验", 2)
    add_body(doc, "CVS可为航空器、船舶和遇险信标提供异常提示或物理一致性证据，例如ICAO地址不变但应答机已经更换、AIS中的MMSI与历史物理RF链不一致，或同一干扰源在不同海域重复出现。安全关键系统不能依据一次RFF不匹配直接丢弃消息或拒绝设备；合理输出应包括accept、defer和high risk，并与协议身份、雷达、SAR、光学、位置和维护记录联合使用。")

    add_heading(doc, "5.9 卫星下行发射机识别作为旁证研究", 2)
    add_body(doc, "PAST-AI和SatIQ使用真实Iridium下行数据研究卫星发射机认证，说明硬件相关特征可以穿过真实卫星传播链并被地面接收机观测[11][12]。但它们研究的是“卫星发射机→地面接收机”，而CVS当前Phase2主场景是“地面终端→目标星载接收机”。两者的发射对象、接收域和部署约束不同，因此只能作为真实星地链路中存在可观测硬件指纹的外部旁证，不能与当前上行任务共享同一套类别语义。")

    add_heading(doc, "5.10 场景总表", 2)
    add_table(doc,
              ["场景", "阶段", "标签来源", "价值与边界"],
              [
                  ["新星载接收机校准", "Stage2-B", "已登记旧终端受控发送", "当前核心；解决接收域变化。"],
                  ["受认证NTN新终端注册", "Stage2-C", "主认证、运营授权和registration episode", "当前主场景；不声称3GPP已定义RFF流程。"],
                  ["TT&C上行RF链管理", "Stage2-B/C", "任务计划、地面站和维护记录", "标签可信度高，类别规模较小。"],
                  ["DCP／远程IoT", "Stage2-C", "地址、频率、发送计划和设备记录", "行政身份需与物理RF链分离。"],
                  ["跨卫星持续一致性", "Stage2-B", "历史登记和多接收域观测", "研究共享身份空间与局部状态。"],
                  ["维修和重新注册", "Stage2-B/C", "维修记录和受控复测", "处理硬件更换与长期漂移。"],
                  ["未知干扰和匿名关联", "Phase3", "定位、多源证据和现场确权", "未来工作；unknown query不能直接训练。"],
                  ["ADS-B／AIS辅助核验", "Phase3／系统融合", "协议、轨迹和独立传感器", "只能作为辅助证据，不直接拒绝安全消息。"],
                  ["卫星下行RFFI", "平行研究／旁证", "真实下行数据集", "方向相反，不能替代上行验证。"],
              ],
              [2350, 1550, 2780, 2680], font_size=8.25)

    add_heading(doc, "6. 为什么需要星上计算", 1, page_break=True)
    add_body(doc, "CVS不需要声称所有训练都必须在天上完成。更准确的定位是：地面完成大规模表征学习，目标卫星执行推理、质量评估和受约束的轻量少样本适配或注册。当地面馈电链路容量充足时，把support特征下传到地面计算delta后再安全上注，应作为ground-assisted baseline。")
    for text in [
        "原始I/Q体量大，不适合持续回传；",
        "异常判断需要早于下一次地面链路；",
        "馈电链路间歇或暂时不可用；",
        "目标接收机的局部状态不宜全部集中上传；",
        "不同卫星需要维护少量receiver-specific状态；",
        "星上预筛选可以减少原始无线数据的长期集中存储。",
    ]:
        add_bullet(doc, text, bullet_id)
    add_body(doc, "ESA的PhiSat-2已展示在轨AI对观测数据进行筛选和压缩，只回传有用结果，从而减少传输负担并加快决策[10]。这只能证明星上边缘处理具有工程合理性，CVS仍需独立报告模型大小、峰值RAM、状态字节、推理时延、更新时延、能耗、异常恢复和回滚能力。")

    add_heading(doc, "7. 统一科学问题与可观测效果", 1)
    research_questions = [
        ("RQ1 跨接收机表征", "在地面标签有限、source receiver多样但目标星载接收机不可见时，能否学习保留发射机差异、降低receiver leakage的表征？"),
        ("RQ2 目标域少样本适配", "在不访问source IQ、source replay和query真值的条件下，K-shot旧类support能否校准残余目标域偏移？"),
        ("RQ3 新RF链注册", "能否利用受认证、受授权的K个独立burst注册新类，同时保持旧类准确率、逐类floor和分数校准？"),
        ("RQ4 可信Enrollment", "主认证成功后，如何防止重放、relay、错误标签和registration window污染？"),
        ("RQ5 长期身份生命周期", "如何区分接收域变化、时间漂移、硬件维修和真正的新物理RF链？"),
        ("RQ6 开放世界扩展", "如何区分真正未知设备与因receiver／channel shift而看似未知的已注册设备？"),
        ("RQ7 星上可部署性", "方法在状态字节、RAM、时延、功耗和故障恢复方面是否优于地面辅助或完全冻结基线？"),
    ]
    for title, body in research_questions:
        add_body(doc, f"{title}：{body}", bold_prefix=title + "：")

    add_heading(doc, "7.1 因果状态命名", 2)
    add_body(doc, "联合研究域适应和新类注册时，应使用四个显式状态，避免以模糊的“前／后”掩盖两种干预：")
    add_table(doc,
              ["状态", "域适应", "新类注册", "可报告内容"],
              [
                  ["DA0_REG0", "前", "前", "旧类直接迁移、资源和时延；新类指标为N/A。"],
                  ["DA1_REG0", "后", "前", "旧类目标域适配效果；新类指标为N/A。"],
                  ["DA0_REG1", "前", "后", "不做DA时的新类注册与旧类保持。"],
                  ["DA1_REG1", "后", "后", "联合适配和注册后的旧／新类表现。"],
              ],
              [1700, 1500, 1500, 4660], font_size=8.8)

    add_heading(doc, "8. 数据证据阶梯与声明边界", 1, page_break=True)
    add_table(doc,
              ["证据层", "用途", "能够证明", "不能证明"],
              [
                  ["WiSig／ManySig地面代理", "跨接收机、跨日期算法研究", "真实receiver/date shift下的方法差异", "真实星载前端或在轨性能"],
                  ["数字LEO启发压力代理", "多普勒、CFO、衰落和低SNR压力测试", "对指定数字扰动的敏感性", "真实H_LEO与R_t联合响应"],
                  ["自建多SDR／目标频段数据", "控制设备、接收机、日期、温度和K-shot", "物理设备区分与变量归因", "飞行级载荷和真实轨道环境"],
                  ["RF硬件在环", "信道模拟器、工程接收链和真实RF传输", "端到端target-like接收性能", "真实在轨温度、辐射和几何"],
                  ["受控在轨采集", "已知终端在授权窗口上行", "真实H_LEO+R_t下的适配与注册", "长期大规模商业可运营性"],
              ],
              [1900, 2400, 2500, 2560], font_size=8.4)
    add_callout(doc, "严格边界", "WiSig／ManySig是真实地面OTA数据，不是卫星数据；对已接收地面I/Q追加数字LEO变换只能称为satellite-inspired link stress proxy，不能称为真实目标卫星数据或在轨验证。")
    add_body(doc, "当前可以声明：研究弱标注跨接收机域泛化、LEO弱信道压力下旧类少样本适配、具有合法K-shot support的新类注册与旧类保持，并为Phase3定义安全交接。当前不能声明：CVS已经完成真实在轨验证、自动未知语义确权、多星协同闭环，或RFFI可以替代密码学认证、定位和现场调查。")

    add_heading(doc, "9. 结论", 1, page_break=True)
    add_body(doc, "CVS最有说服力的研究价值，不是单独证明神经网络能够区分若干发射机，而是把五项现实约束放入同一可审计问题：地面标签有限；目标星载接收机在训练阶段不可见；部署后不能访问source样本；运营网络只能产生少量可信目标域support；物理身份库还会随终端接入、维修和时间持续演化。")
    add_body(doc, "由此形成三个互不替代的科学轴：接收域轴研究如何分离发射机身份与接收机、信道和时间扰动；类别轴研究如何注册新RF链并保持已有身份；开放世界轴研究标签缺失时如何拒识、关联和确权而不污染模型。受控NTN终端注册、TT&C、DCP、跨卫星切换、维修重注册、未知干扰、ADS-B／AIS和卫星下行RFFI等场景可以同时写入项目，但必须明确每个场景的信号方向、标签来源、研究阶段和证据边界。")
    add_callout(doc, "最终表述", "CVS研究的不是“从太空给无线设备贴标签”，而是在大规模星地融合网络中，如何把网络认证与运营登记建立的逻辑身份信任延伸到具体物理RF发射链，并在跨卫星、跨接收机和长期运行条件下持续维护这一绑定。")

    add_heading(doc, "参考资料", 1, page_break=True)
    references = [
        ("[1] 3GPP，Non-Terrestrial Networks (NTN) overview。", "https://www.3gpp.org/technologies/ntn-overview"),
        ("[2] 3GPP，Authentication and Key Management for Applications (AKMA) in 5G。", "https://www.3gpp.org/technologies/akma"),
        ("[3] Iridium，Iridium NTN Direct。", "https://www.iridium.com/services/iridium-ntn-direct"),
        ("[4] Iridium，On-Air Trials Underway: Iridium NTN Direct Prepares to Enter Beta，2026-01-21。", "https://investor.iridium.com/2026-01-21-On-Air-Trials-Underway-Iridium-NTN-Direct-Prepares-to-Enter-Beta-as-Testing-Continues"),
        ("[5] UCLA CORES Lab，WiSig Overview；Hanna et al.，WiSig: A Large-Scale WiFi Signal Dataset for Receiver and Channel Agnostic RF Fingerprinting。", "https://cores.ee.ucla.edu/wisig/overview/"),
        ("[6] ESA，Estrack: ESA's global ground station network。", "https://www.esa.int/Enabling_Support/Operations/ESA_Ground_Stations/Estrack_ESA_s_global_ground_station_network"),
        ("[7] NOAA GOES-R，Concept of Operations，Data Collection System。", "https://goes-r.noaa.gov/syseng/docs/CONOPS.pdf"),
        ("[8] ITU-R SM.2355-2，Spectrum monitoring evolution。", "https://www.itu.int/pub/R-REP-SM.2355-2-2023"),
        ("[9] HawkEye 360，Geolocating RF Interference to the NOAA GOES Satellite System。", "https://www.he360.com/resources/geolocating-rf-interference-to-the-noaa-goes-satellite-system/"),
        ("[10] ESA，Introducing PhiSat-2。", "https://www.esa.int/Applications/Observing_the_Earth/Phsat-2/Introducing_Phsat-2"),
        ("[11] Oligeri et al.，PAST-AI: Physical-Layer Authentication of Satellite Transmitters via Deep Learning。", "https://arxiv.org/abs/2010.05470"),
        ("[12] Saeed et al.，SatIQ: extensible and stable satellite authentication using hardware fingerprinting。", "https://ora.ox.ac.uk/objects/uuid%3A19963898-f543-4595-bc7f-296485c22760"),
    ]
    for label, url in references:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_hyperlink(p, label, url)

    # Core properties and field update hint
    doc.core_properties.title = "CVS天基射频指纹识别：NTN基础、现实机制与分阶段研究场景设计"
    doc.core_properties.subject = "CVS-RFFI／CV-SincNet研究场景与证据边界"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def structural_check(path: Path) -> None:
    assert path.exists() and path.stat().st_size > 50_000
    with ZipFile(path) as zf:
        names = set(zf.namelist())
        assert "word/document.xml" in names
        xml = zf.read("word/document.xml").decode("utf-8")
        for required in (
            "NTN是什么",
            "透明转发载荷",
            "再生式载荷",
            "受认证NTN新终端的物理RF链注册",
            "未知干扰源发现与跨时空匿名关联",
            "参考资料",
        ):
            assert required in xml, required
        assert "turn" not in xml
        assert "cite" not in xml
        assert "w:tblW" in xml and "w:tblGrid" in xml and "w:tcW" in xml
        assert "w:numPr" in xml


if __name__ == "__main__":
    output = build_document()
    structural_check(output)
    print(output)
